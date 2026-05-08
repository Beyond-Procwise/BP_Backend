"""Synthesize realistic DOCX / XLSX / CSV procurement fixtures for the
extraction golden-set. Run once; commits the output under
tests/structural_extractor/fixtures/docs/.

Fixtures:
- invoice_acme.docx     — Invoice with native table of line items
- invoice_megamart.xlsx — Invoice on a single sheet
- po_widgetco.xlsx      — Purchase Order with header block + line-items block
- line_items_bulk.csv   — CSV bulk line-items export
"""
from __future__ import annotations

import io
from pathlib import Path

import openpyxl
from docx import Document

OUT = Path(__file__).resolve().parents[1] / "docs"
OUT.mkdir(exist_ok=True, parents=True)


def build_docx_invoice() -> None:
    d = Document()
    d.add_paragraph("ACME PROFESSIONAL SERVICES LTD")
    d.add_paragraph("42 Innovation Road, Cambridge, CB2 1TN")
    d.add_paragraph("")
    d.add_paragraph("INVOICE")
    d.add_paragraph("")
    d.add_paragraph("Invoice No: ACME-2024-0417")
    d.add_paragraph("Invoice Date: 15 March 2024")
    d.add_paragraph("PO No: PO-77821")
    d.add_paragraph("")
    d.add_paragraph("Bill To:")
    d.add_paragraph("Assurity Ltd")
    d.add_paragraph("10 Redkiln Way, Horsham, West Sussex RH13 5QH")
    d.add_paragraph("")
    t = d.add_table(rows=4, cols=4)
    t.rows[0].cells[0].text = "Description"
    t.rows[0].cells[1].text = "Qty"
    t.rows[0].cells[2].text = "Unit Price"
    t.rows[0].cells[3].text = "Line Total"
    t.rows[1].cells[0].text = "Technical Consulting"
    t.rows[1].cells[1].text = "20"
    t.rows[1].cells[2].text = "150.00"
    t.rows[1].cells[3].text = "3000.00"
    t.rows[2].cells[0].text = "Project Management"
    t.rows[2].cells[1].text = "10"
    t.rows[2].cells[2].text = "180.00"
    t.rows[2].cells[3].text = "1800.00"
    t.rows[3].cells[0].text = "Documentation"
    t.rows[3].cells[1].text = "5"
    t.rows[3].cells[2].text = "100.00"
    t.rows[3].cells[3].text = "500.00"
    d.add_paragraph("")
    d.add_paragraph("Subtotal: 5300.00")
    d.add_paragraph("Tax (20%): 1060.00")
    d.add_paragraph("Total: 6360.00")
    d.add_paragraph("")
    d.add_paragraph("Payment Terms: Net 30")
    d.add_paragraph("Currency: GBP")
    d.save(str(OUT / "invoice_acme.docx"))


def build_xlsx_invoice() -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Invoice"
    # Header block
    ws["A1"] = "MEGAMART SUPPLIES INC"
    ws["A2"] = "Invoice No"
    ws["B2"] = "MM-INV-9921"
    ws["A3"] = "Invoice Date"
    ws["B3"] = "2024-05-10"
    ws["A4"] = "PO No"
    ws["B4"] = "PO-33501"
    ws["A5"] = "Bill To"
    ws["B5"] = "Assurity Ltd"
    ws["A6"] = "Currency"
    ws["B6"] = "USD"
    ws["A7"] = ""
    # Line-item table
    ws["A8"] = "Description"
    ws["B8"] = "Qty"
    ws["C8"] = "Unit Price"
    ws["D8"] = "Line Total"
    ws["A9"] = "Office Chair"
    ws["B9"] = 5
    ws["C9"] = 120.00
    ws["D9"] = 600.00
    ws["A10"] = "Standing Desk"
    ws["B10"] = 5
    ws["C10"] = 340.00
    ws["D10"] = 1700.00
    # Totals
    ws["A12"] = "Subtotal"
    ws["B12"] = 2300.00
    ws["A13"] = "Tax (10%)"
    ws["B13"] = 230.00
    ws["A14"] = "Total"
    ws["B14"] = 2530.00
    ws["A15"] = "Payment Terms"
    ws["B15"] = "Net 14"
    wb.save(str(OUT / "invoice_megamart.xlsx"))


def build_xlsx_po() -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PO"
    ws["A1"] = "PURCHASE ORDER"
    ws["A3"] = "PO No"
    ws["B3"] = "PO-88512"
    ws["A4"] = "PO Date"
    ws["B4"] = "2024-06-20"
    ws["A5"] = "Supplier"
    ws["B5"] = "WidgetCo Ltd"
    ws["A6"] = "Currency"
    ws["B6"] = "EUR"
    ws["A8"] = "Description"
    ws["B8"] = "Qty"
    ws["C8"] = "Unit Price"
    ws["D8"] = "Line Total"
    ws["A9"] = "Widget Type A"
    ws["B9"] = 100
    ws["C9"] = 12.50
    ws["D9"] = 1250.00
    ws["A10"] = "Widget Type B"
    ws["B10"] = 50
    ws["C10"] = 25.00
    ws["D10"] = 1250.00
    ws["A12"] = "Subtotal"
    ws["B12"] = 2500.00
    ws["A13"] = "Tax (19%)"
    ws["B13"] = 475.00
    ws["A14"] = "Total"
    ws["B14"] = 2975.00
    wb.save(str(OUT / "po_widgetco.xlsx"))


def build_csv_line_items() -> None:
    content = (
        "description,quantity,unit_price,line_total\n"
        "Server Rack,2,850.00,1700.00\n"
        "Network Switch,4,240.00,960.00\n"
        "Cable Bundle,10,45.50,455.00\n"
    )
    (OUT / "line_items_bulk.csv").write_text(content)


def main():
    build_docx_invoice()
    build_xlsx_invoice()
    build_xlsx_po()
    build_csv_line_items()
    print(f"Fixtures written to {OUT}")


if __name__ == "__main__":
    main()
