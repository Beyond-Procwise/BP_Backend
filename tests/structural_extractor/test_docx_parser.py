import io
from docx import Document
from src.services.structural_extractor.parsing.docx_parser import parse_docx
from src.services.structural_extractor.parsing.model import NodeRef, Token


def _build_sample_docx() -> bytes:
    d = Document()
    d.add_paragraph("Invoice No: INV-001")
    d.add_paragraph("Invoice Date: 01/07/2022")
    t = d.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Description"
    t.rows[0].cells[1].text = "Qty"
    t.rows[0].cells[2].text = "Amount"
    t.rows[1].cells[0].text = "Widget"
    t.rows[1].cells[1].text = "10"
    t.rows[1].cells[2].text = "100.00"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_parse_docx_paragraphs_have_noderef():
    doc = parse_docx(_build_sample_docx(), "test.docx")
    assert doc.source_format == "docx"
    para_tokens = [t for t in doc.tokens if isinstance(t.anchor, NodeRef) and t.anchor.kind == "paragraph"]
    assert len(para_tokens) >= 2
    assert any(t.text == "INV-001" for t in para_tokens)


def test_parse_docx_table_has_cell_noderef():
    doc = parse_docx(_build_sample_docx(), "test.docx")
    assert len(doc.tables) == 1
    cell_tokens = [t for t in doc.tokens if isinstance(t.anchor, NodeRef) and t.anchor.kind == "table_cell"]
    assert any(t.text == "Widget" for t in cell_tokens)
