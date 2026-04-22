from io import BytesIO
from docx import Document
from src.services.structural_extractor.parsing.model import (
    NodeRef, Token, Region, ParsedDocument, Table
)
from src.services.structural_extractor.exceptions import DocxParseError


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    try:
        d = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise DocxParseError(f"python-docx failed for {filename}: {exc}") from exc

    tokens: list[Token] = []
    regions: list[Region] = []
    tables: list[Table] = []
    full_text_parts: list[str] = []
    order = 0

    for p_idx, para in enumerate(d.paragraphs):
        para_tokens: list[Token] = []
        for word in para.text.split():
            tok = Token(
                text=word,
                anchor=NodeRef(kind="paragraph", paragraph_index=p_idx),
                order=order,
            )
            order += 1
            tokens.append(tok)
            para_tokens.append(tok)
        if para_tokens:
            regions.append(Region(tokens=para_tokens, kind="paragraph"))
            full_text_parts.append(para.text)

    for t_idx, tbl in enumerate(d.tables):
        rows: list[list[Region]] = []
        for r_idx, row in enumerate(tbl.rows):
            row_regions: list[Region] = []
            for c_idx, cell in enumerate(row.cells):
                cell_tokens: list[Token] = []
                for word in cell.text.split():
                    tok = Token(
                        text=word,
                        anchor=NodeRef(
                            kind="table_cell", table_index=t_idx, row=r_idx, col=c_idx
                        ),
                        order=order,
                    )
                    order += 1
                    tokens.append(tok)
                    cell_tokens.append(tok)
                row_regions.append(Region(tokens=cell_tokens, kind="cell"))
            rows.append(row_regions)
            full_text_parts.append(" | ".join(c.text for c in row.cells))
        tables.append(Table(rows=rows, header_row_index=0 if rows else None))

    return ParsedDocument(
        source_format="docx",
        filename=filename,
        tokens=tokens,
        regions=regions,
        tables=tables,
        pages_or_sheets=1,
        full_text="\n".join(full_text_parts),
        raw_bytes=file_bytes,
    )
