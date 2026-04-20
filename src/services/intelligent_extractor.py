"""Intelligent Document Extractor — structure-preserving, self-verifying.

Replaces the flat-text-to-LLM approach with a structured pipeline:

1. PARSE   — Parse document into structured representation (tables, metadata, totals)
             preserving cell-level structure for Excel/DOCX, table detection for PDF
2. EXTRACT — Send structured representation to LLM with schema mapping
3. VERIFY  — Cross-check every extracted value against source structure
4. RETRY   — If verification finds errors, re-extract with targeted guidance

The key insight: don't convert structured data to flat text and hope the LLM
reconstructs it. Preserve the structure and let the LLM map it to our schema.
"""

from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
EXTRACTION_MODEL = os.getenv(
    "PROCWISE_EXTRACTION_MODEL", "BeyondProcwise/AgentNick:extract"
)


class DocumentStructure:
    """Structured representation of a parsed document."""

    def __init__(self):
        self.metadata: List[Dict[str, str]] = []   # [{"label": ..., "value": ...}]
        self.tables: List[Dict[str, Any]] = []      # [{"headers": [...], "rows": [[...]]}]
        self.totals: Dict[str, Any] = {}            # {"Subtotal": 100000, "VAT": 20000}
        self.raw_text: str = ""                     # fallback full text
        self.file_type: str = ""                    # xlsx, pdf, docx, jpeg
        self.notes: List[str] = []                  # additional text (terms, notes)

    def to_prompt_text(self) -> str:
        """Convert to a clear, structured text for LLM consumption."""
        parts = []

        # Metadata section
        if self.metadata:
            parts.append("=== DOCUMENT METADATA ===")
            for entry in self.metadata:
                label = entry.get("label", "")
                value = entry.get("value", "")
                if label and value:
                    parts.append(f"  {label}: {value}")
                elif value:
                    parts.append(f"  {value}")

        # Tables section
        for t_idx, table in enumerate(self.tables):
            headers = table.get("headers", [])
            rows = table.get("rows", [])
            if not rows:
                continue

            parts.append(f"\n=== LINE ITEMS TABLE ===")
            if headers:
                parts.append("| " + " | ".join(str(h) for h in headers) + " |")
                parts.append("|" + "|".join("---" for _ in headers) + "|")

            for row_idx, row in enumerate(rows):
                if headers:
                    # Label each cell with header for clarity
                    labeled = []
                    for i, cell in enumerate(row):
                        hdr = headers[i] if i < len(headers) else f"Col{i+1}"
                        labeled.append(f"{hdr}: {cell}")
                    parts.append(f"  Row {row_idx + 1}: " + " | ".join(labeled))
                else:
                    parts.append(
                        f"  Row {row_idx + 1}: "
                        + " | ".join(str(c) for c in row)
                    )

        # Totals
        if self.totals:
            parts.append("\n=== TOTALS ===")
            for label, value in self.totals.items():
                parts.append(f"  {label}: {value}")

        # Notes
        if self.notes:
            parts.append("\n=== NOTES ===")
            for note in self.notes:
                parts.append(f"  {note}")

        return "\n".join(parts)


class IntelligentExtractor:
    """Structure-preserving, self-verifying document extraction engine."""

    # Column header synonyms for intelligent mapping
    _HEADER_SYNONYMS = {
        "quantity": ["quantity", "qty", "qty.", "no.", "units", "count"],
        "description": [
            "description", "item", "product", "service", "details",
            "particulars", "goods", "item description",
        ],
        "unit_price": [
            "unit price", "price", "rate", "cost", "price per unit",
            "unit cost", "each", "price each", "unit price (£)",
            "unit price (gbp)", "price (£)",
        ],
        "line_total": [
            "total", "amount", "line total", "total price", "ext price",
            "extended", "line amount", "total price (£)",
            "total price (gbp)", "amount (£)", "net",
        ],
        "item_id": [
            "item code", "sku", "part no", "product code", "ref",
            "item #", "code", "cat no", "catalog", "part number",
        ],
        "uom": [
            "uom", "unit", "unit of measure", "measure", "pack",
        ],
    }

    _TOTAL_LABELS = {
        "subtotal", "sub-total", "sub total", "net total",
        "total before tax", "total (ex vat)", "total ex vat",
        "amount before tax",
    }
    _TAX_LABELS = {"vat", "tax", "gst", "sales tax", "vat (20%)", "vat @20%"}
    _GRAND_TOTAL_LABELS = {
        "total", "grand total", "total (gbp)", "total payable",
        "amount due", "total including vat", "total incl vat",
        "total (inc vat)", "invoice total",
    }

    def __init__(self, direct_service, pattern_store=None):
        """Initialize with a DirectExtractionService for text extraction methods."""
        self._service = direct_service
        self._pattern_store = pattern_store

    # ------------------------------------------------------------------
    # Phase 1: PARSE — Structure-preserving document parsing
    # ------------------------------------------------------------------

    def parse_document(self, file_bytes: bytes, ext: str) -> DocumentStructure:
        """Parse any document format into a structured representation."""
        doc = DocumentStructure()
        doc.file_type = ext.lstrip(".")

        if ext in (".xlsx", ".xls"):
            self._parse_excel(file_bytes, doc)
        elif ext == ".csv":
            self._parse_csv(file_bytes, doc)
        elif ext == ".docx":
            self._parse_docx(file_bytes, doc)
        elif ext == ".pdf":
            self._parse_pdf(file_bytes, doc)
        elif ext in (".jpeg", ".jpg", ".png"):
            self._parse_image(file_bytes, doc)
        else:
            doc.raw_text = f"Unsupported format: {ext}"

        return doc

    def _parse_excel(self, file_bytes: bytes, doc: DocumentStructure):
        """Parse Excel into structured metadata + tables + totals."""
        import openpyxl

        try:
            wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
        except Exception:
            logger.warning("Failed to load Excel workbook", exc_info=True)
            doc.raw_text = self._service._tabular_to_text(file_bytes, ".xlsx")
            return

        for sheet in wb.worksheets:
            if not sheet.max_row or sheet.max_row == 0:
                continue

            # Collect all rows as (col_idx, value) pairs
            all_rows = []
            for row in sheet.iter_rows(
                min_row=1, max_row=min(sheet.max_row, 300), values_only=False
            ):
                cells = []
                for c in row:
                    if c.value is not None:
                        val = str(c.value).strip()
                        if val:
                            cells.append((c.column - 1, val))
                all_rows.append(cells)

            # Detect header row
            header_idx = self._find_header_row(all_rows)

            if header_idx is None:
                # No table structure detected — store as raw text
                doc.raw_text = self._service._tabular_to_text(file_bytes, ".xlsx")
                return

            # Parse metadata (rows above header)
            for i in range(header_idx):
                row_cells = all_rows[i]
                if not row_cells:
                    continue
                self._parse_metadata_row(row_cells, doc)

            # Build header map
            header_cells = all_rows[header_idx]
            header_map = {col: text for col, text in header_cells}
            headers = [text for _, text in sorted(header_cells)]

            # Parse line items (rows after header, before totals)
            table_rows = []
            for i in range(header_idx + 1, len(all_rows)):
                row_cells = all_rows[i]
                if not row_cells:
                    continue

                texts = [t for _, t in row_cells]

                # Check if this is a totals row
                if self._is_total_label(texts):
                    # Parse remaining rows as totals
                    for j in range(i, len(all_rows)):
                        rc = all_rows[j]
                        if not rc:
                            continue
                        t = [v for _, v in rc]
                        self._parse_total_row(t, doc)
                    break

                # Check if row has meaningful content (needs description)
                has_desc = self._row_has_description(row_cells, header_map)
                if not has_desc:
                    continue

                # Build row aligned to headers
                row_values = []
                for col_idx, hdr_text in sorted(header_map.items()):
                    cell_val = ""
                    for c_idx, c_val in row_cells:
                        if c_idx == col_idx:
                            cell_val = c_val
                            break
                    row_values.append(cell_val)
                table_rows.append(row_values)

            if table_rows:
                doc.tables.append({"headers": headers, "rows": table_rows})

    def _parse_csv(self, file_bytes: bytes, doc: DocumentStructure):
        """Parse CSV into structured table."""
        import pandas as pd

        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                df = pd.read_csv(BytesIO(file_bytes), encoding=enc)
                break
            except (UnicodeDecodeError, Exception):
                continue
        else:
            return

        headers = [str(c) for c in df.columns]
        rows = []
        for _, row in df.iterrows():
            rows.append([str(v) if pd.notna(v) else "" for v in row.values])
        doc.tables.append({"headers": headers, "rows": rows})

    def _parse_docx(self, file_bytes: bytes, doc: DocumentStructure):
        """Parse DOCX into structured metadata + tables.

        Handles split/watermarked documents by merging all tables into
        one logical structure and deduplicating rows.
        """
        try:
            from docx import Document

            docx_doc = Document(BytesIO(file_bytes))

            # Extract paragraphs as metadata
            for p in docx_doc.paragraphs:
                text = p.text.strip()
                if text:
                    if ":" in text:
                        parts = text.split(":", 1)
                        doc.metadata.append(
                            {"label": parts[0].strip(), "value": parts[1].strip()}
                        )
                    else:
                        doc.metadata.append({"label": "", "value": text})

            # Collect ALL rows from ALL tables with deduplication
            # DOCX split tables (watermarks, page breaks) create multiple
            # table objects for what is logically one table.
            seen_rows: set[str] = set()
            all_data_rows: list[list[str]] = []
            metadata_rows: list[list[str]] = []
            largest_col_count = 0

            for table in docx_doc.tables:
                col_count = len(table.columns)
                if col_count > largest_col_count:
                    largest_col_count = col_count

                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    row_key = " | ".join(cells)

                    # Skip empty rows
                    non_empty = [c for c in cells if c]
                    if not non_empty:
                        continue

                    # Skip duplicate rows
                    if row_key in seen_rows:
                        continue
                    seen_rows.add(row_key)

                    # Classify: totals, metadata (small tables), or data
                    if self._is_total_label(non_empty):
                        self._parse_total_row(non_empty, doc)
                    elif col_count <= 2 and len(non_empty) <= 2:
                        # Small table = likely metadata/header info
                        for cell in non_empty:
                            if "\n" in cell:
                                # Multi-line cell = address or company info
                                doc.metadata.append({"label": "", "value": cell})
                            elif ":" in cell:
                                parts = cell.split(":", 1)
                                doc.metadata.append(
                                    {"label": parts[0].strip(),
                                     "value": parts[1].strip()}
                                )
                            else:
                                doc.metadata.append({"label": "", "value": cell})
                    else:
                        all_data_rows.append(cells)

            # Identify the line items table from collected data rows
            # The line items table has the most columns and numeric data
            if all_data_rows:
                # Find rows that look like line items (have numeric values)
                line_item_rows = []
                header_row = None
                _col_header_words = {
                    "qty", "quantity", "description", "item", "price",
                    "total", "amount", "unit", "product", "service",
                    "no", "ref", "uom",
                }
                for row in all_data_rows:
                    has_number = any(
                        cell.replace(",", "").replace(".", "").replace("£", "")
                        .replace("$", "").strip().isdigit()
                        for cell in row if cell
                    )
                    if has_number:
                        line_item_rows.append(row)
                    elif not header_row:
                        # Check if this row looks like column headers
                        row_words = set()
                        for cell in row:
                            if cell:
                                row_words.update(cell.lower().split())
                        if len(row_words & _col_header_words) >= 2:
                            header_row = row
                        else:
                            # Not a header — treat as metadata
                            non_empty = [c for c in row if c]
                            for cell in non_empty:
                                doc.metadata.append({"label": "", "value": cell})
                    else:
                        non_empty = [c for c in row if c]
                        for cell in non_empty:
                            doc.metadata.append({"label": "", "value": cell})

                if line_item_rows:
                    headers = header_row or []
                    doc.tables.append(
                        {"headers": headers, "rows": line_item_rows}
                    )

            # Parse totals from any remaining numeric-only small cells
            # (e.g., standalone £129,200 values)
            for entry in list(doc.metadata):
                val = entry.get("value", "")
                cleaned = re.sub(r"[£$€,\s]", "", val)
                if cleaned and cleaned.replace(".", "").isdigit():
                    try:
                        num = float(cleaned)
                        if num > 100:  # likely a total, not metadata
                            doc.totals[f"Amount_{len(doc.totals)+1}"] = num
                            doc.metadata.remove(entry)
                    except ValueError:
                        pass

        except Exception:
            logger.debug("python-docx parsing failed", exc_info=True)
            doc.raw_text = self._service._extract_docx_text(file_bytes)

    def _parse_pdf(self, file_bytes: bytes, doc: DocumentStructure):
        """Parse PDF with table extraction + text extraction.

        Uses pdfplumber's extract_tables() for structured data, but
        validates that tables are meaningful (3+ columns, numeric data).
        Falls back to raw text extraction when table detection fails.
        """
        import pdfplumber

        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                all_text_parts = []
                good_tables = []

                for page in pdf.pages:
                    # Extract text (always — used for metadata and fallback)
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        all_text_parts.append(page_text)

                    # Try structured table extraction
                    page_tables = page.extract_tables() or []
                    for table in page_tables:
                        if not table or len(table) < 2:
                            continue
                        # Validate: good tables have 3+ columns and numeric data
                        col_count = max(len(row) for row in table)
                        if col_count < 3:
                            continue
                        has_numbers = any(
                            any(
                                c and re.search(r"\d{2,}", str(c))
                                for c in row
                            )
                            for row in table[1:]
                        )
                        if not has_numbers:
                            continue
                        # This looks like a real data table
                        headers = [
                            str(c).strip() if c else "" for c in table[0]
                        ]
                        rows = []
                        for row in table[1:]:
                            cleaned = [
                                str(c).strip() if c else "" for c in row
                            ]
                            if any(cleaned):
                                rows.append(cleaned)
                        if rows:
                            good_tables.append(
                                {"headers": headers, "rows": rows}
                            )

                # Always store raw text
                if all_text_parts:
                    doc.raw_text = "\n".join(all_text_parts)

                # Use structured tables if they passed validation
                if good_tables:
                    for table in good_tables:
                        item_rows = []
                        for row in table["rows"]:
                            non_empty = [c for c in row if c]
                            if self._is_total_label(non_empty):
                                self._parse_total_row(non_empty, doc)
                            else:
                                item_rows.append(row)
                        if item_rows:
                            doc.tables.append(
                                {"headers": table["headers"], "rows": item_rows}
                            )

                # Parse metadata from text
                if doc.raw_text:
                    for line in doc.raw_text.split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        if ":" in line:
                            parts = line.split(":", 1)
                            label = parts[0].strip()
                            value = parts[1].strip()
                            if label and value and len(label) < 40:
                                doc.metadata.append(
                                    {"label": label, "value": value}
                                )

        except Exception:
            logger.debug("pdfplumber extraction failed", exc_info=True)

        # If no structured tables and no raw text, use full fallback
        if not doc.tables and not doc.raw_text:
            doc.raw_text = self._service._extract_pdf_text(file_bytes)

    def _parse_image(self, file_bytes: bytes, doc: DocumentStructure):
        """Parse image via OCR into structured text."""
        ocr_text = self._service._extract_image_text(file_bytes)
        doc.raw_text = ocr_text

        # Try to detect table structure from OCR text
        lines = ocr_text.split("\n")
        header_idx = None
        header_words = {"quantity", "qty", "description", "price", "total", "amount"}

        for i, line in enumerate(lines):
            words = set(line.lower().split())
            if len(words & header_words) >= 3:
                header_idx = i
                break

        if header_idx is not None:
            # Parse metadata above header
            for i in range(header_idx):
                line = lines[i].strip()
                if line:
                    if ":" in line:
                        parts = line.split(":", 1)
                        doc.metadata.append(
                            {"label": parts[0].strip(), "value": parts[1].strip()}
                        )
                    else:
                        doc.metadata.append({"label": "", "value": line})

    # ------------------------------------------------------------------
    # Phase 2: EXTRACT — Intelligent LLM extraction
    # ------------------------------------------------------------------

    def extract(
        self,
        doc: DocumentStructure,
        doc_type: str,
        schema: Dict[str, Any],
    ) -> Optional[Dict[str, Any]]:
        """Extract structured data from parsed document using LLM.

        Checks for known patterns first to inject proven extraction hints.
        After successful extraction, captures the pattern for future reuse.
        """
        from services.ollama_client import ollama_generate

        # Check for known extraction pattern
        pattern_hints = ""
        matched_pattern = None
        if self._pattern_store:
            matched_pattern = self._pattern_store.find_matching_pattern(
                doc, doc.file_type, doc_type
            )
            if matched_pattern and matched_pattern.get("extraction_hints"):
                pattern_hints = (
                    f"\nPREVIOUS SUCCESSFUL PATTERN:\n"
                    f"{matched_pattern['extraction_hints']}\n"
                )
            if matched_pattern and matched_pattern.get("column_mapping"):
                mapping = matched_pattern["column_mapping"]
                pattern_hints += (
                    f"\nKNOWN COLUMN MAPPING from previous extraction of similar document:\n"
                    + "\n".join(
                        f"  {k} → {v}" for k, v in mapping.items()
                    )
                    + "\n"
                )

        prompt = self._build_intelligent_prompt(
            doc, doc_type, schema, extra_context=pattern_hints
        )

        raw = ollama_generate(
            prompt,
            model=EXTRACTION_MODEL,
            num_predict=4096,
        )
        if not raw:
            logger.error("LLM extraction returned empty response")
            return None

        result = self._parse_response(raw, schema)
        if not result:
            return None

        # Phase 3: VERIFY
        issues = self._verify(result, doc, doc_type)

        # Phase 4: RETRY if critical issues found
        if issues.get("critical_count", 0) > 0:
            logger.info(
                "Verification found %d critical issues — retrying with guidance",
                issues["critical_count"],
            )
            retry_prompt = self._build_retry_prompt(
                doc, doc_type, schema, result, issues
            )
            raw2 = ollama_generate(
                retry_prompt,
                model=EXTRACTION_MODEL,
                num_predict=4096,
            )
            if raw2:
                retry_result = self._parse_response(raw2, schema)
                if retry_result:
                    retry_issues = self._verify(retry_result, doc, doc_type)
                    if retry_issues.get("critical_count", 0) < issues["critical_count"]:
                        logger.info("Retry improved extraction — using retry result")
                        result = retry_result
                        issues = retry_issues

        result["_verification"] = issues

        # Phase 5: LEARN — capture pattern on successful extraction
        if issues.get("critical_count", 0) == 0 and self._pattern_store:
            header = result.get("header", {})
            supplier = (
                header.get("supplier_id")
                or header.get("supplier_name")
                or ""
            )
            # Build column mapping from what the LLM extracted
            col_mapping = {}
            if doc.tables:
                table_headers = doc.tables[0].get("headers", [])
                for h in table_headers:
                    if h:
                        col_mapping[h] = self._guess_schema_field(h)

            hints = (
                f"Supplier: {supplier}\n"
                f"Line items: {len(result.get('line_items', []))}\n"
                f"Header fields: {len(header)}\n"
                f"File type: {doc.file_type}"
            )
            self._pattern_store.capture_pattern(
                doc_structure=doc,
                file_type=doc.file_type,
                doc_type=doc_type,
                supplier_name=supplier,
                column_mapping=col_mapping,
                extraction_hints=hints,
            )

        return result

    @classmethod
    def _guess_schema_field(cls, header_name: str) -> str:
        """Map a document column header to the most likely schema field."""
        norm = header_name.lower().strip()
        for schema_field, synonyms in cls._HEADER_SYNONYMS.items():
            if any(syn in norm for syn in synonyms):
                return schema_field
        return header_name

    def _build_intelligent_prompt(
        self, doc: DocumentStructure, doc_type: str, schema: Dict,
        extra_context: str = "",
    ) -> str:
        """Build extraction prompt using structured document representation."""
        from services.direct_extraction_service import DirectExtractionService

        # Get procurement context
        context = DirectExtractionService._PROCUREMENT_CONTEXT.get(doc_type, "")

        # Filter audit columns
        skip = {
            "created_date", "created_by", "last_modified_date",
            "last_modified_by", "exchange_rate_to_usd", "converted_amount_usd",
        }
        header_cols = [
            c for c in schema["header_columns"] if c not in skip
        ]
        line_cols = [
            c for c in schema.get("line_columns", {}) if c not in skip
        ]

        header_spec = "\n".join(
            f"  - {col} ({schema['header_columns'][col]})"
            for col in header_cols
        )
        line_spec = ""
        if line_cols:
            line_spec = (
                "\nLINE ITEM COLUMNS:\n"
                + "\n".join(
                    f"  - {col} ({schema['line_columns'][col]})"
                    for col in line_cols
                )
            )

        # Build document content — use structured representation when tables exist,
        # otherwise include raw text so the LLM has full document content
        if doc.tables:
            doc_content = doc.to_prompt_text()
        elif doc.raw_text:
            # No structured tables — give LLM the full raw text
            doc_content = doc.raw_text
        else:
            doc_content = doc.to_prompt_text() or ""

        return f"""You are ProcWise, an expert procurement document extraction system.
Extract ALL data from this {doc_type} document with absolute accuracy.

{context}

TARGET SCHEMA — use these EXACT field names in your JSON:
HEADER FIELDS:
{header_spec}
{line_spec}

EXTRACTION RULES:
1. EXTRACT EXACTLY what the document says — never invent or compute values
2. Dates → YYYY-MM-DD format. ONLY extract dates that are explicitly written in the document. Do NOT compute dates from payment terms
3. Amounts → numbers only, strip currency symbols (£1,234.56 → 1234.56)
4. Currency → 3-letter ISO code (GBP, USD, EUR)
5. tax_percent → the percentage NUMBER (20% → 20)
6. Extract EVERY line item. STOP at subtotal/total/tax rows
7. line_total / line_amount: extract the ACTUAL value from the document. Do NOT compute quantity × unit_price
8. quantity: COUNT of items. unit_price: cost PER SINGLE ITEM
9. The SUPPLIER is the company that CREATED/SENT this document (letterhead, logo, top of document)
10. The BUYER is the RECIPIENT company. buyer_id MUST be a COMPANY NAME (e.g., "Assurity Ltd", "Horizon Retail Group Ltd"). NEVER put an address in buyer_id. If "Bill To" shows only a department/address, look for the company name elsewhere
11. supplier_id / supplier_name MUST be a COMPANY NAME. Never put an address in these fields
16. ADDRESSES: Capture full multi-line addresses in the appropriate address fields (supplier_address, buyer_address, delivery_address_line1/2, delivery_city, postal_code). Combine multiple address lines into one value. Addresses and company names go in SEPARATE fields
12. item_id: extract product code/SKU/part number if present, otherwise OMIT
13. unit_of_measure: extract if explicitly stated, otherwise OMIT
14. If duplicate line items exist (from watermarks/split pages), include each item ONLY ONCE
15. If a field is NOT in the document, OMIT it — do not guess or compute it

Return ONLY this JSON:
{{
  "header": {{ ... }},
  "line_items": [ {{ ... }}, ... ]
}}
{extra_context}
DOCUMENT:
{doc_content}"""

    def _build_retry_prompt(
        self,
        doc: DocumentStructure,
        doc_type: str,
        schema: Dict,
        prev_result: Dict,
        issues: Dict,
    ) -> str:
        """Build a retry prompt with specific guidance about what went wrong."""
        base_prompt = self._build_intelligent_prompt(doc, doc_type, schema)

        issue_text = "\n".join(
            f"  - {issue}" for issue in issues.get("messages", [])
        )

        return f"""{base_prompt}

IMPORTANT — Your previous extraction had these errors:
{issue_text}

Please correct these issues. Pay careful attention to:
- Extract line_total / line_amount as written in the document, do NOT compute
- Ensure you capture ALL line items (count them in the source)
- Do not duplicate any line items
- Verify supplier vs buyer identification"""

    # ------------------------------------------------------------------
    # Phase 3: VERIFY — Cross-check extraction against source
    # ------------------------------------------------------------------

    def _verify(
        self, result: Dict, doc: DocumentStructure, doc_type: str
    ) -> Dict[str, Any]:
        """Verify extracted data against source document structure."""
        issues = {"messages": [], "critical_count": 0, "warning_count": 0}

        header = result.get("header", {})
        line_items = result.get("line_items", [])

        # 1. Check line item count against source table
        for table in doc.tables:
            expected_count = len(table.get("rows", []))
            actual_count = len(line_items)
            if expected_count > 0 and actual_count < expected_count:
                issues["messages"].append(
                    f"LINE_COUNT_MISMATCH: Source has {expected_count} rows "
                    f"but only {actual_count} line items extracted"
                )
                issues["critical_count"] += 1

        # 2. Check line item math (qty × price vs total)
        lt_field = "line_amount" if doc_type == "Invoice" else "line_total"
        for i, item in enumerate(line_items):
            qty = self._to_float(item.get("quantity"))
            price = self._to_float(item.get("unit_price"))
            total = self._to_float(item.get(lt_field))
            if qty and price and total:
                expected = round(qty * price, 2)
                if abs(expected - total) > 1.0:
                    issues["messages"].append(
                        f"LINE_MATH: Item {i+1} qty({qty}) × price({price}) "
                        f"= {expected} but total = {total}"
                    )
                    issues["warning_count"] += 1

        # 3. Check line sum vs header subtotal
        subtotal_field = (
            "invoice_amount" if doc_type == "Invoice" else "total_amount"
        )
        header_total = self._to_float(header.get(subtotal_field))
        if header_total and line_items:
            line_sum = sum(
                self._to_float(item.get(lt_field)) or 0 for item in line_items
            )
            if line_sum > 0 and abs(line_sum - header_total) > 1.0:
                diff_pct = abs(line_sum - header_total) / header_total * 100
                severity = "critical" if diff_pct > 5 else "warning"
                issues["messages"].append(
                    f"SUM_MISMATCH: Line items sum = {line_sum} "
                    f"but header {subtotal_field} = {header_total} "
                    f"(diff {diff_pct:.1f}%)"
                )
                if severity == "critical":
                    issues["critical_count"] += 1
                else:
                    issues["warning_count"] += 1

        # 4. Check required fields present
        required = {
            "Invoice": ["invoice_id", "supplier_id"],
            "Purchase_Order": ["po_id", "supplier_name"],
            "Quote": ["quote_id", "supplier_id"],
        }
        for field in required.get(doc_type, []):
            val = header.get(field)
            if not val or (isinstance(val, str) and not val.strip()):
                issues["messages"].append(f"MISSING_FIELD: {field} is empty")
                issues["critical_count"] += 1

        # 5. Value anchoring — check that key values appear in source
        source_text = doc.raw_text or doc.to_prompt_text()
        for field in ["supplier_id", "supplier_name", "buyer_id"]:
            val = header.get(field, "")
            if val and isinstance(val, str) and len(val) > 3:
                # Check if any significant word from the value is in source
                words = [w for w in val.split() if len(w) > 3]
                if words and not any(w in source_text for w in words):
                    issues["messages"].append(
                        f"NOT_ANCHORED: {field}='{val}' — "
                        f"key words not found in source text"
                    )
                    issues["warning_count"] += 1

        return issues

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _parse_response(
        self, raw: str, schema: Dict
    ) -> Optional[Dict[str, Any]]:
        """Parse LLM JSON response."""
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*\n?", "", cleaned)
            cleaned = re.sub(r"\n?```\s*$", "", cleaned)

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            # Find the outermost balanced JSON object
            start = cleaned.find("{")
            if start >= 0:
                depth = 0
                end = start
                for i in range(start, len(cleaned)):
                    if cleaned[i] == "{":
                        depth += 1
                    elif cleaned[i] == "}":
                        depth -= 1
                        if depth == 0:
                            end = i + 1
                            break
                try:
                    data = json.loads(cleaned[start:end])
                except json.JSONDecodeError:
                    logger.error("Failed to parse LLM JSON response")
                    return None
            else:
                return None

        header = data.get("header", {})
        line_items = data.get("line_items", [])

        # Filter to schema columns only
        valid_header = {
            k: v
            for k, v in header.items()
            if k in schema["header_columns"] and v is not None
        }
        line_cols = set(schema.get("line_columns", {}).keys())
        valid_lines = []
        for item in line_items:
            valid = {
                k: v for k, v in item.items() if k in line_cols and v is not None
            }
            if valid:
                valid_lines.append(valid)

        return {"header": valid_header, "line_items": valid_lines}

    def _find_header_row(self, all_rows) -> Optional[int]:
        """Find the table header row by keyword matching."""
        header_words = {
            "quantity", "qty", "description", "item", "product",
            "unit price", "total", "amount", "price", "service", "rate",
        }
        for i, row_cells in enumerate(all_rows):
            texts = [t.lower().strip() for _, t in row_cells]
            matches = sum(
                1 for t in texts if any(kw in t for kw in header_words)
            )
            if matches >= 3:
                return i
        return None

    def _parse_metadata_row(
        self, row_cells: List[Tuple[int, str]], doc: DocumentStructure
    ):
        """Parse a metadata row into key-value pairs."""
        values = [v for _, v in row_cells]

        # Check for label-value pairs (consecutive cells where one is a label)
        i = 0
        while i < len(values):
            val = values[i]
            # Check if this looks like a label (short, ends with common patterns)
            is_label = (
                len(val) < 30
                and not val.replace(",", "").replace(".", "").isdigit()
                and i + 1 < len(values)
            )
            if is_label and ":" in val:
                parts = val.split(":", 1)
                doc.metadata.append(
                    {"label": parts[0].strip(), "value": parts[1].strip()}
                )
                i += 1
            elif is_label and i + 1 < len(values):
                doc.metadata.append({"label": val, "value": values[i + 1]})
                i += 2
            else:
                doc.metadata.append({"label": "", "value": val})
                i += 1

    def _is_total_label(self, texts: List[str]) -> bool:
        """Check if texts contain a totals-section label."""
        if len([t for t in texts if t.strip()]) > 3:
            return False
        all_labels = self._TOTAL_LABELS | self._TAX_LABELS | self._GRAND_TOTAL_LABELS
        for t in texts:
            norm = t.lower().strip()
            if any(norm == lbl or norm.startswith(lbl) for lbl in all_labels):
                return True
        return False

    def _parse_total_row(self, texts: List[str], doc: DocumentStructure):
        """Parse a totals row into the totals dict."""
        if len(texts) >= 2:
            label = texts[0].strip()
            value = texts[-1].strip()
            # Try to parse value as number
            cleaned = re.sub(r"[£$€,\s]", "", value)
            try:
                doc.totals[label] = float(cleaned)
            except ValueError:
                doc.totals[label] = value
        elif len(texts) == 1:
            # Single value — might be just a number
            doc.totals[texts[0]] = texts[0]

    def _row_has_description(
        self,
        row_cells: List[Tuple[int, str]],
        header_map: Dict[int, str],
    ) -> bool:
        """Check if a row has content in a description-like column."""
        desc_keywords = {"description", "item", "product", "service", "details"}
        for col_idx, val in row_cells:
            header = header_map.get(col_idx, "").lower()
            if any(kw in header for kw in desc_keywords):
                return bool(val.strip())
        # If no description column found, accept rows with 3+ cells
        return len(row_cells) >= 3

    @staticmethod
    def _to_float(val) -> Optional[float]:
        if val is None:
            return None
        if isinstance(val, (int, float)):
            return float(val)
        try:
            cleaned = re.sub(r"[£$€,\s]", "", str(val))
            return float(cleaned)
        except (ValueError, TypeError):
            return None
