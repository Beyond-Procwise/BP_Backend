"""Direct extraction service — schema-driven, LLM-powered, accurate.

Extracts document content, sends to AgentNick with the exact bp_ table
schema, and persists directly to the database. No temp files, no
multi-strategy fallbacks, no staging tables.

Flow:
1. Download file from S3
2. Extract text (pdfplumber → OCR fallback → S3 canonical fallback)
3. Send text + exact column definitions to AgentNick LLM
4. Parse JSON response
5. INSERT/UPDATE directly into bp_ tables with audit columns
"""

from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import requests

logger = logging.getLogger(__name__)

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
EXTRACTION_MODEL = os.getenv("PROCWISE_EXTRACTION_MODEL", "BeyondProcwise/AgentNick:extract")

# Exact table schemas from the database — source of truth
TABLE_SCHEMAS = {
    "Invoice": {
        "header_table": "proc.bp_invoice",
        "line_table": "proc.bp_invoice_line_items",
        "pk": "invoice_id",
        "line_pk": "invoice_line_id",
        "line_fk": "invoice_id",
        "line_seq": "line_no",
        "header_columns": {
            "invoice_id": "text", "po_id": "text", "supplier_id": "text",
            "buyer_id": "text", "requisition_id": "text", "requested_by": "text",
            "requested_date": "date", "invoice_date": "date", "due_date": "date",
            "invoice_paid_date": "date", "payment_terms": "text",
            "currency": "varchar", "invoice_amount": "numeric",
            "tax_percent": "numeric", "tax_amount": "numeric",
            "invoice_total_incl_tax": "numeric", "exchange_rate_to_usd": "numeric",
            "converted_amount_usd": "numeric", "country": "text", "region": "text",
            "invoice_status": "text",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
        "line_columns": {
            "invoice_line_id": "text", "invoice_id": "text", "line_no": "integer",
            "item_id": "text", "item_description": "text", "quantity": "integer",
            "unit_of_measure": "text", "unit_price": "numeric",
            "line_amount": "numeric", "tax_percent": "numeric",
            "tax_amount": "numeric", "total_amount_incl_tax": "numeric",
            "po_id": "text", "delivery_date": "date", "country": "text",
            "region": "text",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
    },
    "Purchase_Order": {
        "header_table": "proc.bp_purchase_order",
        "line_table": "proc.bp_po_line_items",
        "pk": "po_id",
        "line_pk": "po_line_id",
        "line_fk": "po_id",
        "line_seq": "line_number",
        "header_columns": {
            "po_id": "text", "supplier_name": "text", "supplier_id": "text",
            "buyer_id": "text",
            "requisition_id": "text", "requested_by": "text",
            "requested_date": "date", "currency": "varchar",
            "order_date": "date", "expected_delivery_date": "date",
            "ship_to_country": "text", "delivery_region": "text",
            "incoterm": "text", "incoterm_responsibility": "text",
            "total_amount": "numeric",
            "tax_percent": "numeric", "tax_amount": "numeric",
            "total_amount_incl_tax": "numeric",
            "delivery_address_line1": "text",
            "delivery_address_line2": "text", "delivery_city": "text",
            "postal_code": "text", "payment_terms": "varchar",
            "po_status": "varchar", "contract_id": "text",
            "exchange_rate_to_usd": "numeric", "converted_amount_usd": "numeric",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
        "line_columns": {
            "po_line_id": "text", "po_id": "text", "line_number": "integer",
            "item_id": "text", "item_description": "text", "quote_number": "text",
            "quantity": "numeric", "unit_price": "numeric",
            "unit_of_measure": "text", "currency": "varchar",
            "line_total": "numeric", "tax_percent": "smallint",
            "tax_amount": "numeric", "total_amount": "numeric",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
    },
    "Quote": {
        "header_table": "proc.bp_quote",
        "line_table": "proc.bp_quote_line_items",
        "pk": "quote_id",
        "line_pk": "quote_line_id",
        "line_fk": "quote_id",
        "line_seq": "line_number",
        "header_columns": {
            "quote_id": "text", "deal_id": "text", "supplier_id": "text",
            "buyer_id": "text", "supplier_address": "text",
            "buyer_address": "text", "quote_date": "date",
            "validity_date": "date", "currency": "varchar",
            "total_amount": "numeric", "tax_percent": "numeric",
            "tax_amount": "numeric", "total_amount_incl_tax": "numeric",
            "po_id": "text", "country": "text", "region": "text",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
        "line_columns": {
            "quote_line_id": "text", "quote_id": "text",
            "line_number": "integer", "item_id": "text",
            "item_description": "text", "quantity": "integer",
            "unit_of_measure": "text", "unit_price": "numeric",
            "line_total": "numeric", "tax_percent": "numeric",
            "tax_amount": "numeric", "total_amount": "numeric",
            "currency": "varchar",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
    },
    "Contract": {
        "header_table": "proc.bp_contracts",
        "line_table": None,
        "pk": "contract_id",
        "line_pk": None,
        "line_fk": None,
        "line_seq": None,
        "header_columns": {
            "contract_id": "text", "contract_title": "text",
            "contract_type": "text", "supplier_id": "text",
            "buyer_org_id": "text", "contract_start_date": "date",
            "contract_end_date": "date", "currency": "text",
            "total_contract_value": "numeric", "spend_category": "text",
            "business_unit_id": "text", "cost_centre_id": "text",
            "is_amendment": "text", "parent_contract_id": "text",
            "auto_renew_flag": "text", "renewal_term": "text",
            "contract_lifecycle_status": "text",
            "created_date": "timestamp", "created_by": "text",
            "last_modified_by": "text", "last_modified_date": "timestamp",
        },
        "line_columns": {},
    },
}

CATEGORY_MAP = {
    "invoice": "Invoice",
    "po": "Purchase_Order",
    "purchase_order": "Purchase_Order",
    "quote": "Quote",
    "quotes": "Quote",
    "contract": "Contract",
    "contracts": "Contract",
}


def _decode_pdf_string(raw: str) -> str:
    """Decode PDF octal escapes in StructTree text (e.g. \\243 → £)."""
    import re

    def _octal_repl(m: re.Match) -> str:
        return chr(int(m.group(1), 8))

    return re.sub(r"\\(\d{3})", _octal_repl, raw)


class DirectExtractionService:
    """Schema-driven extraction: text → LLM → bp_ tables."""

    def __init__(self, agent_nick) -> None:
        self._agent_nick = agent_nick
        self._s3_bucket = agent_nick.settings.s3_bucket_name

    def extract_and_persist(
        self,
        file_path: str,
        category: str,
        *,
        user_id: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Extract document content and persist to bp_ tables.

        Returns dict with status, extracted fields, and any errors.
        """
        doc_type = CATEGORY_MAP.get(category.strip().lower(), "")
        if not doc_type:
            return {"status": "error", "error": f"Unknown category: {category}"}

        schema = TABLE_SCHEMAS.get(doc_type)
        if not schema:
            return {"status": "error", "error": f"No schema for doc_type: {doc_type}"}

        # Step 1: Download and extract text
        text, file_bytes = self._get_document_text(file_path)
        if not text.strip():
            return {
                "status": "error",
                "error": f"No text could be extracted from {file_path}",
                "file_path": file_path,
            }

        logger.info(
            "Extracted %d chars from %s for %s extraction",
            len(text), file_path, doc_type,
        )

        # Step 2: Send to LLM with exact schema
        extraction = self._llm_extract(text, doc_type, schema)
        if not extraction:
            return {
                "status": "error",
                "error": "LLM extraction returned no data",
                "file_path": file_path,
            }

        header = extraction.get("header", {})
        line_items = extraction.get("line_items", [])

        # Step 3: Persist to bp_ tables
        now = datetime.now(timezone.utc)
        audit = {
            "created_date": now,
            "created_by": user_id or "AgentNick",
            "last_modified_by": user_id or "AgentNick",
            "last_modified_date": now,
        }

        pk_value = header.get(schema["pk"], "")
        if not pk_value:
            return {
                "status": "error",
                "error": f"Missing primary key '{schema['pk']}' in extraction",
                "header": header,
            }

        header_result = self._persist_header(header, schema, audit)
        line_result = self._persist_line_items(
            line_items, schema, pk_value, audit
        )

        return {
            "status": "success",
            "file_path": file_path,
            "doc_type": doc_type,
            "pk": pk_value,
            "header_fields": len(header),
            "line_items": len(line_items),
            "header_persisted": header_result,
            "lines_persisted": line_result,
        }

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------
    def _get_document_text(self, file_path: str) -> Tuple[str, bytes]:
        """Download from S3 and extract text. Try canonical fallback paths."""
        ext = os.path.splitext(file_path)[1].lower()
        file_bytes = self._download_s3(file_path)

        # If direct download failed, try canonical S3 paths before giving up
        if not file_bytes:
            file_bytes = self._try_canonical_download(file_path)
            if not file_bytes:
                return "", b""

        # CSV/Excel — convert to text representation
        if ext in (".csv", ".xlsx", ".xls"):
            return self._tabular_to_text(file_bytes, ext), file_bytes

        # PDF/DOCX/Image
        text = self._extract_text_from_bytes(file_bytes, ext)

        # Fallback: try canonical S3 path if text extraction was empty
        if not text.strip():
            alt_text, alt_bytes = self._try_canonical_text(file_path, ext)
            if alt_text:
                return alt_text, alt_bytes

        return text, file_bytes

    def _try_canonical_download(self, file_path: str) -> Optional[bytes]:
        """Try canonical S3 paths when direct download fails.

        Handles both prefixed paths (documents/quote/file.xlsx) and
        bare filenames (file.xlsx) by searching known S3 prefixes.
        """
        canonical_map = {
            "documents/invoice/": ["Invoice/", "PO_Invoice/"],
            "documents/po/": ["Purchase_Order/", "PO_Invoice/"],
            "documents/quote/": ["Quote/", "PO_Invoice/"],
        }

        filename = file_path
        s3_prefixes_to_try: list[str] = []

        # Extract filename and determine which S3 prefixes to search
        for prefix, s3_prefixes in canonical_map.items():
            if file_path.startswith(prefix):
                filename = file_path[len(prefix):]
                s3_prefixes_to_try = s3_prefixes
                break
        else:
            # Bare filename — try all canonical prefixes
            s3_prefixes_to_try = [
                p for prefixes in canonical_map.values() for p in prefixes
            ]

        for s3_prefix in s3_prefixes_to_try:
            alt_bytes = self._download_s3(s3_prefix + filename)
            if alt_bytes:
                logger.info(
                    "Canonical S3 download '%s%s' succeeded",
                    s3_prefix, filename,
                )
                return alt_bytes

        return None

    def _try_canonical_text(self, file_path: str, ext: str) -> Tuple[str, bytes]:
        """Try canonical S3 paths when text extraction produced empty result."""
        canonical_map = {
            "documents/invoice/": ["Invoice/", "PO_Invoice/"],
            "documents/po/": ["Purchase_Order/", "PO_Invoice/"],
            "documents/quote/": ["Quote/", "PO_Invoice/"],
        }

        filename = file_path
        s3_prefixes_to_try: list[str] = []

        for prefix, s3_prefixes in canonical_map.items():
            if file_path.startswith(prefix):
                filename = file_path[len(prefix):]
                s3_prefixes_to_try = s3_prefixes
                break
        else:
            s3_prefixes_to_try = [
                p for prefixes in canonical_map.values() for p in prefixes
            ]

        for s3_prefix in s3_prefixes_to_try:
            alt_bytes = self._download_s3(s3_prefix + filename)
            if alt_bytes:
                alt_text = self._extract_text_from_bytes(alt_bytes, ext)
                if alt_text.strip():
                    logger.info(
                        "Canonical fallback '%s%s' succeeded (%d chars)",
                        s3_prefix, filename, len(alt_text),
                    )
                    return alt_text, alt_bytes

            # Try fuzzy match
            alt_text, alt_bytes = self._search_s3_fuzzy(
                s3_prefix, filename, ext
            )
            if alt_text:
                return alt_text, alt_bytes

        return "", b""

    def _search_s3_fuzzy(
        self, s3_prefix: str, filename: str, ext: str
    ) -> tuple[str, bytes]:
        """Search S3 prefix for a file matching the key identifiers in filename."""
        # Extract key identifiers: supplier name and document number
        # e.g., "GOMEZ, GOOD ETC PO507269 for QUT104683 .pdf" → search for "GOMEZ" + "PO507269"
        name_part = os.path.splitext(filename)[0].strip()
        # Extract the document ID (PO/INV/QUT number)
        import re
        doc_id_match = re.search(
            r"((?:PO|INV|QUT|INV-)\s*[\d\-]+)", name_part, re.IGNORECASE
        )
        doc_id = doc_id_match.group(1).replace(" ", "") if doc_id_match else ""
        # Extract supplier name (first word group before the doc ID)
        supplier_hint = name_part.split()[0] if name_part.split() else ""

        if not doc_id and not supplier_hint:
            return "", b""

        try:
            s3 = self._agent_nick.reserve_s3_connection().__enter__()
            result = s3.list_objects_v2(
                Bucket=self._s3_bucket, Prefix=s3_prefix, MaxKeys=200
            )
            for obj in result.get("Contents", []):
                key = obj["Key"]
                key_upper = key.upper()
                # Match by document ID (most specific)
                if doc_id and doc_id.upper() in key_upper:
                    alt_bytes = self._download_s3(key)
                    if alt_bytes:
                        alt_text = self._extract_text_from_bytes(alt_bytes, ext)
                        if alt_text.strip():
                            logger.info(
                                "Fuzzy S3 fallback matched '%s' via doc_id '%s' (%d chars)",
                                key, doc_id, len(alt_text),
                            )
                            return alt_text, alt_bytes
                # Match by supplier name
                elif supplier_hint and supplier_hint.upper() in key_upper:
                    alt_bytes = self._download_s3(key)
                    if alt_bytes:
                        alt_text = self._extract_text_from_bytes(alt_bytes, ext)
                        if alt_text.strip():
                            logger.info(
                                "Fuzzy S3 fallback matched '%s' via supplier '%s' (%d chars)",
                                key, supplier_hint, len(alt_text),
                            )
                            return alt_text, alt_bytes
        except Exception:
            logger.debug("Fuzzy S3 search failed for %s%s", s3_prefix, filename)

        return "", b""

    def _download_s3(self, key: str) -> Optional[bytes]:
        """Download a file from S3."""
        try:
            s3 = self._agent_nick.reserve_s3_connection().__enter__()
            obj = s3.get_object(Bucket=self._s3_bucket, Key=key)
            return obj["Body"].read()
        except Exception:
            logger.debug("S3 download failed for '%s'", key)
            return None

    def _extract_text_from_bytes(self, file_bytes: bytes, ext: str) -> str:
        """Extract text from PDF/DOCX/image bytes."""
        if ext == ".pdf":
            return self._extract_pdf_text(file_bytes)
        elif ext == ".docx":
            return self._extract_docx_text(file_bytes)
        elif ext in (".png", ".jpg", ".jpeg"):
            return self._extract_image_text(file_bytes)
        return ""

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pdfplumber, with OCR and StructTree fallbacks."""
        import pdfplumber
        text_parts: list[str] = []
        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
        except Exception:
            logger.debug("pdfplumber failed", exc_info=True)

        if text_parts:
            return "\n".join(text_parts)

        # OCR fallback for scanned PDFs — multi-engine with preprocessing
        try:
            from pdf2image import convert_from_bytes
            import pytesseract
            images = convert_from_bytes(file_bytes, dpi=300)

            for img in images:
                # Try Tesseract first on raw image
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    text_parts.append(ocr_text)
                    continue

                # Tesseract failed — try with image preprocessing
                try:
                    import cv2
                    import numpy as np
                    img_array = np.array(img)
                    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                    # Increase contrast with CLAHE
                    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                    gray = clahe.apply(gray)
                    # Denoise
                    gray = cv2.fastNlMeansDenoising(gray, h=10)
                    # Adaptive threshold for binarization
                    binary = cv2.adaptiveThreshold(
                        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                        cv2.THRESH_BINARY, 31, 10,
                    )
                    from PIL import Image
                    preprocessed = Image.fromarray(binary)
                    ocr_text = pytesseract.image_to_string(preprocessed)
                    if ocr_text.strip():
                        text_parts.append(ocr_text)
                        continue
                except Exception:
                    logger.debug("Preprocessed Tesseract failed", exc_info=True)

                # Both Tesseract attempts failed — try easyOCR (GPU)
                try:
                    import easyocr
                    import numpy as np
                    reader = easyocr.Reader(["en"], gpu=True, verbose=False)
                    img_array = np.array(img)
                    results = reader.readtext(img_array, detail=0, paragraph=True)
                    if results:
                        ocr_text = "\n".join(results)
                        text_parts.append(ocr_text)
                        logger.info("easyOCR GPU recovered %d text blocks", len(results))
                except Exception:
                    logger.debug("easyOCR fallback failed", exc_info=True)

            if not text_parts:
                logger.warning(
                    "All OCR engines produced no text (%d pages converted)",
                    len(images),
                )
        except Exception:
            logger.warning("OCR fallback failed for scanned PDF", exc_info=True)

        if text_parts:
            return "\n".join(text_parts)

        # StructTree fallback for corrupted PDFs (e.g. Canva-generated with
        # broken zlib content streams).  Tagged PDFs store text in /T and /E
        # attributes of /StructElem objects even when the page stream is
        # unreadable.
        #
        # Key improvement: reconstruct table rows by detecting H3 column
        # headers then grouping subsequent P elements into rows matching
        # the header column count.
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")

            elements: list[tuple[str, str]] = []  # (struct_type, text)
            for i in range(1, doc.xref_length()):
                obj_str = doc.xref_object(i)
                if "/StructElem" not in obj_str:
                    continue

                # Get structure type (/S)
                s_type = ""
                s_start = obj_str.find("/S /")
                if s_start != -1:
                    s_start += 4
                    s_end = obj_str.find("\n", s_start)
                    s_type = obj_str[s_start:s_end].strip()

                # Get text (/T or /E)
                text = ""
                for tag in ("/T (", "/E ("):
                    start = obj_str.find(tag)
                    if start == -1:
                        continue
                    start += len(tag)
                    end = obj_str.find(")", start)
                    if end != -1:
                        text = _decode_pdf_string(obj_str[start:end].strip())
                        break

                if text:
                    elements.append((s_type, text))

            doc.close()

            if not elements:
                return ""

            # Reconstruct structured text with table awareness.
            # Strategy: detect table header keywords, then group subsequent
            # P elements into rows. Handles both H3-header and P-header
            # patterns, and multi-line descriptions.
            import re as _re

            _TABLE_HEADERS = {
                "description", "item", "item description", "items",
                "product", "product description", "service", "services",
                "qty", "quantity", "qty.", "no.", "no",
                "price", "rate", "unit cost",
                "unit price", "unit price (£)", "unit price ($)",
                "subtotal", "sub total", "sub-total",
                "total", "total (£)", "total ($)", "total price",
                "amount", "net amount", "gross amount",
                "line total", "line total (£)", "line amount",
                "cost", "ext. price", "extended price",
                "total price", "total price (£)",
                "discount", "discount %", "vat", "vat %", "tax",
                "uom", "unit", "units",
            }
            _SUMMARY_STARTS = (
                "sub-total", "subtotal", "tax", "total", "vat",
                "payable", "bank", "account", "notes", "payment",
                "delivery", "grand total", "net total", "amount due",
                "balance", "remittance",
            )
            _MONEY_RE = _re.compile(r"^[£$€¥]?\s*[\d,]+\.?\d*$")
            _NUMBER_RE = _re.compile(r"^\d+\.?\d*$")

            output_lines: list[str] = []
            i = 0
            while i < len(elements):
                s_type, text = elements[i]
                text_lower = text.lower().strip()

                # Detect table header: element whose text is a known header keyword
                if text_lower in _TABLE_HEADERS:
                    headers = [text]
                    j = i + 1
                    while j < len(elements) and elements[j][1].lower().strip() in _TABLE_HEADERS:
                        headers.append(elements[j][1])
                        j += 1

                    if len(headers) >= 2:
                        output_lines.append(" | ".join(headers))
                        col_count = len(headers)

                        # Allowed struct types in table body
                        _TABLE_BODY_TYPES = {"P", "H5", "H6", "Span", "TD", "LBody"}

                        # Group subsequent elements into table rows.
                        # A cell is either a money value (£/$X), a number, or
                        # a text description. Multi-line descriptions (2 P
                        # elements like "Staedtler Pen" + "Black Ink") are
                        # merged if neither looks like a number/money.
                        while j < len(elements):
                            _, next_text = elements[j]
                            if next_text.lower().strip().startswith(_SUMMARY_STARTS):
                                break
                            if elements[j][0] not in _TABLE_BODY_TYPES:
                                break

                            row_cells: list[str] = []
                            while len(row_cells) < col_count and j < len(elements):
                                _, cell = elements[j]
                                if cell.lower().strip().startswith(_SUMMARY_STARTS):
                                    break
                                if elements[j][0] not in _TABLE_BODY_TYPES:
                                    break

                                # Merge multi-line text descriptions
                                if (row_cells
                                    and not _MONEY_RE.match(cell)
                                    and not _NUMBER_RE.match(cell.strip())
                                    and not _MONEY_RE.match(row_cells[-1])
                                    and not _NUMBER_RE.match(row_cells[-1].strip())):
                                    row_cells[-1] += " " + cell
                                else:
                                    row_cells.append(cell)
                                j += 1

                            if row_cells:
                                output_lines.append(" | ".join(row_cells))
                            # Allow partial rows (don't break on incomplete rows)
                            if len(row_cells) == 0:
                                break

                        i = j
                        continue

                output_lines.append(text)
                i += 1

            if output_lines:
                logger.info(
                    "StructTree fallback recovered %d lines (table-aware) from corrupted PDF",
                    len(output_lines),
                )
                return "\n".join(output_lines)
        except Exception:
            logger.warning("StructTree fallback failed", exc_info=True)

        return ""

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        """Extract text from DOCX including paragraphs AND tables.

        Falls back to raw XML parsing for complex/watermarked documents
        that python-docx cannot handle.
        """
        # Primary: python-docx
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            parts: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        line = " | ".join(cells)
                        parts.append(line)
                        # After the first row (header), add a separator
                        if row_idx == 0:
                            parts.append("--- | " * (len(cells) - 1) + "---")
            text = "\n".join(parts)
            if text.strip():
                return text
        except Exception:
            logger.debug("python-docx extraction failed", exc_info=True)

        # Fallback: raw XML extraction from DOCX zip
        # Handles watermarked, complex, or partially corrupted DOCX files
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            zf = zipfile.ZipFile(BytesIO(file_bytes))
            parts: list[str] = []

            # Extract from word/document.xml (main content)
            for xml_path in ["word/document.xml", "word/document2.xml"]:
                if xml_path in zf.namelist():
                    xml_data = zf.read(xml_path)
                    root = ET.fromstring(xml_data)
                    # Strip namespace prefixes for easier parsing
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for para in root.findall(".//w:p", ns):
                        texts = [t.text for t in para.findall(".//w:t", ns) if t.text]
                        if texts:
                            parts.append("".join(texts))

            # Also extract from tables
            for xml_path in zf.namelist():
                if "word/" in xml_path and xml_path.endswith(".xml"):
                    try:
                        xml_data = zf.read(xml_path)
                        root = ET.fromstring(xml_data)
                        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                        for tbl in root.findall(".//w:tbl", ns):
                            for tr in tbl.findall(".//w:tr", ns):
                                cells = []
                                for tc in tr.findall(".//w:tc", ns):
                                    cell_texts = [t.text for t in tc.findall(".//w:t", ns) if t.text]
                                    if cell_texts:
                                        cells.append(" ".join(cell_texts))
                                if cells:
                                    parts.append(" | ".join(cells))
                    except Exception:
                        continue

            zf.close()
            text = "\n".join(parts)
            if text.strip():
                logger.info("DOCX raw XML fallback recovered %d lines", len(parts))
                return text
        except Exception:
            logger.debug("DOCX raw XML fallback failed", exc_info=True)

        return ""

    def _extract_image_text(self, file_bytes: bytes) -> str:
        """Extract text from image with comprehensive OCR pipeline.

        Procurement-optimized: handles scanned invoices, photographed
        receipts, low-quality images. Preserves table structure.

        Pipeline:
        1. Image preprocessing (deskew, upscale, contrast, denoise)
        2. Tesseract with multiple PSM modes for best result
        3. easyOCR GPU with layout detection as fallback
        4. Merge best result
        """
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            import pytesseract
            import cv2
            import numpy as np

            img = Image.open(BytesIO(file_bytes)).convert("RGB")
            img_array = np.array(img)

            # Step 1: Image preprocessing
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)

            # Upscale small images (phone photos, thumbnails)
            h, w = gray.shape
            if max(h, w) < 1500:
                scale = 2.0
                gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
                logger.info("Image upscaled %.0fx (was %dx%d)", scale, w, h)

            # Deskew: detect and correct rotation
            try:
                coords = np.column_stack(np.where(gray < 128))
                if len(coords) > 100:
                    angle = cv2.minAreaRect(coords)[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle
                    if abs(angle) > 0.5 and abs(angle) < 15:
                        center = (gray.shape[1] // 2, gray.shape[0] // 2)
                        M = cv2.getRotationMatrix2D(center, angle, 1.0)
                        gray = cv2.warpAffine(
                            gray, M, (gray.shape[1], gray.shape[0]),
                            flags=cv2.INTER_CUBIC,
                            borderMode=cv2.BORDER_REPLICATE,
                        )
                        logger.info("Image deskewed by %.1f degrees", angle)
            except Exception:
                pass  # Deskew is best-effort

            # Contrast enhancement (CLAHE)
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)

            # Denoise
            denoised = cv2.fastNlMeansDenoising(enhanced, h=12)

            # Step 2: Try multiple Tesseract modes and pick the best
            results = []

            # Mode A: Adaptive threshold + Tesseract PSM 6 (block of text)
            binary_adaptive = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY, 31, 10,
            )
            try:
                pil_adaptive = Image.fromarray(binary_adaptive)
                text_a = pytesseract.image_to_string(
                    pil_adaptive, config="--psm 6 --oem 3"
                )
                if text_a.strip():
                    results.append(("tesseract_psm6", text_a.strip()))
            except Exception:
                pass

            # Mode B: Otsu threshold + Tesseract PSM 3 (auto page segmentation)
            _, binary_otsu = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            try:
                pil_otsu = Image.fromarray(binary_otsu)
                text_b = pytesseract.image_to_string(
                    pil_otsu, config="--psm 3 --oem 3"
                )
                if text_b.strip():
                    results.append(("tesseract_psm3", text_b.strip()))
            except Exception:
                pass

            # Mode C: Enhanced grayscale directly (no binarization — preserves more detail)
            try:
                pil_gray = Image.fromarray(denoised)
                text_c = pytesseract.image_to_string(
                    pil_gray, config="--psm 4 --oem 3"
                )
                if text_c.strip():
                    results.append(("tesseract_psm4", text_c.strip()))
            except Exception:
                pass

            # Step 3: easyOCR with layout detection (GPU)
            try:
                import easyocr
                reader = easyocr.Reader(["en"], gpu=True, verbose=False)
                # Use detail mode to get bounding boxes for table reconstruction
                ocr_results = reader.readtext(img_array, detail=1, paragraph=False)
                if ocr_results:
                    # Sort by vertical position then horizontal for reading order
                    ocr_results.sort(key=lambda r: (r[0][0][1], r[0][0][0]))

                    # Group into lines by Y proximity
                    lines = []
                    current_line = []
                    last_y = -100
                    for bbox, text_val, conf in ocr_results:
                        y = bbox[0][1]
                        if abs(y - last_y) > 15 and current_line:
                            lines.append(" ".join(current_line))
                            current_line = []
                        current_line.append(text_val)
                        last_y = y
                    if current_line:
                        lines.append(" ".join(current_line))

                    easy_text = "\n".join(lines)
                    if easy_text.strip():
                        results.append(("easyocr_gpu", easy_text.strip()))
                        logger.info("easyOCR extracted %d lines from image", len(lines))
            except Exception:
                logger.debug("easyOCR failed for image", exc_info=True)

            # Step 4: Pick the best result (most content with readable structure)
            if not results:
                logger.warning("All OCR engines produced no text from image")
                return ""

            # Score each result: prefer more lines, more alphanumeric chars, fewer garbage chars
            def _score(text: str) -> float:
                lines = len(text.split("\n"))
                alnum = sum(1 for c in text if c.isalnum())
                total = len(text) or 1
                alnum_ratio = alnum / total
                return lines * alnum_ratio * len(text)

            best_method, best_text = max(results, key=lambda r: _score(r[1]))
            logger.info(
                "Image OCR: best method=%s (%d chars, %d lines)",
                best_method, len(best_text), len(best_text.split("\n")),
            )
            return best_text

        except Exception:
            logger.warning("Image text extraction failed", exc_info=True)
        return ""

    def _tabular_to_text(self, file_bytes: bytes, ext: str) -> str:
        """Convert CSV/Excel to structured text for LLM processing.

        For Excel files, reads ALL cells including metadata rows above
        the main table (quote ID, supplier, dates etc.) to give the LLM
        full context.
        """
        try:
            if ext == ".csv":
                for enc in ("utf-8", "latin-1", "cp1252"):
                    try:
                        df = pd.read_csv(BytesIO(file_bytes), encoding=enc)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    return ""
                header = " | ".join(str(c) for c in df.columns)
                rows = []
                for _, row in df.iterrows():
                    rows.append(" | ".join(
                        str(v) for v in row.values if pd.notna(v)
                    ))
                return f"COLUMNS: {header}\n\n" + "\n".join(rows)

            # Excel: read raw cells to capture metadata above the table
            import openpyxl
            try:
                wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
            except Exception:
                # Try ZIP repair for corrupted xlsx files
                logger.warning("openpyxl failed, attempting ZIP repair for xlsx")
                try:
                    import struct, zipfile
                    cd_pos = file_bytes.find(b'PK\x01\x02')
                    eocd_pos = file_bytes.rfind(b'PK\x05\x06')
                    if cd_pos > 0 and eocd_pos > cd_pos:
                        fixed = bytearray(file_bytes)
                        struct.pack_into('<I', fixed, eocd_pos + 16, cd_pos)
                        struct.pack_into('<I', fixed, eocd_pos + 12, eocd_pos - cd_pos)
                        wb = openpyxl.load_workbook(BytesIO(bytes(fixed)), data_only=True)
                        logger.info("ZIP repair succeeded for corrupted xlsx")
                    else:
                        raise ValueError("No valid ZIP central directory found")
                except Exception:
                    # Try pandas as last resort (handles some corrupt formats)
                    try:
                        df = pd.read_excel(BytesIO(file_bytes), engine="openpyxl")
                        if not df.empty:
                            header = " | ".join(str(c) for c in df.columns)
                            rows = [" | ".join(str(v) for v in row if pd.notna(v))
                                    for _, row in df.iterrows()]
                            return f"COLUMNS: {header}\n\n" + "\n".join(rows)
                    except Exception:
                        pass
                    logger.warning("Excel file is corrupted and cannot be repaired")
                    return ""
            return self._excel_to_structured_text(wb)
        except Exception:
            logger.warning("Tabular text extraction failed", exc_info=True)
            return ""

    # ------------------------------------------------------------------
    # Excel structure detection
    # ------------------------------------------------------------------

    _HEADER_KEYWORDS = frozenset({
        "quantity", "qty", "description", "item", "product", "unit price",
        "total", "amount", "price", "uom", "unit of measure", "part",
        "sku", "code", "no", "ref", "service", "rate",
    })
    _TOTAL_KEYWORDS = frozenset({
        "subtotal", "sub-total", "sub total", "total", "vat", "tax",
        "grand total", "net", "gross", "discount",
    })

    @classmethod
    def _is_header_row(cls, cells_text: list[str]) -> bool:
        """True when 3+ cells look like column headers."""
        matches = 0
        for ct in cells_text:
            norm = ct.lower().strip()
            if any(kw in norm for kw in cls._HEADER_KEYWORDS):
                matches += 1
        return matches >= 3

    @classmethod
    def _is_total_row(cls, cells_text: list[str]) -> bool:
        """True when a row is a summary/total row, not a line item.

        A total row has 1-2 non-empty cells where one is a label like
        "Subtotal" or "VAT (20%)" and the other is a number.  Line item
        rows typically have 3+ cells (qty, description, price, total).

        The keyword must appear as a standalone label — not embedded
        inside a longer item description like "Delivery & Installation
        Services".
        """
        non_empty = [ct for ct in cells_text if ct.strip()]
        if len(non_empty) > 2:
            return False
        for ct in non_empty:
            norm = ct.lower().strip()
            if any(kw == norm or norm.startswith(kw) for kw in cls._TOTAL_KEYWORDS):
                return True
        return False

    def _excel_to_structured_text(self, wb) -> str:
        """Convert openpyxl workbook to structured text with metadata/table separation.

        Detects the header row that defines line-item columns, then:
        - Rows above header → DOCUMENT METADATA section
        - Header row + data rows → LINE ITEMS TABLE with column-labeled values
        - Subtotal/Total/VAT rows → TOTALS section
        """
        parts: list[str] = []

        for sheet in wb.worksheets:
            if not sheet.max_row or sheet.max_row == 0:
                continue

            if len(wb.worksheets) > 1:
                parts.append(f"--- Sheet: {sheet.title} ---")

            # Collect all rows with their cell values and column positions
            all_rows: list[list[tuple[int, str]]] = []  # [(col_idx, text), ...]
            for row in sheet.iter_rows(
                min_row=1, max_row=min(sheet.max_row, 200), values_only=False
            ):
                cells = []
                for c in row:
                    if c.value is not None:
                        val = str(c.value).strip()
                        if val:
                            cells.append((c.column - 1, val))  # 0-indexed
                all_rows.append(cells)

            # Find header row
            header_row_idx = None
            header_map: dict[int, str] = {}  # col_idx -> header name
            for i, row_cells in enumerate(all_rows):
                texts = [t for _, t in row_cells]
                if self._is_header_row(texts):
                    header_row_idx = i
                    header_map = {col: text for col, text in row_cells}
                    break

            if header_row_idx is None:
                # No header detected — fall back to raw pipe-delimited output
                for row_cells in all_rows:
                    if row_cells:
                        parts.append(" | ".join(t for _, t in row_cells))
                continue

            # === DOCUMENT METADATA ===
            parts.append("\n=== DOCUMENT METADATA ===")
            for i in range(header_row_idx):
                row_cells = all_rows[i]
                if not row_cells:
                    continue
                parts.append(" | ".join(t for _, t in row_cells))

            # === LINE ITEMS TABLE ===
            parts.append("\n=== LINE ITEMS TABLE ===")
            in_totals = False

            # Identify description column indices for phantom-row filtering
            _DESC_HINTS = {"description", "item", "product", "service"}
            desc_col_idxs = {
                col for col, hdr in header_map.items()
                if any(h in hdr.lower() for h in _DESC_HINTS)
            }

            for i in range(header_row_idx + 1, len(all_rows)):
                row_cells = all_rows[i]
                if not row_cells:
                    continue

                texts = [t for _, t in row_cells]

                # Check if we've hit totals section
                if not in_totals and self._is_total_row(texts):
                    in_totals = True
                    parts.append("\n--- TOTALS ---")

                if in_totals:
                    parts.append(" | ".join(texts))
                else:
                    # Filter phantom rows: if we know which columns are
                    # descriptions, skip rows that have no text in any
                    # description column (e.g. rows with only "0" values).
                    if desc_col_idxs:
                        has_desc = any(
                            col_idx in desc_col_idxs and val.strip()
                            for col_idx, val in row_cells
                        )
                        if not has_desc:
                            continue

                    # Label each cell with its column header
                    labeled = []
                    for col_idx, val in row_cells:
                        hdr = header_map.get(col_idx, "")
                        if hdr:
                            labeled.append(f"{hdr}: {val}")
                        else:
                            labeled.append(val)
                    if labeled:
                        parts.append(" | ".join(labeled))

        return "\n".join(parts) if parts else ""

    # ------------------------------------------------------------------
    # LLM extraction
    # ------------------------------------------------------------------
    def _llm_extract(
        self, text: str, doc_type: str, schema: Dict
    ) -> Optional[Dict[str, Any]]:
        """Send document text + exact schema to AgentNick for extraction."""
        header_cols = list(schema["header_columns"].keys())
        line_cols = list(schema.get("line_columns", {}).keys())

        prompt = self._build_extraction_prompt(
            text, doc_type, header_cols, line_cols, schema
        )

        try:
            from services.ollama_client import ollama_generate
            raw = ollama_generate(
                prompt,
                model=EXTRACTION_MODEL,
                num_predict=4096,
            )
            if not raw:
                logger.error("LLM extraction returned empty response")
                return None
            return self._parse_llm_response(raw, schema)
        except Exception as exc:
            logger.exception("LLM extraction failed: %s", exc)
            return None

    # Procurement domain context per document type
    _PROCUREMENT_CONTEXT = {
        "Invoice": """PROCUREMENT CONTEXT — INVOICE:
An invoice is a payment request from a SUPPLIER to a BUYER for goods/services delivered.
- supplier_id / supplier_name: The company that ISSUED this invoice (the SELLER). Look for "From", "Vendor", "Supplier", company letterhead, or the name at the TOP of the document. This is NOT the "Bill To" or "Ship To" company.
- buyer_id: The company being BILLED (the purchaser). Look for "Bill To", "Invoice To", "Customer".
- invoice_id: The unique invoice reference number. Look for "Invoice No", "Invoice #", "Inv No", "Reference". This is the document's own ID, not a PO or order number.
- po_id: The Purchase Order this invoice relates to. Look for "PO Number", "Order Ref", "Your Ref".
- invoice_date: When the invoice was ISSUED. Look for "Invoice Date", "Date", "Date of Issue".
- due_date: When payment is DUE. Look for "Due Date", "Payment Due", "Terms Date". Must be AFTER invoice_date.
- invoice_amount: The SUBTOTAL before tax (net amount). Look for "Subtotal", "Net Total", "Amount Before Tax".
- tax_amount: The tax/VAT amount. Look for "VAT", "Tax", "GST". Must be LESS than the subtotal.
- tax_percent: The tax RATE as a percentage (e.g., 20 for 20% VAT). Look for "VAT @20%", "Tax Rate". If the document says "20%" the value is 20, not 0.20. Do NOT confuse line item quantities or reference numbers with tax percentage.
- invoice_total_incl_tax: The FINAL total including tax. Look for "Total", "Amount Due", "Grand Total", "Total Payable". This is the largest amount on the invoice.
- payment_terms: e.g. "Net 30", "Due on receipt". Look for "Payment Terms", "Terms".""",

        "Purchase_Order": """PROCUREMENT CONTEXT — PURCHASE ORDER:
A PO is issued BY the buyer TO a supplier, authorizing purchase of goods/services.
- po_id: The PO number. Look for "PO Number", "Purchase Order No", "Order #".
- supplier_name: The company RECEIVING the order — the VENDOR who will FULFIL it. Look for "To", "Vendor", "Supplier", "Ship From", "Deliver From". This is the company the goods/services are being ordered FROM.
- buyer_id: The company that CREATED/ISSUED the PO — the one PLACING the order. Look for letterhead, "From", "Issued By", "Buyer", "Ordered By", "Ship To", "Bill To". This is the company PAYING for the goods/services.
- order_date: When the PO was issued.
- expected_delivery_date: When goods/services are expected.
- total_amount: The SUBTOTAL before tax — sum of all line items BEFORE VAT/tax is added. Look for "Subtotal", "Net Total", "Total Before Tax". This is NOT the grand total including VAT.
- tax_amount: The VAT/tax amount added on top of the subtotal.
- tax_percent: The VAT/tax rate as a number (e.g. 20 for 20%). Extract from labels like "VAT @20%", "Tax 20%".
- total_amount_incl_tax: The GRAND TOTAL including tax. Look for "Total", "Grand Total", "Total Payable", "Total (GBP)".
- payment_terms: e.g. "Net 30", "Due on receipt".
CRITICAL: total_amount is the SUBTOTAL (before tax). If you see a "Grand Total" or "Total Including VAT", that goes in total_amount_incl_tax, NOT total_amount.""",

        "Quote": """PROCUREMENT CONTEXT — QUOTE/QUOTATION:
A quote is a pricing proposal from a SUPPLIER to a prospective BUYER.
- quote_id: The quotation reference number. Look for "Quote No", "Quotation #", "Quote Number", "Ref", "QTE-", "QUT".
- supplier_id: The company that CREATED and SENT this quote — the SELLER. This is the company whose name, logo, letterhead, or address appears at the TOP of the document. For Excel quotes, the supplier name is in the first few rows above the table (e.g., "PeopleFirst HR Solutions Ltd", "SupplyX Ltd"). This is NOT the "Prepared For" or "Bill To" company.
- buyer_id: The company the quote was SENT TO — the prospective BUYER/CUSTOMER. Look for "Prepared For", "Customer", "Attention", "Bill To", "Invoice Address". For Excel quotes, this is the company in the "Prepared For" or "Invoice Address" section.
- supplier_address: The supplier's full address (from letterhead/header).
- buyer_address: The buyer's address (from "Invoice Address", "Delivery Address", "Bill To").
- quote_date: When the quote was issued. Look for "Quote Date", "Date", "Issued".
- validity_date: When the quote expires. Look for "Valid Until", "Expiry", "Quote valid for X days". Must be AFTER quote_date.
- total_amount: The SUBTOTAL before tax. Look for "Subtotal", "Net Total", "Total (ex VAT)". NOT the final total with VAT.
- tax_amount: The VAT/tax amount. Look for "VAT", "Tax".
- tax_percent: The VAT/tax rate as a number (e.g. 20 for 20%).
- total_amount_incl_tax: The FINAL total including tax. Look for "Total (GBP)", "Grand Total", "Total Payable".""",

        "Contract": """PROCUREMENT CONTEXT — CONTRACT:
A contract is a binding agreement between parties for goods/services over a period.
- contract_id: The contract reference number.
- supplier_id: The contracting VENDOR/supplier.
- contract_start_date / contract_end_date: The contract period.
- total_contract_value: The full monetary value of the contract.""",
    }

    def _build_extraction_prompt(
        self,
        text: str,
        doc_type: str,
        header_cols: List[str],
        line_cols: List[str],
        schema: Dict,
    ) -> str:
        """Build a procurement-intelligent extraction prompt.

        Embeds deep domain knowledge about procurement document types,
        field semantics, and column-mapping rules so the LLM understands
        the business context and maps fields correctly.
        """
        # Filter out audit/system columns the LLM shouldn't extract
        _SKIP_COLS = {
            "created_date", "created_by", "last_modified_date", "last_modified_by",
            "exchange_rate_to_usd", "converted_amount_usd",
        }
        header_cols = [c for c in header_cols if c not in _SKIP_COLS]
        line_cols = [c for c in line_cols if c not in _SKIP_COLS]

        header_spec = "\n".join(
            f"  - {col} ({schema['header_columns'][col]})"
            for col in header_cols
        )
        line_spec = ""
        if line_cols:
            line_spec = "\nLINE ITEM COLUMNS (use these exact field names):\n" + "\n".join(
                f"  - {col} ({schema['line_columns'][col]})"
                for col in line_cols
            )

        context = self._PROCUREMENT_CONTEXT.get(doc_type, "")

        return f"""You are ProcWise, an expert procurement document extraction system.
Extract ALL data from this {doc_type} document with absolute accuracy.

{context}

DATABASE COLUMNS — use these EXACT field names in your JSON response:
HEADER FIELDS:
{header_spec}
{line_spec}

CRITICAL RULES:
1. EXTRACT EXACTLY what the document says — never invent, guess, or compute values
2. If a value appears in the document, extract it verbatim (converted to the correct type)
3. If a field is NOT in the document, OMIT it entirely — do not guess
4. Dates: convert to YYYY-MM-DD format
5. Amounts: numbers only, strip currency symbols (£1,234.56 → 1234.56)
6. Currency: 3-letter ISO code (GBP, USD, EUR, AUD, etc.)
7. tax_percent: the percentage NUMBER (20% → 20). Do NOT confuse quantities, reference numbers, or other digits with tax_percent
8. Line items: extract EVERY line item row. STOP at subtotal/total/tax summary rows — those are NOT line items. Rows with quantity=0 or total=0 with no description should be SKIPPED
9. quantity: a COUNT of items (typically small: 1, 2, 5, 10, 100). NOT a price or amount
10. unit_price: cost PER SINGLE ITEM. NOT the total line amount
11. line_amount / line_total: the total for that line (usually quantity × unit_price)
12. The SUPPLIER/VENDOR is the company that ISSUED/SENT this document — their name/logo/address is at the TOP
13. The BUYER is the company RECEIVING the document — look for "Prepared For", "Bill To", "Customer", "Ship To"
14. item_id: If the document shows a product code, SKU, part number, catalog number, or item reference for a line item, extract it as item_id. Look for columns like "Item Code", "SKU", "Part No", "Product Code", "Ref", "Item #". If no product code exists in the document, OMIT item_id
15. unit_of_measure: Extract the unit if present (e.g., "each", "box", "kg", "hours", "months", "days", "per annum", "set", "licence"). Look for columns like "UOM", "Unit", "Measure". If not explicitly stated, OMIT — do not guess
16. For EXCEL/spreadsheet documents: The "DOCUMENT METADATA" section above the table contains header information (supplier company, buyer, dates, quote/PO number). The "LINE ITEMS TABLE" section contains products/services with column-labeled values. Extract header fields from metadata and line items from the table. The "TOTALS" section has subtotal, tax, and total values

RESPONSE — return ONLY this JSON structure, nothing else:
{{
  "header": {{ ... fields using exact column names above ... }},
  "line_items": [ {{ ... line item fields ... }}, ... ]
}}

DOCUMENT TEXT:
{text}"""

    def _parse_llm_response(
        self, raw: str, schema: Dict
    ) -> Optional[Dict[str, Any]]:
        """Parse LLM JSON response, handling markdown code blocks."""
        # Strip markdown code fences
        cleaned = raw.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\s*\n?", "", cleaned)
            cleaned = re.sub(r"\n?```\s*$", "", cleaned)

        try:
            data = json.loads(cleaned)
        except json.JSONDecodeError:
            # Try to find JSON object in the response
            match = re.search(r"\{[\s\S]*\}", cleaned)
            if match:
                try:
                    data = json.loads(match.group())
                except json.JSONDecodeError:
                    logger.error("Failed to parse LLM response as JSON")
                    return None
            else:
                logger.error("No JSON found in LLM response")
                return None

        header = data.get("header", {})
        line_items = data.get("line_items", [])

        # Validate: only keep columns that exist in schema
        valid_header = {}
        for k, v in header.items():
            if k in schema["header_columns"] and v is not None:
                valid_header[k] = v

        valid_lines = []
        line_cols = set(schema.get("line_columns", {}).keys())
        for item in line_items:
            valid_item = {k: v for k, v in item.items() if k in line_cols and v is not None}
            if valid_item:
                valid_lines.append(valid_item)

        return {"header": valid_header, "line_items": valid_lines}

    # ------------------------------------------------------------------
    # Database persistence
    # ------------------------------------------------------------------
    def _persist_header(
        self,
        header: Dict[str, Any],
        schema: Dict,
        audit: Dict[str, Any],
    ) -> bool:
        """Insert or update header record in the bp_ table."""
        table = schema["header_table"]
        pk_col = schema["pk"]
        pk_value = header.get(pk_col)
        if not pk_value:
            return False

        schema_name, table_name = table.split(".", 1)
        columns = dict(schema["header_columns"])

        # Build payload with type coercion
        payload = {}
        for col, val in header.items():
            if col not in columns:
                continue
            payload[col] = self._coerce_value(val, columns[col])

        # Add audit columns (already datetime objects from caller)
        for k, v in audit.items():
            if v is not None:
                payload[k] = v

        # Final sanitization: convert any remaining empty strings to None
        for k, v in list(payload.items()):
            if isinstance(v, str) and not v.strip():
                payload[k] = None

        try:
            conn = self._agent_nick.get_db_connection()
            conn.autocommit = True
            with conn.cursor() as cur:
                col_names = list(payload.keys())
                placeholders = ["%s"] * len(col_names)
                values = [payload[c] for c in col_names]

                # UPSERT: insert or update on PK conflict
                insert_sql = (
                    f'INSERT INTO {table} ({", ".join(col_names)}) '
                    f'VALUES ({", ".join(placeholders)})'
                )

                # Check if PK has unique constraint
                cur.execute(
                    "SELECT 1 FROM pg_indexes WHERE schemaname=%s AND tablename=%s "
                    "AND indexdef LIKE %s LIMIT 1",
                    (schema_name, table_name, f"%{pk_col}%"),
                )
                has_pk = cur.fetchone() is not None

                if has_pk:
                    # On update: preserve created_date/created_by, update last_modified
                    preserve_on_update = {"created_date", "created_by"}
                    update_cols = [
                        c for c in col_names
                        if c != pk_col and c not in preserve_on_update
                    ]
                    if update_cols:
                        update_clause = ", ".join(
                            f"{c} = EXCLUDED.{c}" for c in update_cols
                        )
                        insert_sql += (
                            f" ON CONFLICT ({pk_col}) DO UPDATE SET {update_clause}"
                        )
                    else:
                        insert_sql += f" ON CONFLICT ({pk_col}) DO NOTHING"
                else:
                    insert_sql += " ON CONFLICT DO NOTHING"

                cur.execute(insert_sql, values)
            conn.close()
            logger.info(
                "Persisted %s header: %s=%s (%d fields)",
                table, pk_col, pk_value, len(payload),
            )
            return True
        except Exception as exc:
            logger.exception("Failed to persist header to %s: %s", table, exc)
            return False

    def _persist_line_items(
        self,
        line_items: List[Dict[str, Any]],
        schema: Dict,
        pk_value: str,
        audit: Dict[str, Any],
    ) -> int:
        """Replace line items in the bp_ line items table.

        Deletes existing line items for this PK first, then inserts fresh.
        This ensures re-extractions fully replace stale data and recovers
        from external deletions (e.g. manual pgAdmin cleanup).
        """
        line_table = schema.get("line_table")
        if not line_table or not line_items:
            return 0

        line_fk = schema["line_fk"]
        line_pk = schema["line_pk"]
        line_seq = schema["line_seq"]
        columns = dict(schema.get("line_columns", {}))
        count = 0

        try:
            conn = self._agent_nick.get_db_connection()
            conn.autocommit = True
            with conn.cursor() as cur:
                # Delete existing line items for this parent record
                # so re-extractions fully replace stale data
                cur.execute(
                    f"DELETE FROM {line_table} WHERE {line_fk} = %s",
                    (pk_value,),
                )
                deleted = cur.rowcount
                if deleted:
                    logger.debug(
                        "Cleared %d old line items from %s for %s=%s",
                        deleted, line_table, line_fk, pk_value,
                    )

                for idx, item in enumerate(line_items, start=1):
                    payload = {}
                    for col, val in item.items():
                        if col not in columns:
                            continue
                        payload[col] = self._coerce_value(val, columns[col])

                    # Sanitize empty strings to None BEFORE auto-generation
                    # so that empty-string PKs from extraction don't block it
                    for k, v in list(payload.items()):
                        if isinstance(v, str) and not v.strip():
                            payload[k] = None

                    # Set FK and sequence
                    payload[line_fk] = pk_value
                    if line_seq and not payload.get(line_seq):
                        payload[line_seq] = idx
                    if line_pk and not payload.get(line_pk):
                        payload[line_pk] = f"{pk_value}-{idx}"

                    # Add audit columns
                    for k, v in audit.items():
                        if v is not None:
                            payload[k] = v

                    col_names = list(payload.keys())
                    placeholders = ["%s"] * len(col_names)
                    values = [payload[c] for c in col_names]

                    # UPSERT: insert or update on PK conflict
                    insert_sql = (
                        f'INSERT INTO {line_table} ({", ".join(col_names)}) '
                        f'VALUES ({", ".join(placeholders)})'
                    )
                    # Fresh insert — old rows already deleted above
                    if line_pk:
                        insert_sql += f" ON CONFLICT ({line_pk}) DO NOTHING"
                    cur.execute(insert_sql, values)
                    count += 1

            conn.close()
            logger.info(
                "Persisted %d line items to %s for %s=%s",
                count, line_table, line_fk, pk_value,
            )
        except Exception as exc:
            logger.exception("Failed to persist line items to %s: %s", line_table, exc)

        return count

    @staticmethod
    def _coerce_value(val: Any, col_type: str) -> Any:
        """Coerce a value to the target column type."""
        if val is None:
            return None

        # Empty strings should be NULL for non-text types
        if isinstance(val, str) and not val.strip():
            if col_type in ("text", "varchar", "character varying"):
                return None
            return None

        if col_type in ("numeric", "integer", "smallint"):
            if isinstance(val, (int, float)):
                return val
            text = str(val).strip()
            # Handle European format
            if re.match(r"^\d{1,3}(?:\.\d{3})+,\d{1,2}$", text):
                text = text.replace(".", "").replace(",", ".")
            elif "," in text and "." in text:
                if text.rfind(",") > text.rfind("."):
                    text = text.replace(".", "").replace(",", ".")
                else:
                    text = text.replace(",", "")
            elif "," in text:
                parts = text.split(",")
                if len(parts) == 2 and len(parts[1]) <= 2:
                    text = text.replace(",", ".")
                else:
                    text = text.replace(",", "")
            text = re.sub(r"[£$€¥₹%\s]", "", text)
            # Handle parenthesized negatives
            paren = re.match(r"^\((.+)\)$", text)
            if paren:
                text = "-" + paren.group(1)
            try:
                if col_type in ("integer", "smallint"):
                    return int(float(text))
                return float(text)
            except (ValueError, TypeError):
                return None

        if col_type == "date":
            if isinstance(val, str):
                text = val.strip()
                # Try common formats
                for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y",
                            "%m/%d/%Y", "%d %B %Y", "%d %b %Y", "%B %d, %Y",
                            "%b %d, %Y"):
                    try:
                        from datetime import datetime as dt
                        return dt.strptime(text, fmt).date()
                    except ValueError:
                        continue
                # Fallback to dateutil
                try:
                    from dateutil import parser
                    return parser.parse(text, dayfirst=True).date()
                except Exception:
                    return None
            return val

        if col_type in ("timestamp", "timestamp without time zone"):
            if isinstance(val, datetime):
                return val
            if isinstance(val, str):
                text = val.strip()
                if not text:
                    return None
                try:
                    from dateutil import parser
                    return parser.parse(text)
                except Exception:
                    return None
            return val

        # Text types — just convert to string, empty → None
        if val is not None:
            s = str(val).strip()
            return s if s else None
        return None
