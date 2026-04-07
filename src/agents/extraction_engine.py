
from __future__ import annotations

import argparse
import calendar
import glob
import json
import logging
import os
import re
import sys
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from io import BytesIO
from typing import Any, Callable, Optional, TypedDict

import fitz
import numpy as np
import pytesseract
import requests
import spacy
from docx import Document
from PIL import Image, ImageFilter, ImageEnhance, ImageOps
from spacy.language import Language


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
MAX_FILE_SIZE_MB = 50

TESSERACT_CMD = os.getenv("TESSERACT_CMD", r"C:\Program Files\Tesseract-OCR\tesseract.exe")
OCR_LANG = "eng"
OCR_DPI = 300
OCR_MIN_TEXT_LENGTH = 50

OCR_RESIZE_MIN_WIDTH = 1000
OCR_DENOISE_STRENGTH = 10
OCR_BINARIZE_BLOCK_SIZE = 31
OCR_BINARIZE_C = 10

SPACY_MODEL = "en_core_web_sm"
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
NUEXTRACT_MODEL = "nuextract:3.8b"

LINE_ITEM_ROW_Y_TOLERANCE = 8
LINE_ITEM_COL_X_GAP = 40
VENDOR_NAME_MAX_LINES = 5

LAYOUT_ROW_Y_TOLERANCE = 6.0
LAYOUT_COLUMN_GAP_RATIO = 0.20
LAYOUT_MIN_BLOCKS = 5

# Database connection function for supplier lookups (set via set_db_connection_func)
_db_connection_func: Optional[Callable] = None


def set_db_connection_func(func: Callable) -> None:
    """Set the database connection factory for supplier name lookups."""
    global _db_connection_func
    _db_connection_func = func


class BaseExtractor(ABC):

    @abstractmethod
    def extract(self, context: dict[str, Any]) -> dict[str, Any]:
        ...


logger = logging.getLogger(__name__)


_spacy_nlp: Optional[Language] = None


def get_spacy() -> Language:
    global _spacy_nlp
    if _spacy_nlp is None:
        logger.info("Loading spaCy model: %s", SPACY_MODEL)
        try:
            _spacy_nlp = spacy.load(SPACY_MODEL)
        except OSError:
            logger.warning("spaCy model not found – downloading %s", SPACY_MODEL)
            from spacy.cli import download
            download(SPACY_MODEL)
            _spacy_nlp = spacy.load(SPACY_MODEL)
        logger.info("spaCy model loaded.")
    return _spacy_nlp


class FileDetector:

    @staticmethod
    def detect(file_path: str) -> str:

        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if size_mb > MAX_FILE_SIZE_MB:
            raise ValueError(
                f"File too large ({size_mb:.1f} MB). Max allowed: {MAX_FILE_SIZE_MB} MB"
            )

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}"
            )

        return ext.lstrip(".")


class FileLoader:

    @staticmethod
    def load(file_path: str) -> bytes:
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        with open(file_path, "rb") as fh:
            return fh.read()


class JSONWriter:

    @staticmethod
    def to_string(data: dict[str, Any], indent: int = 2) -> str:
        return json.dumps(data, indent=indent, ensure_ascii=False)

    @staticmethod
    def to_file(data: dict[str, Any], output_path: str, indent: int = 2) -> str:

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        with open(output_path, "w", encoding="utf-8") as fh:
            json.dump(data, fh, indent=indent, ensure_ascii=False)

        abs_path = os.path.abspath(output_path)
        logger.info("JSON output written to %s", abs_path)
        return abs_path


class DOCXParser:

    @staticmethod
    def extract_text(file_path: str) -> str:

        doc = Document(file_path)
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n".join(parts)

    @staticmethod
    def extract_from_bytes(data: bytes) -> str:

        doc = Document(BytesIO(data))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n".join(parts)

    @staticmethod
    def extract_tables(file_path: str) -> list[list[list[str]]]:

        doc = Document(file_path)
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
        return tables


def _pdf_collapse_spaced_text(text: str) -> str:
    def _collapse_match(m: re.Match) -> str:
        collapsed = m.group(0).replace(" ", "")
        return collapsed

    text = re.sub(
        r"(?<!\S)([A-Za-z0-9:,] ){1,}[A-Za-z0-9:,](?!\S)",
        _collapse_match,
        text,
    )
    return text


def _pdf_clean_text(text: str) -> str:
    text = text.replace("\ufffd", "\u00a3")
    text = _pdf_collapse_spaced_text(text)
    return text


class PDFParser:

    @staticmethod
    def extract_text(file_path: str) -> str:

        text_parts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = page.get_text("text")
                if page_text:
                    text_parts.append(page_text.strip())
        return _pdf_clean_text("\n".join(text_parts))

    @staticmethod
    def extract_text_by_page(file_path: str) -> list[str]:

        pages: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                pages.append(page.get_text("text").strip())
        return pages

    @staticmethod
    def is_scanned(file_path: str, min_chars: int = 50) -> bool:

        with fitz.open(file_path) as doc:
            for page in doc:
                if len(page.get_text("text").strip()) >= min_chars:
                    return False
        return True


if os.path.isfile(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


class WordBox(TypedDict):
    text: str
    x: int
    y: int
    w: int
    h: int


def _preprocess(img: Image.Image) -> Image.Image:
    img = img.convert("L")

    if img.width < OCR_RESIZE_MIN_WIDTH:
        scale = OCR_RESIZE_MIN_WIDTH / img.width
        img = img.resize(
            (int(img.width * scale), int(img.height * scale)),
            Image.LANCZOS,
        )

    img = ImageEnhance.Contrast(img).enhance(1.8)

    img = img.filter(ImageFilter.MedianFilter(size=3))

    arr = np.array(img)
    threshold = _otsu_threshold(arr)
    arr = ((arr > threshold) * 255).astype(np.uint8)
    img = Image.fromarray(arr)

    return img


def _otsu_threshold(gray: np.ndarray) -> int:

    hist, _ = np.histogram(gray.ravel(), bins=256, range=(0, 256))
    total = gray.size
    current_sum = 0.0
    total_sum = float(np.dot(np.arange(256), hist))
    weight_bg = 0.0
    max_variance = 0.0
    best_thresh = 0

    for t in range(256):
        weight_bg += hist[t]
        if weight_bg == 0:
            continue
        weight_fg = total - weight_bg
        if weight_fg == 0:
            break
        current_sum += t * hist[t]
        mean_bg = current_sum / weight_bg
        mean_fg = (total_sum - current_sum) / weight_fg
        variance = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
        if variance > max_variance:
            max_variance = variance
            best_thresh = t

    return best_thresh


def _pdf_to_images(file_path: str, dpi: int = OCR_DPI) -> list[Image.Image]:

    images: list[Image.Image] = []
    zoom = dpi / 72
    mat = fitz.Matrix(zoom, zoom)
    with fitz.open(file_path) as doc:
        for page in doc:
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            images.append(img)
    return images


def run_ocr(file_path: str) -> str:
    images = _pdf_to_images(file_path)
    text_parts: list[str] = []

    for img in images:
        processed = _preprocess(img)
        text = pytesseract.image_to_string(processed, lang=OCR_LANG)
        if text.strip():
            text_parts.append(text.strip())

    return "\n".join(text_parts)


def run_ocr_with_boxes(file_path: str) -> tuple[str, list[WordBox]]:
    images = _pdf_to_images(file_path)
    all_text_parts: list[str] = []
    all_boxes: list[WordBox] = []
    y_offset = 0

    for img in images:
        processed = _preprocess(img)

        text = pytesseract.image_to_string(processed, lang=OCR_LANG)
        if text.strip():
            all_text_parts.append(text.strip())

        data = pytesseract.image_to_data(processed, lang=OCR_LANG, output_type=pytesseract.Output.DICT)
        n = len(data["text"])
        for i in range(n):
            word = data["text"][i].strip()
            if not word:
                continue
            all_boxes.append(
                WordBox(
                    text=word,
                    x=data["left"][i],
                    y=data["top"][i] + y_offset,
                    w=data["width"][i],
                    h=data["height"][i],
                )
            )

        y_offset += img.height

    full_text = "\n".join(all_text_parts)
    return full_text, all_boxes


def spacy_extract_entities(text: str) -> dict[str, list[str]]:
    nlp = get_spacy()
    doc = nlp(text)

    entities: dict[str, list[str]] = {}
    for ent in doc.ents:
        label = ent.label_
        value = ent.text.strip()
        if not value:
            continue
        entities.setdefault(label, [])
        if value not in entities[label]:
            entities[label].append(value)

    return entities


INVOICE_TEMPLATE = {
    "invoice number": "",
    "purchase order number": "",
    "invoice date": "",
    "due date": "",
    "vendor or seller name": "",
    "vendor or seller address": "",
    "bill to or invoice to company name": "",
    "bill to or invoice to address": "",
    "subtotal amount before tax": "",
    "tax amount": "",
    "total amount after tax": "",
    "currency code": "",
    "bank name": "",
    "bank account number": "",
    "bank sort code": "",
    "phone number": "",
    "email address": "",
}

_LABEL_MAP_INVOICE = {
    "invoice number": "invoice number",
    "purchase order number": "purchase order number",
    "invoice date": "date",
    "due date": "due date",
    "vendor or seller name": "vendor name",
    "vendor or seller address": "address",
    "bill to or invoice to company name": "company name",
    "bill to or invoice to address": "bill_to_address",
    "subtotal amount before tax": "subtotal",
    "tax amount": "tax amount",
    "total amount after tax": "total amount",
    "currency code": "currency",
    "bank name": "bank name",
    "bank account number": "bank account",
    "bank sort code": "sort code",
    "phone number": "phone number",
    "email address": "email address",
}


def nuextract_extract_entities_invoice(text: str, labels: list[str] | None = None) -> dict[str, list[str]]:
    chunks = _split_text(text, max_chars=3000)
    entities: dict[str, list[str]] = {}

    for chunk in chunks:
        try:
            extracted = _call_nuextract_invoice(chunk)
        except Exception as exc:
            logger.warning("NuExtract extraction failed on chunk: %s", exc)
            continue

        for key, value in extracted.items():
            if key not in _LABEL_MAP_INVOICE:
                continue

            label = _LABEL_MAP_INVOICE[key]
            values = value if isinstance(value, list) else [value]

            for v in values:
                v = str(v).strip()
                if not v:
                    continue
                entities.setdefault(label, [])
                if v not in entities[label]:
                    entities[label].append(v)

    return entities


def _call_nuextract_invoice(text: str) -> dict:
    template_str = json.dumps(INVOICE_TEMPLATE, indent=2)

    prompt = f"""<|input|>
{text}
<|template|>
{template_str}
<|output|>"""

    response = requests.post(
        f"{OLLAMA_BASE_URL}/api/generate",
        json={
            "model": NUEXTRACT_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0,
                "num_predict": 2048,
                "num_gpu": 0,
            },
        },
        timeout=120,
    )
    response.raise_for_status()

    raw_output = response.json().get("response", "").strip()
    logger.debug("NuExtract raw output: %s", raw_output[:500])

    return _parse_response(raw_output)


PO_TEMPLATE = {
    "purchase order number": "",
    "order date": "",
    "quote number": "",
    "buyer or ordering company name": "",
    "buyer or ordering company address": "",
    "supplier or recipient name": "",
    "supplier or recipient address": "",
    "subtotal amount before tax": "",
    "discount amount": "",
    "tax amount": "",
    "total amount after tax": "",
    "currency code": "",
    "payment terms": "",
    "delivery address": "",
    "phone number": "",
    "email address": "",
}

_LABEL_MAP_PO = {
    "purchase order number": "po_number",
    "order date": "order_date",
    "quote number": "quote_number",
    "buyer or ordering company name": "buyer_name",
    "buyer or ordering company address": "buyer_address",
    "supplier or recipient name": "supplier_name",
    "supplier or recipient address": "supplier_address",
    "subtotal amount before tax": "subtotal",
    "discount amount": "discount",
    "tax amount": "tax_amount",
    "total amount after tax": "total_amount",
    "currency code": "currency",
    "payment terms": "payment_terms",
    "delivery address": "delivery_address",
    "phone number": "phone_number",
    "email address": "email_address",
}


def nuextract_extract_entities_po(text: str, labels: list[str] | None = None) -> dict[str, list[str]]:
    chunks = _split_text(text, max_chars=3000)
    entities: dict[str, list[str]] = {}

    for chunk in chunks:
        try:
            extracted = _call_nuextract_po(chunk)
        except Exception as exc:
            logger.warning("NuExtract extraction failed on chunk: %s", exc)
            continue

        for key, value in extracted.items():
            if key not in _LABEL_MAP_PO:
                continue

            label = _LABEL_MAP_PO[key]
            values = value if isinstance(value, list) else [value]

            for v in values:
                v = str(v).strip()
                if not v:
                    continue
                entities.setdefault(label, [])
                if v not in entities[label]:
                    entities[label].append(v)

    return entities


def _call_nuextract_po(text: str) -> dict:
    template_str = json.dumps(PO_TEMPLATE, indent=2)

    prompt = f"""<|input|>
{text}
<|template|>
{template_str}
<|output|>"""

    response = requests.post(
        f"{OLLAMA_BASE_URL}/api/generate",
        json={
            "model": NUEXTRACT_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0,
                "num_predict": 2048,
                "num_gpu": 0,
            },
        },
        timeout=120,
    )
    response.raise_for_status()

    raw_output = response.json().get("response", "").strip()
    logger.debug("NuExtract raw output: %s", raw_output[:500])

    return _parse_response(raw_output)


def _parse_response(raw: str) -> dict:

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    m = re.search(r"```(?:json)?\s*\n?(.*?)```", raw, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1).strip())
        except json.JSONDecodeError:
            pass

    brace_start = raw.find("{")
    brace_end = raw.rfind("}")
    if brace_start != -1 and brace_end > brace_start:
        try:
            return json.loads(raw[brace_start:brace_end + 1])
        except json.JSONDecodeError:
            pass

    logger.warning("Could not parse NuExtract output as JSON: %s", raw[:200])
    return {}


def _split_text(text: str, max_chars: int = 3000) -> list[str]:

    lines = text.split("\n")
    chunks: list[str] = []
    current: list[str] = []
    length = 0

    for line in lines:
        if length + len(line) > max_chars and current:
            chunks.append("\n".join(current))
            current = []
            length = 0
        current.append(line)
        length += len(line) + 1

    if current:
        chunks.append("\n".join(current))

    return chunks


SOURCE_BASE_SCORES: dict[str, float] = {
    "nuextract":        0.72,
    "spacy":            0.68,
    "regex_context":    0.65,
    "regex_fallback":   0.55,
    "ocr_table":        0.75,
    "heuristic":        0.60,
    "text_table":       0.70,
    "layout":           0.78,
}


@dataclass
class Candidate:

    value: str
    confidence: float
    source: str

    def __repr__(self) -> str:
        return f"Candidate({self.value!r}, conf={self.confidence:.2f}, src={self.source!r})"


_COMPANY_SUFFIXES = re.compile(
    r"\b(ltd|limited|llc|inc|incorporated|corp|corporation|plc|gmbh|"
    r"pvt|private|co|company|sa|srl|ag|pty)\b",
    re.IGNORECASE,
)

_CLEAN_MONEY_RE = re.compile(r"^[£$€¥₹\u00a3]?\s*[\d,]+\.\d{2}$")
_HAS_CURRENCY_RE = re.compile(r"[£$€¥₹\u00a3]")
_DATE_FULL_RE = re.compile(r"\d{1,2}[/\-]\d{1,2}[/\-]\d{4}")
_DATE_SHORT_RE = re.compile(r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2}")


def adjust_vendor_name(c: Candidate) -> Candidate:
    v = c.value.strip()
    conf = c.confidence

    word_count = len(v.split())
    if word_count >= 2:
        conf += 0.10
    elif word_count == 1 and len(v) <= 3:
        conf -= 0.10

    if _COMPANY_SUFFIXES.search(v):
        conf += 0.12

    if v.isupper() and word_count >= 1:
        conf += 0.03

    if re.search(r"\d", v):
        conf -= 0.15

    return Candidate(c.value, _clamp(conf), c.source)


def adjust_invoice_number(c: Candidate) -> Candidate:

    v = c.value.strip()
    conf = c.confidence

    if re.match(r"^\d+$", v):
        conf += 0.08

    if re.match(r"^(INV|INV-|#)\d+", v, re.IGNORECASE):
        conf += 0.10

    if len(v) < 3:
        conf -= 0.15

    if len(v) > 20:
        conf -= 0.10

    if " " in v:
        conf -= 0.08

    return Candidate(c.value, _clamp(conf), c.source)


def adjust_date(c: Candidate) -> Candidate:

    v = c.value.strip()
    conf = c.confidence

    if _DATE_FULL_RE.match(v):
        conf += 0.12
    elif _DATE_SHORT_RE.match(v):
        conf += 0.05

    if re.match(r"^\d{4}-\d{2}-\d{2}$", v):
        conf += 0.15

    if re.search(r"[a-zA-Z]{4,}", v):
        conf -= 0.10

    numeric_date = re.match(r"^(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})$", v)
    if numeric_date:
        d, m, y = int(numeric_date.group(1)), int(numeric_date.group(2)), int(numeric_date.group(3))
        if m > 12 or d > 31 or (m == 0 or d == 0):
            conf -= 0.50
        if y < 100 and y > 50:
            conf -= 0.20

    return Candidate(c.value, _clamp(conf), c.source)


def adjust_money(c: Candidate) -> Candidate:

    v = c.value.strip()
    conf = c.confidence

    if _CLEAN_MONEY_RE.match(v):
        conf += 0.15

    if _HAS_CURRENCY_RE.search(v):
        conf += 0.08

    if re.match(r"^[£$€¥₹\u00a3]?\s*[\d,]+$", v):
        conf -= 0.05

    stripped = re.sub(r"[£$€¥₹\u00a3,.\s]", "", v)
    if not stripped.isdigit():
        conf -= 0.20

    return Candidate(c.value, _clamp(conf), c.source)


def adjust_address(c: Candidate) -> Candidate:

    v = c.value.strip()
    conf = c.confidence

    if len(v) > 40:
        conf += 0.10
    elif len(v) > 20:
        conf += 0.05
    elif len(v) < 10:
        conf -= 0.15

    if re.search(r"[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}", v, re.IGNORECASE):
        conf += 0.08
    if re.search(r"\b\d{5}(?:-\d{4})?\b", v):
        conf += 0.08

    if "," in v:
        conf += 0.05

    return Candidate(c.value, _clamp(conf), c.source)


def adjust_line_item(c: Candidate) -> Candidate:

    conf = c.confidence
    return Candidate(c.value, _clamp(conf), c.source)


def pick_best(candidates: list[Candidate]) -> str:
    if not candidates:
        return ""

    valid = [c for c in candidates if c.value.strip()]
    if not valid:
        return ""

    valid.sort(key=lambda c: (c.confidence, len(c.value)), reverse=True)

    winner = valid[0]
    if len(valid) > 1:
        logger.debug(
            "Candidate ranking: winner=%r (%.2f), runner-up=%r (%.2f)",
            winner.value, winner.confidence,
            valid[1].value, valid[1].confidence,
        )

    return winner.value


def pick_best_money(candidates: list[Candidate]) -> str:
    if not candidates:
        return ""

    valid = [c for c in candidates if c.value.strip()]
    if not valid:
        return ""

    def _numeric(v: str) -> float:
        m = re.search(r"[\d,]+\.?\d*", v)
        if m:
            try:
                return float(m.group(0).replace(",", ""))
            except ValueError:
                pass
        return 0.0

    valid.sort(
        key=lambda c: (round(c.confidence, 1), _numeric(c.value)),
        reverse=True,
    )

    winner = valid[0]
    logger.debug("Money candidate winner: %r (conf=%.2f, val=%.2f)",
                 winner.value, winner.confidence, _numeric(winner.value))
    return winner.value


def _clamp(v: float, lo: float = 0.0, hi: float = 1.0) -> float:
    return max(lo, min(hi, v))


def base_score(source: str) -> float:

    return SOURCE_BASE_SCORES.get(source, 0.50)


_CURR = r"[£$€¥₹\u00a3]"


_INVOICE_NUM_PATTERNS = [
    re.compile(
        r"(?:^|\n)\s*INVOICE\s*\n\s*(\d{3,20})\s*(?:\n|$)",
    ),
    re.compile(
        r"(?:invoice\s*(?:no|number|#|num))\s*[.:\s]*\n?\s*"
        r"(\d[\w\-/]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:invoice\s*(?:no|number|#|num))\s*[.:\s]*\n?\s*"
        r"([A-Z][\w\-/]*\d[\w\-/]*)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:inv\s*(?:no|number|#|num)?)\s*[.:\s]+\n?\s*"
        r"([A-Z][\w\-/]*\d[\w\-/]*)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:inv[\s.#:]+)"
        r"(\d[\w\-/]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:bill\s*(?:no|number|#)[\s.:]*)"
        r"(\d[\w\-/]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:^|\n)\s*(INV[\-/]?[\w\-/]{3,20})\s*(?:\n|$)",
    ),
]


def find_invoice_number(text: str) -> Optional[str]:
    for pat in _INVOICE_NUM_PATTERNS:
        m = pat.search(text)
        if m:
            return m.group(1).strip()
    return None


_PO_NUM_PATTERNS = [
    re.compile(
        r"(?:purchase\s*order\s*(?:no|number|#)?)\s*[:\s]*\n?\s*"
        r"(?:PO\s*)?(\d[\w\-/]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:p\.?o\.?\s*(?:no|number|#|num|ref(?:erence)?)?)\s*[.:\s]+(?:PO\s*)?(\d[\w\-/]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:p\.?o\.?\s*(?:no|number|#|num|ref(?:erence)?)?)\s*[.:\s]*\n\s*(?:PO\s*)?(\d[\w\-/]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(\d{4,20})\s*\n\s*(?:p\.?o\.?\s*(?:no|number|#|num|ref(?:erence)?)?)\s*[.:\s]*(?!\d)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:^|\n)\s*PO\s*(\d{5,20})\s*(?:\n|$)",
    ),
]


def find_po_number(text: str) -> Optional[str]:
    for pat in _PO_NUM_PATTERNS:
        m = pat.search(text)
        if m:
            return m.group(1).strip()
    return None


_DATE_PATTERNS = [
    re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})\b"),
    re.compile(
        r"\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?"
        r"\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?"
        r",?\s+\d{4})\b",
        re.IGNORECASE,
    ),
    re.compile(r"\b(\d{4}-\d{2}-\d{2})\b"),
    re.compile(
        r"\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?"
        r"\s+\d{4})\b",
        re.IGNORECASE,
    ),
]


def find_all_dates(text: str) -> list[str]:
    dates: list[str] = []
    for pat in _DATE_PATTERNS:
        for m in pat.finditer(text):
            val = m.group(1).strip()
            if val not in dates:
                dates.append(val)
    return dates


_INVOICE_DATE_PATTERNS = [
    re.compile(
        r"(?:invoice\s*date|date\s*(?:of\s*)?invoice|date\s*issued|issue\s*date|issued\s*(?:on|date))"
        r"[\s.:]*\n?([\s\S]{0,200})",
        re.IGNORECASE,
    ),
    re.compile(r"(?:date)[\s.:]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})", re.IGNORECASE),
]


def find_invoice_date(text: str) -> Optional[str]:
    for pat in _INVOICE_DATE_PATTERNS:
        m = pat.search(text)
        if m:
            raw = m.group(1).strip()
            for dp in _DATE_PATTERNS:
                dm = dp.search(raw)
                if dm:
                    return dm.group(1)
    return None


_DUE_DATE_PATTERNS = [
    re.compile(
        r"(?:due\s*date|payment\s*due|date\s*due|pay\s*by)"
        r"[\s.:]*\n?([\s\S]{0,80})",
        re.IGNORECASE,
    ),
]


def find_due_date(text: str) -> Optional[str]:
    for pat in _DUE_DATE_PATTERNS:
        m = pat.search(text)
        if m:
            raw = m.group(1).strip()
            for dp in _DATE_PATTERNS:
                dm = dp.search(raw)
                if dm:
                    return dm.group(1)
    return None


_CURRENCY_SYMBOLS = {
    "£": "GBP", "\u00a3": "GBP",
    "$": "USD", "€": "EUR", "¥": "JPY", "₹": "INR",
}

_MONEY_PATTERN = re.compile(rf"({_CURR})\s*([\d,]+\.?\d*)")


def detect_currency(text: str) -> Optional[str]:
    m = _MONEY_PATTERN.search(text)
    if m:
        return _CURRENCY_SYMBOLS.get(m.group(1), m.group(1))
    for code in ("GBP", "USD", "EUR", "JPY", "INR"):
        if code in text.upper():
            return code
    return None


_TOTAL_PATTERNS = [
    re.compile(
        rf"(?:^|\n)\s*TOTAL\s*[:\s]*\n\s*({_CURR}\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
    re.compile(
        rf"(?:^|\n)\s*TOTAL\s+({_CURR}\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
    re.compile(
        rf"(?:grand[\s\-]*total|amount\s*due|balance\s*due)\s*[:\s]*\n?\s*({_CURR}?\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
    re.compile(
        rf"(?:total\s*(?:amount|due|payable))\s*[:\s]*\n?\s*({_CURR}?\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
]


_SUBTOTAL_PATTERNS = [
    re.compile(
        rf"(?:sub[\s\-]total|net\s*(?:amount|total))\s*[:\s]+({_CURR}?\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
    re.compile(
        rf"(?:sub[\s\-]total)\s*\n"
        rf"(?:.*\n)*?"
        rf"\s*({_CURR}\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
]


_TAX_PATTERNS = [
    re.compile(
        rf"(?:tax|vat|gst|sales\s*tax)\s*(?:\([\d.]+%\))?[: \t]+({_CURR}\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
]


def _extract_amount(patterns: list[re.Pattern], text: str) -> Optional[str]:
    for pat in patterns:
        m = pat.search(text)
        if m:
            raw = m.group(1).strip()
            if raw:
                return raw
    return None


def find_total(text: str) -> Optional[str]:
    return _extract_amount(_TOTAL_PATTERNS, text)


def find_subtotal(text: str) -> Optional[str]:
    result = _extract_amount(_SUBTOTAL_PATTERNS, text)
    if result:
        return result

    m = re.search(
        rf"Sub[\s\-]?Total\s*\n\s*(?:Tax|VAT|GST).*?\n\s*({_CURR}\s*[\d,]+\.?\d*)\s*\n",
        text,
        re.IGNORECASE,
    )
    if m:
        return m.group(1).strip()

    return None


def find_tax(text: str) -> Optional[str]:
    result = _extract_amount(_TAX_PATTERNS, text)
    if result:
        return result

    m = re.search(
        rf"Sub[\s\-]?Total\s*\n\s*(?:Tax|VAT|GST).*?\n"
        rf"\s*{_CURR}\s*[\d,]+\.?\d*\s*\n"
        rf"\s*({_CURR}\s*[\d,]+\.?\d*)",
        text,
        re.IGNORECASE,
    )
    if m:
        return m.group(1).strip()

    return None


_DISCOUNT_PATTERNS = [
    re.compile(
        rf"(?:discount)\s*(?:\([\d.]+%\))?\s*[:\s]+({_CURR}\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
    re.compile(
        rf"(?:discount)\s*(?:\([\d.]+%\))?\s*[:\s]*\n\s*({_CURR}\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
]


def find_discount(text: str) -> Optional[str]:
    return _extract_amount(_DISCOUNT_PATTERNS, text)


_BILL_TO_PATTERN = re.compile(
    r"(?:bill(?:ing)?\s*(?:to|ed\s*to)|customer\s*information)\s*[:\s]*\n?((?:.+\n?){1,10})",
    re.IGNORECASE,
)

_SHIP_TO_PATTERN = re.compile(
    r"(?:ship\s*to|deliver\s*to|delivery\s*address)\s*[:\s]*\n?((?:.+\n?){1,5})",
    re.IGNORECASE,
)


def find_bill_to(text: str) -> Optional[str]:
    m = _BILL_TO_PATTERN.search(text)
    if m:
        block = m.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        filtered: list[str] = []
        for line in lines:
            if re.match(r"^(payable|notes|payment|bank|account|total|description|item\s+description|purchase\s*order|p\.?o\.?\s*(?:no|number|#|ref))", line, re.IGNORECASE):
                break
            if re.match(r"^(date\s*issued|invoice\s*no|po\s*no)\s*[:.]?\s*$", line, re.IGNORECASE):
                continue
            if re.match(r"^(invoice(\s+to)?|bill(\s+to)?|billed(\s+to)?|receipt|statement)\s*[:.]?\s*$", line, re.IGNORECASE):
                continue
            if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}$", line):
                continue
            if re.match(r"^(https?://|www\.|.*@)", line, re.IGNORECASE):
                break
            if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", line):
                break
            filtered.append(line)
        if filtered:
            return ", ".join(filtered[:4])
    return None


def find_ship_to(text: str) -> Optional[str]:
    m = _SHIP_TO_PATTERN.search(text)
    if m:
        block = m.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()][:5]
        return ", ".join(lines)
    return None


_RECIPIENT_PATTERN = re.compile(
    r"(?:recipient|attention|attn|for\s+the\s+attention\s+of)\s*[:\s]*\n?((?:.+\n?){1,5})",
    re.IGNORECASE,
)


def find_recipient(text: str) -> Optional[str]:
    m = _RECIPIENT_PATTERN.search(text)
    if m:
        block = m.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        filtered: list[str] = []
        for line in lines:
            if re.match(r"^(payable|notes|payment|bank|account|total|description)", line, re.IGNORECASE):
                break
            if re.match(r"^(https?://|www\.|.*@)", line, re.IGNORECASE):
                break
            filtered.append(line)
        if filtered:
            return ", ".join(filtered[:3])
    return None


_LINE_ITEM_PATTERN = re.compile(
    r"^(.+?)\s+"
    rf"({_CURR}?\d[\d,]*\.?\d{{0,2}})\s+"
    r"(\d+)\s+"
    rf"({_CURR}?\d[\d,]*\.?\d{{0,2}})\s*$",
    re.MULTILINE,
)

_LINE_ITEM_PATTERN_ALT = re.compile(
    r"^(.+?)\s+"
    r"(\d+)\s+"
    rf"({_CURR}?\d[\d,]*\.?\d{{0,2}})\s+"
    rf"({_CURR}?\d[\d,]*\.?\d{{0,2}})\s*$",
    re.MULTILINE,
)


def find_line_items_from_text(text: str) -> list[dict[str, str]]:
    filtered_text = _remove_phone_lines(text)

    all_sets: list[list[dict[str, str]]] = []

    items: list[dict[str, str]] = []
    for m in _LINE_ITEM_PATTERN.finditer(filtered_text):
        desc = m.group(1).strip()
        if _is_summary_label(desc) or _is_phone_fragment(desc):
            continue
        items.append({
            "item": desc,
            "unit_price": m.group(2),
            "quantity": m.group(3),
            "amount": m.group(4),
        })
    if items:
        all_sets.append(items)

    items = []
    for m in _LINE_ITEM_PATTERN_ALT.finditer(filtered_text):
        desc = m.group(1).strip()
        if _is_summary_label(desc) or _is_phone_fragment(desc):
            continue
        items.append({
            "item": desc,
            "quantity": m.group(2),
            "unit_price": m.group(3),
            "amount": m.group(4),
        })
    if items:
        all_sets.append(items)

    items = _parse_newline_separated_table(filtered_text)
    if items:
        all_sets.append(items)

    if not all_sets:
        return []

    def _set_quality(item_set: list[dict[str, str]]) -> tuple[int, float]:
        items_with_amounts = sum(1 for it in item_set if it.get("amount", "").strip())
        avg_desc_len = sum(len(it.get("item", "")) for it in item_set) / len(item_set)
        return (items_with_amounts, avg_desc_len)

    all_sets.sort(key=_set_quality, reverse=True)
    best = all_sets[0]

    for other_set in all_sets[1:]:
        for other_item in other_set:
            other_amt = re.sub(r"[£$€¥₹\u00a3,\s]", "", other_item.get("amount", ""))
            other_desc = other_item.get("item", "")
            for best_item in best:
                best_amt = re.sub(r"[£$€¥₹\u00a3,\s]", "", best_item.get("amount", ""))
                best_desc = best_item.get("item", "")
                if (best_amt and best_amt == other_amt
                        and len(other_desc) > len(best_desc)
                        and best_desc.lower() in other_desc.lower()):
                    best_item["item"] = other_desc

    return best


def _parse_newline_separated_table(text: str) -> list[dict[str, str]]:
    lines = text.split("\n")

    header_idx = _find_table_header(lines)
    if header_idx is None:
        return []

    header_info = _detect_header_columns(lines, header_idx)
    if not header_info:
        return []

    col_count, start_idx = header_info
    data_lines = lines[start_idx:]

    items: list[dict[str, str]] = []
    i = 0
    skipped = 0
    while i + col_count - 1 < len(data_lines):
        row = [data_lines[i + j].strip() for j in range(col_count)]

        desc = row[0]
        if not desc or _is_summary_label(desc) or _looks_like_number(desc) or _is_phone_fragment(desc):
            i += 1
            skipped += 1
            if skipped > col_count + 2:
                break
            continue

        if col_count == 4:
            if not _looks_like_number(row[1]) and i + col_count < len(data_lines):
                extra_desc = row[1]
                qty, price = row[2], row[3]
                amount_line = data_lines[i + col_count].strip()
                if not (_looks_like_number(qty) or _looks_like_number(amount_line)):
                    i += 1
                    continue
                items.append({
                    "item": f"{desc} {extra_desc}",
                    "quantity": qty,
                    "unit_price": price,
                    "amount": amount_line,
                })
                skipped = 0
                i += col_count + 1
            else:
                if not any(_looks_like_number(row[j]) for j in range(1, 4)):
                    i += 1
                    skipped += 1
                    if skipped > col_count * 2:
                        break
                    continue
                items.append({
                    "item": desc,
                    "unit_price": row[1],
                    "quantity": row[2],
                    "amount": row[3],
                })
                skipped = 0
                i += col_count
        elif col_count == 3:
            if not _looks_like_number(row[1]) and i + col_count < len(data_lines):
                extra_desc = row[1]
                qty = row[2]
                amount_line = data_lines[i + col_count].strip()
                if not (_looks_like_number(qty) or _looks_like_number(amount_line)):
                    i += 1
                    skipped += 1
                    if skipped > col_count * 2:
                        break
                    continue
                items.append({
                    "item": f"{desc} {extra_desc}",
                    "quantity": qty,
                    "unit_price": "",
                    "amount": amount_line,
                })
                skipped = 0
                i += col_count + 1
            else:
                if not any(_looks_like_number(row[j]) for j in range(1, 3)):
                    i += 1
                    skipped += 1
                    if skipped > col_count * 2:
                        break
                    continue
                items.append({
                    "item": desc,
                    "quantity": row[1],
                    "unit_price": "",
                    "amount": row[2],
                })
                skipped = 0
                i += col_count
        else:
            i += 1

    return items


def _find_table_header(lines: list[str]) -> Optional[int]:

    header_starters = {"description", "item", "particular", "particulars",
                       "product", "service", "details"}
    for i, line in enumerate(lines):
        if line.strip().lower() in header_starters:
            return i
    return None


def _detect_header_columns(lines: list[str], header_start: int) -> Optional[tuple[int, int]]:

    header_keywords = {
        "description", "item", "particular", "particulars", "product",
        "service", "details", "price", "rate", "unit price", "cost",
        "qty", "quantity", "units", "count",
        "amount", "subtotal", "sub-total", "total", "line total", "extended",
    }

    col_count = 0
    i = header_start
    while i < len(lines):
        stripped = lines[i].strip().lower()
        if stripped in header_keywords:
            col_count += 1
            i += 1
        else:
            break

    if col_count >= 3:
        return col_count, i

    return None


_FP_PHONE_RE = re.compile(
    r"(?:\+\d{1,3}[\s\-]?)?"
    r"(?:\(?\d{2,5}\)?[\s\-]?)?"
    r"\d{3,4}[\s\-]\d{3,4}"
    r"(?:[\s\-]\d{2,4})?",
)


def _is_phone_fragment(s: str) -> bool:

    cleaned = s.strip()
    if cleaned.startswith("+"):
        return True
    if _FP_PHONE_RE.fullmatch(cleaned):
        return True
    return False


def _remove_phone_lines(text: str) -> str:

    lines = text.split("\n")
    filtered = []
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("+") and re.match(r"^\+[\d\s\-()]{7,20}$", stripped):
            continue
        if _FP_PHONE_RE.fullmatch(stripped):
            continue
        filtered.append(line)
    return "\n".join(filtered)


def _looks_like_number(s: str) -> bool:

    cleaned = re.sub(rf"[£$€¥₹\u00a3,\s]", "", s)
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def _is_summary_label(label: str) -> bool:

    lower = label.lower().strip().rstrip(":")
    skip_exact = {"total", "sub-total", "subtotal", "tax", "vat", "gst", "discount",
                  "shipping", "freight", "balance", "amount due", "grand total", "net",
                  "total amount", "vat (20%)", "vat (10%)", "vat (5%)"}
    if lower in skip_exact:
        return True

    skip_starts = ["total amount", "vat ", "vat(", "sales tax", "grand total",
                   "tax ", "tax("]
    for prefix in skip_starts:
        if lower.startswith(prefix):
            return True

    skip_contains = ["account no", "account number", "account name", "sort code",
                     "bank", "iban", "swift", "bic", "payable to", "remit to",
                     "po no", "invoice no", "date issued", "vat no",
                     "pay by", "payment", "thank you",
                     "invoice reference", "this invoice"]
    for kw in skip_contains:
        if kw in lower:
            return True

    return False


_PAYMENT_TERMS_PATTERNS = [
    re.compile(
        r"(?:payment|payable|pay)\s+(?:\w+\s+){0,4}(?:required|expected|needed|made)\s+within\s+(\d+)\s+days",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:payment|payable|pay)\s+(?:is\s+)?due\s+within\s+(\d+)\s+days",
        re.IGNORECASE,
    ),
    re.compile(
        r"\bNet\s+(\d+)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:payment\s*)?terms?\s*[:\s]+(\d+)\s*days",
        re.IGNORECASE,
    ),
    re.compile(
        r"due\s+in\s+(\d+)\s+days",
        re.IGNORECASE,
    ),
    re.compile(
        r"within\s+(\d+)\s+days\s+(?:of|from)",
        re.IGNORECASE,
    ),
    re.compile(
        r"payable\s+within\s+(\d+)\s+days",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:payment|payable)\s+(?:\w+\s+){0,4}within\s+(\d+)\s+days",
        re.IGNORECASE,
    ),
    re.compile(
        r"within\s+(\d+)\s+days",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:payment|payable)\s+(?:\w+\s+){0,3}(?:immediately|upon\s+receipt)",
        re.IGNORECASE,
    ),
]


def find_payment_terms(text: str) -> Optional[str]:
    for pat in _PAYMENT_TERMS_PATTERNS:
        m = pat.search(text)
        if m:
            if m.lastindex is None or m.lastindex == 0:
                return "0"
            return m.group(1)
    return None


_PO_DATE_PATTERNS = [
    re.compile(
        r"(?:p\.?o\.?\s*date|order\s*date|purchase\s*order\s*date)"
        r"[\s.:]*\n?([\s\S]{0,200})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:date\s*(?:of\s*)?(?:order|purchase))"
        r"[\s.:]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
        re.IGNORECASE,
    ),
]


def find_po_date(text: str) -> Optional[str]:
    for pat in _PO_DATE_PATTERNS:
        m = pat.search(text)
        if m:
            raw = m.group(1).strip()
            for dp in _DATE_PATTERNS:
                dm = dp.search(raw)
                if dm:
                    return dm.group(1)
    return None


_QUOTE_NUM_PATTERNS = [
    re.compile(
        r"(?:quote|quotation|qut)\s*(?:no|number|#|num|ref)?\s*[:\s]*\n?\s*"
        r"([\w\-/]{2,30})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:quote\s*date)\s*[:\s]*\n?\s*"
        r"([\w\-/]{2,30})",
        re.IGNORECASE,
    ),
]


def find_quote_number(text: str) -> Optional[str]:
    for pat in _QUOTE_NUM_PATTERNS:
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]", val):
                continue
            return val
    return None


_RECIPIENT_PATTERNS_PO = [
    re.compile(
        r"(?:recipient|supplier|vendor|deliver\s*to)\s*[:\s]*\n\s*(.+?)(?:\n|$)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:recipient|supplier|vendor)\s*[:\s]+(.+?)(?:\n|$)",
        re.IGNORECASE,
    ),
]


def find_recipient_po(text: str) -> Optional[str]:
    for pat in _RECIPIENT_PATTERNS_PO:
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if val and len(val) > 1:
                val = re.sub(r"\s+\S+@\S+", "", val).strip()
                val = re.sub(r"\s+[\d\+\-\(\)]{7,}", "", val).strip()
                return val
    return None


_DISCOUNT_PATTERNS_PO = [
    re.compile(
        rf"(?:discount)\s*[:\s]*\n?\s*({_CURR}?\s*[\d,]+\.?\d*)",
        re.IGNORECASE,
    ),
]


def find_discount_po(text: str) -> Optional[str]:
    return _extract_amount(_DISCOUNT_PATTERNS_PO, text)


_DELIVERY_ADDR_PATTERN = re.compile(
    r"(?:deliver\s*(?:to|y)\s*(?:address)?|ship\s*(?:to|ping)\s*(?:address)?)"
    r"\s*[:\s]*\n?((?:.+\n?){1,6})",
    re.IGNORECASE,
)


def find_delivery_address(text: str) -> Optional[str]:
    m = _DELIVERY_ADDR_PATTERN.search(text)
    if m:
        block = m.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()][:5]
        return ", ".join(lines)
    return None


@dataclass
class TextBlock:

    text: str
    x0: float
    y0: float
    x1: float
    y1: float
    page: int = 0

    @property
    def mid_x(self) -> float:
        return (self.x0 + self.x1) / 2

    @property
    def mid_y(self) -> float:
        return (self.y0 + self.y1) / 2

    @property
    def width(self) -> float:
        return self.x1 - self.x0

    @property
    def height(self) -> float:
        return self.y1 - self.y0

    def __repr__(self) -> str:
        return (
            f"TextBlock({self.text!r}, "
            f"x={self.x0:.0f}-{self.x1:.0f}, "
            f"y={self.y0:.0f}-{self.y1:.0f})"
        )


def _block_collapse_spaced_text(text: str) -> str:

    def _collapse(m: re.Match) -> str:
        return m.group(0).replace(" ", "")

    return re.sub(
        r"(?<!\S)([A-Za-z0-9:,] ){1,}[A-Za-z0-9:,](?!\S)",
        _collapse,
        text,
    )


def _block_clean_text(text: str) -> str:
    text = text.replace("\ufffd", "\u00a3")
    text = _block_collapse_spaced_text(text)
    return text.strip()


class PDFBlockParser:

    @staticmethod
    def extract_blocks(file_path: str) -> list[TextBlock]:

        blocks: list[TextBlock] = []

        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                page_dict = page.get_text("dict")

                for block in page_dict.get("blocks", []):
                    if block.get("type") != 0:
                        continue

                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")

                        line_text = _block_clean_text(line_text)
                        if not line_text:
                            continue

                        bbox = line.get("bbox", block.get("bbox", [0, 0, 0, 0]))

                        text_stripped = line_text.strip()
                        box_width = bbox[2] - bbox[0]
                        if len(text_stripped) <= 1 and box_width > 40:
                            continue
                        blocks.append(TextBlock(
                            text=line_text,
                            x0=bbox[0],
                            y0=bbox[1],
                            x1=bbox[2],
                            y1=bbox[3],
                            page=page_num,
                        ))

        blocks.sort(key=lambda b: (b.page, b.y0, b.x0))
        return blocks

    @staticmethod
    def page_width(file_path: str) -> float:

        with fitz.open(file_path) as doc:
            if doc.page_count > 0:
                return doc[0].rect.width
        return 595.0


ROW_Y_TOLERANCE = 6.0
COLUMN_GAP_RATIO = 0.20
TABLE_HEADER_KEYWORDS = {
    "description", "item", "particular", "particulars", "product",
    "service", "details", "qty", "quantity", "units", "count",
    "price", "rate", "unit price", "cost", "amount", "subtotal",
    "sub-total", "total", "line total", "extended",
}
TOTALS_KEYWORDS = {
    "sub-total", "sub total", "subtotal", "tax", "vat", "gst", "total",
    "grand total", "grand-total", "amount due", "balance due", "net",
    "discount",
}
PAYMENT_KEYWORDS = {
    "bank", "account", "sort code", "iban", "swift", "bic",
    "payable to", "remit to",
}
BILL_TO_KEYWORDS = {
    "bill to", "billed to", "billing to", "invoice to", "invoiced to",
    "sold to", "customer", "customer information", "ordered by", "buyer",
}

_LA_CURR = r"[£$€¥₹\u00a3]"
_LA_MONEY_RE = re.compile(rf"{_LA_CURR}\s*[\d,]+\.?\d*")


@dataclass
class LayoutRow:

    blocks: list[TextBlock] = field(default_factory=list)

    @property
    def y(self) -> float:
        return min(b.y0 for b in self.blocks) if self.blocks else 0

    @property
    def left_blocks(self) -> list[TextBlock]:
        return sorted(self.blocks, key=lambda b: b.x0)

    @property
    def full_text(self) -> str:
        return "\t".join(b.text for b in self.left_blocks)

    @property
    def texts(self) -> list[str]:
        return [b.text for b in self.left_blocks]

    def column_text(self, page_mid: float, side: str) -> str:
        if side == "left":
            parts = [b.text for b in self.left_blocks if b.mid_x < page_mid]
        else:
            parts = [b.text for b in self.left_blocks if b.mid_x >= page_mid]
        return " ".join(parts)


@dataclass
class InvoiceLayout:

    vendor_zone: list[str] = field(default_factory=list)
    bill_to_zone: list[str] = field(default_factory=list)

    header_fields: dict[str, str] = field(default_factory=dict)

    table_headers: list[str] = field(default_factory=list)
    table_rows: list[list[str]] = field(default_factory=list)

    totals: dict[str, str] = field(default_factory=dict)

    payment_zone: list[str] = field(default_factory=list)

    reconstructed_text: str = ""


def _group_into_rows(
    blocks: list[TextBlock],
    y_tolerance: float = ROW_Y_TOLERANCE,
) -> list[LayoutRow]:
    if not blocks:
        return []

    rows: list[LayoutRow] = []
    current_row = LayoutRow(blocks=[blocks[0]])

    for block in blocks[1:]:
        if abs(block.y0 - current_row.blocks[0].y0) <= y_tolerance:
            current_row.blocks.append(block)
        else:
            rows.append(current_row)
            current_row = LayoutRow(blocks=[block])

    rows.append(current_row)
    return rows


def _is_table_header_row(row: LayoutRow) -> bool:
    matches = 0
    for block in row.blocks:
        block_lower = block.text.strip().lower()
        if block_lower in TABLE_HEADER_KEYWORDS:
            matches += 1
            continue
        for kw in TABLE_HEADER_KEYWORDS:
            if kw in block_lower:
                matches += 1
                break
    return matches >= 2


def _has_label(text: str, keywords: set[str]) -> Optional[str]:
    lower = text.strip().lower().rstrip(":")
    for kw in keywords:
        if lower.startswith(kw):
            return kw
    return None


def _detect_page_mid(blocks: list[TextBlock], page_width: float) -> float:
    if not blocks:
        return page_width / 2

    max_y = max(b.y1 for b in blocks)
    top_blocks = [b for b in blocks if b.y0 < max_y * 0.35]

    if len(top_blocks) < 2:
        return page_width / 2

    sorted_blocks = sorted(top_blocks, key=lambda b: b.x0)
    best_gap = 0
    best_mid = page_width / 2

    for i in range(len(sorted_blocks) - 1):
        gap_start = sorted_blocks[i].x1
        gap_end = sorted_blocks[i + 1].x0
        gap = gap_end - gap_start

        if gap > best_gap and gap > page_width * COLUMN_GAP_RATIO:
            best_gap = gap
            best_mid = (gap_start + gap_end) / 2

    return best_mid


class LayoutAnalyzer:

    def analyze(
        self,
        blocks: list[TextBlock],
        page_width: float = 595.0,
    ) -> InvoiceLayout:
        if not blocks:
            return InvoiceLayout()

        layout = InvoiceLayout()
        rows = _group_into_rows(blocks)
        page_mid = _detect_page_mid(blocks, page_width)

        logger.debug(
            "Layout: %d blocks -> %d rows, page_mid=%.0f",
            len(blocks), len(rows), page_mid,
        )

        table_start = None
        for i, row in enumerate(rows):
            if _is_table_header_row(row):
                table_start = i
                layout.table_headers = [b.text.strip() for b in row.left_blocks]
                break

        table_end = None
        if table_start is not None:
            table_end = self._find_table_end(rows, table_start)

        header_rows = rows[:table_start] if table_start else rows
        self._process_header_zone(header_rows, page_mid, layout, all_rows=rows)

        if table_start is not None and table_end is not None:
            self._process_table_zone(rows, table_start, table_end, layout)

        footer_start = table_end if table_end else (table_start or len(rows))
        if footer_start < len(rows):
            self._process_footer_zone(rows[footer_start:], layout)

        layout.reconstructed_text = self._reconstruct_text(layout, rows, page_mid)

        return layout


    def _process_header_zone(
        self,
        rows: list[LayoutRow],
        page_mid: float,
        layout: InvoiceLayout,
        all_rows: list[LayoutRow] | None = None,
    ) -> None:
        left_lines_before_bt: list[str] = []
        left_lines_after_bt: list[str] = []
        right_lines_before_bt: list[str] = []
        right_lines_after_bt: list[str] = []
        bill_to_label_side: str | None = None
        bill_to_found = False
        bill_to_label_x: float = 0.0

        _skip_vendor = re.compile(
            r"^(invoice|bill|receipt|statement|tax\s*invoice|credit\s*note"
            r"|invoice\s*no|inv\s*no|p\.?o\.?\s*(?:no|ref)|date|due\s*date"
            r"|purchase\s*order|purchase|order|p\.?o\.?\s*date|quote"
            r"|description\s*of\s*service"
            r"|payment\s*(?:instruction|info|method|term))",
            re.IGNORECASE,
        )
        _date_line = re.compile(
            r"^\d{1,2}[/\-]\d{1,2}[/\-]"
            r"|^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
            r"|^\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)",
            re.IGNORECASE,
        )

        for row_i, row in enumerate(rows):
            left_text = row.column_text(page_mid, "left")
            right_text = row.column_text(page_mid, "right")

            extracted_left = False
            extracted_right = False
            for block in row.left_blocks:
                if self._try_extract_header_field(block.text, row, layout,
                                                  all_rows=all_rows, row_idx=row_i):
                    if block.mid_x < page_mid:
                        extracted_left = True
                    else:
                        extracted_right = True

            for block in row.left_blocks:
                if _has_label(block.text.strip().lower(), BILL_TO_KEYWORDS):
                    if block.mid_x < page_mid:
                        bill_to_label_side = "left"
                    else:
                        bill_to_label_side = "right"
                    bill_to_label_x = block.x0
                    bill_to_found = True
                    break

            full = row.full_text.lower()
            if _has_label(full, BILL_TO_KEYWORDS):
                _header_field_labels = re.compile(
                    r"^(invoice\s*no|inv\s*no|p\.?o\.?\s*no|date|po\s*no"
                    r"|po\s*ref)",
                    re.IGNORECASE,
                )
                for block in row.left_blocks:
                    bt = block.text.strip()
                    bt_lower = bt.lower()
                    if _has_label(bt_lower, BILL_TO_KEYWORDS):
                        colon_match = re.search(r":\s*(.+)", bt)
                        if colon_match:
                            inline_val = colon_match.group(1).strip()
                            if inline_val:
                                left_lines_after_bt.append(inline_val)
                        continue
                    if _header_field_labels.match(bt_lower):
                        continue
                    if any(val and val == bt for val in layout.header_fields.values()):
                        continue
                    if block.mid_x < page_mid:
                        left_lines_after_bt.append(bt)
                    else:
                        right_lines_after_bt.append(bt)
                continue

            if not extracted_left:
                if bill_to_found and bill_to_label_side == "left":
                    bt_threshold = bill_to_label_x - 50
                    for block in row.left_blocks:
                        if block.mid_x >= page_mid:
                            continue
                        bt = block.text.strip()
                        if not bt or _skip_vendor.match(bt) or _date_line.match(bt):
                            continue
                        if bill_to_label_x > 100 and block.x0 < bt_threshold:
                            left_lines_before_bt.append(bt)
                        else:
                            left_lines_after_bt.append(bt)
                elif left_text:
                    stripped = left_text.strip()
                    if _date_line.match(stripped):
                        if "invoice_date" not in layout.header_fields:
                            dates = re.findall(
                                r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}", stripped,
                            )
                            if dates:
                                layout.header_fields["invoice_date"] = dates[0]
                            else:
                                layout.header_fields["invoice_date"] = stripped
                    elif not _skip_vendor.match(stripped):
                        left_lines_before_bt.append(left_text)

            if right_text:
                stripped_right = right_text.strip()
                skip_right = False
                for val in layout.header_fields.values():
                    if val and val in right_text:
                        skip_right = True
                        break
                if skip_right:
                    pass
                elif _date_line.match(stripped_right):
                    if "invoice_date" not in layout.header_fields:
                        layout.header_fields["invoice_date"] = stripped_right
                else:
                    if bill_to_found:
                        right_lines_after_bt.append(right_text)
                    else:
                        right_lines_before_bt.append(right_text)

        right_lines = right_lines_before_bt + right_lines_after_bt
        raw_right = [r for r in right_lines if r.strip()]

        _skip_bill_to = re.compile(
            r"^(invoice\s*to|bill\s*to|billed\s*to|billing\s*to|sold\s*to|customer\s*(?:information)?)"
            r"|^(invoice\s*no|inv\s*no|p\.?o\.?\s*no|date|due\s*date"
            r"|purchase\s*order|purchase|order|p\.?o\.?\s*date|quote)",
            re.IGNORECASE,
        )

        def _clean_zone(lines: list[str]) -> list[str]:
            return [l for l in lines if l.strip() and not _skip_bill_to.match(l.strip())]

        if bill_to_label_side == "left":
            layout.vendor_zone = _clean_zone(left_lines_before_bt)
            layout.bill_to_zone = _clean_zone(left_lines_after_bt)
            if not layout.bill_to_zone and raw_right and bill_to_label_x > 0:
                nearby = [r for r in raw_right
                          if r.strip() and not _skip_bill_to.match(r.strip())]
                layout.bill_to_zone = nearby
            if not layout.vendor_zone and raw_right:
                header_vals = set(v.lower() for v in layout.header_fields.values() if v)
                _header_label_re = re.compile(
                    r"^(invoice\s*no|inv\s*no|invoice\s*date|date|po\s*no|due\s*date)",
                    re.IGNORECASE,
                )
                filtered_right = []
                for r in raw_right:
                    rs = r.strip()
                    if not rs:
                        continue
                    if _skip_bill_to.match(rs):
                        continue
                    if _header_label_re.match(rs):
                        continue
                    if rs.lower() in header_vals:
                        continue
                    filtered_right.append(rs)
                layout.vendor_zone = filtered_right
        else:
            all_left = left_lines_before_bt + left_lines_after_bt
            layout.vendor_zone = _clean_zone(all_left)
            if bill_to_label_side == "right" and right_lines_after_bt:
                layout.bill_to_zone = _clean_zone(right_lines_after_bt)
            else:
                layout.bill_to_zone = _clean_zone(raw_right)

        for line in right_lines:
            stripped = line.strip()
            if re.match(r"p\.?o\.?\s*(?:no|ref(?:erence)?)", stripped, re.IGNORECASE) and ":" in stripped:
                m = re.search(r":\s*(\S+)", stripped)
                if m and "po_number" not in layout.header_fields:
                    layout.header_fields["po_number"] = m.group(1)

    def _try_extract_header_field(
        self,
        text: str,
        row: LayoutRow,
        layout: InvoiceLayout,
        all_rows: list[LayoutRow] | None = None,
        row_idx: int = -1,
    ) -> bool:

        field_labels = {
            r"inv(?:oice)?\s*(?:no|number|#|num)": "invoice_number",
            r"no(?:\.|\s*#)": "invoice_number",
            r"p\.?o\.?\s*date": "po_date",
            r"p\.?o\.?\s*(?:no|number|#|num|ref(?:erence)?)?": "po_number",
            r"(?:invoice\s*)?date\s*(?:issued)?|issue\s*date": "invoice_date",
            r"p\.?o\.?\s*date|order\s*date": "po_date",
            r"due\s*date|payment\s*due": "due_date",
            r"quote\s*(?:no|number|#|num|ref|date)": "quote_number",
            r"quotation\s*(?:no|number|#)?": "quote_number",
        }

        lower = text.strip().lower().rstrip(":.")
        for pattern, field_name in field_labels.items():
            if re.match(pattern, lower):
                sep_match = re.search(r"[.:]\s*(.+)", text)
                if sep_match:
                    val = sep_match.group(1).strip()
                    if val and re.search(r"\d", val):
                        if field_name in ("invoice_number", "po_number"):
                            parts = re.split(r"\s*/\s*", val)
                            val = parts[0].strip()
                            if len(parts) > 1 and "invoice_date" not in layout.header_fields:
                                remainder = parts[1].strip()
                                date_m = re.search(
                                    r"\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\.?\s*,?\s*\d{4}"
                                    r"|\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}",
                                    remainder, re.IGNORECASE,
                                )
                                if date_m:
                                    layout.header_fields["invoice_date"] = date_m.group(0).strip()
                        layout.header_fields[field_name] = val
                        return True

                blocks = row.left_blocks
                for i, b in enumerate(blocks):
                    if b.text == text:
                        for j in range(i + 1, len(blocks)):
                            next_text = blocks[j].text.strip()
                            if next_text in (":", ".", ":."):
                                continue
                            if next_text and re.search(r"\d", next_text):
                                layout.header_fields[field_name] = next_text
                                return True
                            break
                label_x = 0.0
                for b in row.left_blocks:
                    if b.text == text:
                        label_x = b.mid_x
                        break
                if all_rows and 0 <= row_idx < len(all_rows) - 1:
                    max_x_dist = 150.0
                    for offset in range(1, min(3, len(all_rows) - row_idx)):
                        candidate_row = all_rows[row_idx + offset]
                        best_block = None
                        best_dist = float("inf")
                        for nb in candidate_row.left_blocks:
                            nt = nb.text.strip()
                            if nt and nt not in (":", ".") and re.search(r"\d", nt):
                                dist = abs(nb.mid_x - label_x)
                                if dist < best_dist and dist <= max_x_dist:
                                    best_dist = dist
                                    best_block = nt
                        if best_block:
                            layout.header_fields[field_name] = best_block
                            return True
                return True
        return False


    def _find_table_end(
        self,
        rows: list[LayoutRow],
        table_start: int,
    ) -> int:

        typical_gap = 50.0
        if table_start + 2 < len(rows):
            first_gap = rows[table_start + 2].y - rows[table_start + 1].y
            if first_gap > 5:
                typical_gap = first_gap
        gap_threshold = max(typical_gap * 2.0, 60.0)

        for i in range(table_start + 1, len(rows)):
            for block in rows[i].left_blocks:
                text_lower = block.text.strip().lower().rstrip(":")
                text_lower = re.sub(r"\s*\(?[\d.]+%\)?", "", text_lower).strip()

                if text_lower in TOTALS_KEYWORDS and i > table_start + 1:
                    return i

                if _has_label(text_lower, PAYMENT_KEYWORDS):
                    return i

            if i > table_start + 2:
                prev_y = rows[i - 1].y
                curr_y = rows[i].y
                if curr_y - prev_y > gap_threshold:
                    return i

        return len(rows)

    def _process_table_zone(
        self,
        rows: list[LayoutRow],
        table_start: int,
        table_end: int,
        layout: InvoiceLayout,
    ) -> None:

        if table_start + 1 >= table_end:
            return

        header_row = rows[table_start]
        header_xs = [b.mid_x for b in header_row.left_blocks]
        n_cols = len(header_xs)

        for i in range(table_start + 1, table_end):
            row = rows[i]
            row_blocks = row.left_blocks

            if not row_blocks:
                continue

            first_text = row_blocks[0].text.strip().lower()
            if first_text in TOTALS_KEYWORDS or first_text.startswith("sub"):
                continue

            cells = [""] * n_cols
            for block in row_blocks:
                best_col = min(
                    range(n_cols),
                    key=lambda c: abs(block.mid_x - header_xs[c]),
                )
                if cells[best_col]:
                    cells[best_col] += " " + block.text
                else:
                    cells[best_col] = block.text

            if not cells[0].strip():
                if layout.table_rows:
                    for c in range(1, n_cols):
                        if cells[c].strip() and not layout.table_rows[-1][c].strip():
                            layout.table_rows[-1][c] = cells[c]
                continue

            has_numeric = any(
                _LA_MONEY_RE.search(cells[c]) or re.match(r"^\d+$", cells[c].strip())
                for c in range(1, n_cols)
                if cells[c].strip()
            )

            if not has_numeric and layout.table_rows:
                prev_has_amounts = any(
                    _LA_MONEY_RE.search(layout.table_rows[-1][c])
                    for c in range(1, n_cols)
                    if layout.table_rows[-1][c].strip()
                )
                y_gap = rows[i].y - rows[i - 1].y if i > table_start else 0
                if prev_has_amounts and y_gap > 20:
                    layout.table_rows.append(cells)
                else:
                    layout.table_rows[-1][0] += " " + cells[0].strip()
            else:
                layout.table_rows.append(cells)


    def _process_footer_zone(
        self,
        rows: list[LayoutRow],
        layout: InvoiceLayout,
    ) -> None:

        for row in rows:
            for block in row.left_blocks:
                self._try_extract_total_from_block(block, row, layout)

            for block in row.left_blocks:
                text_lower = block.text.strip().lower()
                if _has_label(text_lower, PAYMENT_KEYWORDS):
                    layout.payment_zone.append(row.full_text)
                    break

    def _try_extract_total_from_block(
        self,
        block: TextBlock,
        row: LayoutRow,
        layout: InvoiceLayout,
    ) -> None:

        total_labels = {
            "sub-total": "subtotal",
            "sub total": "subtotal",
            "subtotal": "subtotal",
            "net": "subtotal",
            "discount": "discount",
            "tax": "tax",
            "vat": "tax",
            "gst": "tax",
            "total": "total",
            "grand total": "total",
            "grand-total": "total",
            "amount due": "total",
            "balance due": "total",
        }

        text = block.text.strip()
        text_lower = text.lower().rstrip(":")
        text_lower = re.sub(r"\s*\(?[\d.]+%\)?", "", text_lower).strip()

        matched_field = None
        for kw, field_name in total_labels.items():
            if text_lower == kw:
                matched_field = field_name
                break

        if matched_field is None:
            for kw, field_name in total_labels.items():
                if text_lower.startswith(kw):
                    money = _LA_MONEY_RE.search(text)
                    if money:
                        if (field_name == "subtotal"
                                and "subtotal" in layout.totals
                                and "tax" in layout.totals):
                            layout.totals["total"] = money.group(0)
                        else:
                            layout.totals[field_name] = money.group(0)
                        return
                    matched_field = field_name
                    break
            if matched_field is None:
                return

        if (matched_field == "subtotal"
                and "subtotal" in layout.totals
                and "tax" in layout.totals):
            matched_field = "total"

        for other in row.left_blocks:
            if other.x0 > block.x1 and _LA_MONEY_RE.search(other.text):
                layout.totals[matched_field] = other.text.strip()
                return


    def _reconstruct_text(
        self,
        layout: InvoiceLayout,
        rows: list[LayoutRow],
        page_mid: float,
    ) -> str:

        parts: list[str] = []

        if layout.vendor_zone:
            parts.extend(layout.vendor_zone)
            parts.append("")

        if layout.bill_to_zone:
            parts.append("Bill To:")
            parts.extend(layout.bill_to_zone)
            parts.append("")

        field_labels = {
            "invoice_number": "Invoice No",
            "po_number": "PO No",
            "invoice_date": "Date Issued",
            "po_date": "PO Date",
            "due_date": "Due Date",
            "quote_number": "Quote Number",
        }
        for key, label in field_labels.items():
            val = layout.header_fields.get(key, "")
            if val:
                parts.append(f"{label}: {val}")
        if layout.header_fields:
            parts.append("")

        if layout.table_headers:
            parts.append("\t".join(layout.table_headers))
            for row_cells in layout.table_rows:
                parts.append("\t".join(row_cells))
            parts.append("")

        if layout.totals:
            for key in ("subtotal", "discount", "tax", "total"):
                val = layout.totals.get(key, "")
                if val:
                    label = {"subtotal": "Sub-Total", "discount": "Discount",
                             "tax": "Tax", "total": "TOTAL"}[key]
                    parts.append(f"{label}: {val}")
            parts.append("")

        if layout.payment_zone:
            parts.extend(layout.payment_zone)

        text = "\n".join(parts)

        if len(text.strip()) < 50:
            text = "\n".join(row.full_text for row in rows)

        return text


def extract_table_from_boxes(
    word_boxes: list[dict],
    y_tolerance: int = LINE_ITEM_ROW_Y_TOLERANCE,
    x_gap: int = LINE_ITEM_COL_X_GAP,
) -> list[dict[str, str]]:

    if not word_boxes:
        return []

    boxes = sorted(word_boxes, key=lambda b: (b["y"], b["x"]))

    rows: list[list[dict]] = []
    current_row: list[dict] = [boxes[0]]

    for box in boxes[1:]:
        prev_mid_y = current_row[-1]["y"] + current_row[-1]["h"] / 2
        curr_mid_y = box["y"] + box["h"] / 2

        if abs(curr_mid_y - prev_mid_y) <= y_tolerance:
            current_row.append(box)
        else:
            rows.append(current_row)
            current_row = [box]
    rows.append(current_row)

    text_rows: list[list[str]] = []
    for row in rows:
        row_sorted = sorted(row, key=lambda b: b["x"])
        columns: list[list[str]] = [[row_sorted[0]["text"]]]

        for i in range(1, len(row_sorted)):
            prev_right = row_sorted[i - 1]["x"] + row_sorted[i - 1]["w"]
            curr_left = row_sorted[i]["x"]
            gap = curr_left - prev_right

            if gap > x_gap:
                columns.append([row_sorted[i]["text"]])
            else:
                columns[-1].append(row_sorted[i]["text"])

        text_rows.append([" ".join(col) for col in columns])

    header_idx, col_map = _find_header(text_rows)
    if header_idx is None or col_map is None:
        logger.debug("Could not identify table header from OCR boxes.")
        return []

    items: list[dict[str, str]] = []
    for row_cells in text_rows[header_idx + 1 :]:
        item = _parse_data_row(row_cells, col_map)
        if item:
            items.append(item)

    return items


_HEADER_KEYWORDS = {
    "item": ["description", "item", "product", "service", "particular", "details"],
    "quantity": ["qty", "quantity", "units", "no", "count"],
    "unit_price": ["price", "rate", "unit price", "unit cost", "cost", "each"],
    "amount": ["amount", "subtotal", "sub-total", "total", "line total", "ext", "extended"],
}


def _find_header(
    text_rows: list[list[str]],
) -> tuple[Optional[int], Optional[dict[str, int]]]:

    for idx, row in enumerate(text_rows):
        col_map: dict[str, int] = {}
        for col_idx, cell in enumerate(row):
            cell_lower = cell.lower().strip()
            for field, keywords in _HEADER_KEYWORDS.items():
                if any(kw in cell_lower for kw in keywords):
                    if field not in col_map:
                        col_map[field] = col_idx

        if "item" in col_map and len(col_map) >= 2:
            return idx, col_map

    return None, None


_TP_MONEY_RE = re.compile(r"[£$€¥₹]?\s*[\d,]+\.?\d*")
_TP_PHONE_RE = re.compile(
    r"(?:\+\d{1,3}[\s\-]?)?"
    r"(?:\(?\d{2,5}\)?[\s\-]?)?"
    r"\d{3,4}[\s\-]\d{3,4}"
    r"(?:[\s\-]\d{2,4})?",
)
_TP_SKIP_LABELS = {"total", "sub-total", "subtotal", "tax", "vat", "gst", "discount",
                "shipping", "freight", "grand total", "net", "balance", "total amount",
                "amount due", "vat (20%)", "vat (10%)", "vat (5%)"}


def _parse_data_row(
    row_cells: list[str],
    col_map: dict[str, int],
) -> Optional[dict[str, str]]:

    item_idx = col_map.get("item")
    if item_idx is None or item_idx >= len(row_cells):
        return None

    description = row_cells[item_idx].strip()
    desc_lower = description.lower().rstrip(":")
    if not description or desc_lower in _TP_SKIP_LABELS:
        return None

    _skip_starts = ("total amount", "vat ", "vat(", "sales tax", "grand total",
                    "tax ", "tax(", "tax:",
                    "account no", "account number", "sort code", "iban",
                    "swift", "bic", "bank", "payment", "pay by",
                    "thank you", "account name", "this invoice",
                    "invoice reference", "po no", "po ref", "po:")
    if any(desc_lower.startswith(prefix) for prefix in _skip_starts):
        return None

    _skip_contains = ("sort code", "iban:", "swift", "account number",
                      "invoice references", "payment due")
    if any(kw in desc_lower for kw in _skip_contains):
        return None

    if description.startswith("+") or _TP_PHONE_RE.fullmatch(description):
        return None

    result: dict[str, str] = {"item": description}

    for field in ("quantity", "unit_price", "amount"):
        idx = col_map.get(field)
        if idx is not None and idx < len(row_cells):
            result[field] = row_cells[idx].strip()
        else:
            result[field] = ""

    qty_val = result.get("quantity", "")
    if qty_val:
        _qty_money_re = re.compile(
            r"^(\d+)\s+([£$€¥₹]?\s*[\d,]+\.?\d*)$"
        )
        m = _qty_money_re.match(qty_val)
        if m:
            result["quantity"] = m.group(1)
            money_part = m.group(2)
            if not result.get("amount"):
                result["amount"] = money_part
            elif not result.get("unit_price"):
                result["unit_price"] = money_part

    has_number = any(
        _TP_MONEY_RE.search(result.get(f, ""))
        for f in ("quantity", "unit_price", "amount")
    )
    if not has_number:
        return None

    for f in ("quantity", "unit_price", "amount"):
        val = result.get(f, "").strip()
        if val and (val.startswith("+") or _TP_PHONE_RE.fullmatch(val)):
            return None

    return result


_DATE_RE = re.compile(
    r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}"
    r"|\d{4}-\d{2}-\d{2}"
    r"|\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*[\s,]+\d{4}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{1,2},?\s+\d{4}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}",
    re.IGNORECASE,
)
_DIGITS_RE = re.compile(r"\d{2,}")
_MONEY_STRICT_RE = re.compile(r"^[£$€¥₹\u00a3]?\s*[\d,]+\.\d{2}$")
_MONEY_LOOSE_RE = re.compile(r"[£$€¥₹\u00a3]?\s*[\d,]+\.?\d*")


def _has_digits(s: str) -> bool:
    return bool(_DIGITS_RE.search(s))


def _is_clean_ner(s: str) -> bool:
    if "\n" in s:
        return False
    if len(s.strip()) < 2 or len(s.strip()) > 200:
        return False
    label_words = {"sub-total", "subtotal", "tax", "total", "invoice", "purchase order",
                   "description", "payable", "notes", "bank", "account"}
    if s.strip().lower().rstrip(":") in label_words:
        return False
    return True


def _looks_like_date(s: str) -> bool:
    return _is_clean_ner(s) and bool(_DATE_RE.search(s))


def _looks_like_money(s: str) -> bool:
    cleaned = s.strip()
    if not cleaned or not _is_clean_ner(s):
        return False
    if _MONEY_STRICT_RE.match(cleaned):
        return True
    return bool(_MONEY_LOOSE_RE.search(cleaned)) and _has_digits(cleaned)


def _looks_like_address(s: str) -> bool:
    if not _is_clean_ner(s):
        return False
    clean = s.strip()
    alpha_count = sum(1 for c in clean if c.isalpha())
    return len(clean) > 10 and alpha_count > 5


def _clean_value(s: str) -> str:
    return " ".join(s.split()).strip()


class _FPNamespace:
    find_invoice_number = staticmethod(find_invoice_number)
    find_po_number = staticmethod(find_po_number)
    find_all_dates = staticmethod(find_all_dates)
    find_invoice_date = staticmethod(find_invoice_date)
    find_due_date = staticmethod(find_due_date)
    detect_currency = staticmethod(detect_currency)
    find_total = staticmethod(find_total)
    find_subtotal = staticmethod(find_subtotal)
    find_tax = staticmethod(find_tax)
    find_discount = staticmethod(find_discount)
    find_bill_to = staticmethod(find_bill_to)
    find_ship_to = staticmethod(find_ship_to)
    find_recipient = staticmethod(find_recipient)
    find_line_items_from_text = staticmethod(find_line_items_from_text)
    find_payment_terms = staticmethod(find_payment_terms)

fp = _FPNamespace()


class InvoiceExtractor(BaseExtractor):

    def extract(self, context: dict[str, Any]) -> dict[str, Any]:
        text: str = context.get("text", "")
        spacy_ents: dict[str, list[str]] = context.get("spacy_ents", {})
        nuext_ents: dict[str, list[str]] = context.get("nuext_ents", {})
        word_boxes: list[dict] = context.get("word_boxes", [])
        layout_fields: dict[str, Any] = context.get("layout_fields", {})

        result: dict[str, Any] = {"document_type": "invoice"}

        result["vendor_name"] = self._extract_vendor_name(text, spacy_ents, nuext_ents, layout_fields)
        result["vendor_address"] = self._extract_vendor_address(text, spacy_ents, nuext_ents, layout_fields)
        result["invoice_number"] = self._extract_invoice_number(text, spacy_ents, nuext_ents, layout_fields)
        result["po_number"] = self._extract_po_number(text, spacy_ents, nuext_ents, layout_fields)
        result["invoice_date"] = self._extract_invoice_date(text, spacy_ents, nuext_ents, layout_fields)
        result["due_date"] = self._extract_due_date(text, spacy_ents, nuext_ents, layout_fields)
        result["currency"] = self._extract_currency(text, nuext_ents)
        result["bill_to"] = self._extract_bill_to(text, spacy_ents, nuext_ents, layout_fields)
        result["ship_to"] = self._extract_ship_to(text, spacy_ents, nuext_ents)
        result["subtotal"] = self._extract_subtotal(text, spacy_ents, nuext_ents, layout_fields)
        result["tax"] = self._extract_tax(text, spacy_ents, nuext_ents, layout_fields)
        result["total_amount"] = self._extract_total(text, spacy_ents, nuext_ents, layout_fields)
        result["discount"] = self._extract_discount(text, layout_fields)
        result["line_items"] = self._extract_line_items(text, word_boxes, layout_fields)
        result["payment_info"] = self._extract_payment_info(text, nuext_ents)

        full_text = layout_fields.get("_flat_text", text)
        search_texts = [text] if full_text == text else [text, full_text]
        result["payment_terms_days"] = ""
        result["recipient"] = ""
        for t in search_texts:
            if not result["payment_terms_days"]:
                terms = fp.find_payment_terms(t)
                if terms:
                    result["payment_terms_days"] = terms
            if not result["recipient"]:
                recipient = fp.find_recipient(t)
                if recipient:
                    result["recipient"] = recipient

        flat_text = layout_fields.get("_flat_text", text)
        ocr_text = layout_fields.get("_ocr_text", "")
        parts = [text]
        if flat_text != text:
            parts.append(flat_text)
        if ocr_text:
            parts.append(ocr_text)
        result["_raw_text"] = "\n".join(parts)

        return result

    def _extract_vendor_name(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_vendor = layout_fields.get("vendor_name", "")
        if layout_vendor and _is_clean_ner(layout_vendor):
            candidates.append(Candidate(_clean_value(layout_vendor), base_score("layout"), "layout"))
        for v in nuext_ents.get("vendor name", []):
            if _is_clean_ner(v) and not _looks_like_money(v) and not _looks_like_date(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for org in spacy_ents.get("ORG", []):
            if _is_clean_ner(org) and not _looks_like_money(org) and not _looks_like_date(org):
                candidates.append(Candidate(_clean_value(org), base_score("spacy"), "spacy"))
        payable_name = self._name_from_payable_to(text)
        if payable_name:
            candidates.append(Candidate(payable_name, base_score("heuristic") + 0.05, "heuristic"))
        logger.debug("vendor_name candidates: %s", candidates)
        return pick_best(candidates)

    @staticmethod
    def _name_from_payable_to(text):
        m = re.search(r"(?:payable\s*to|remit\s*to)\s*[:\s]*\n?((?:.+\n?){1,5})", text, re.IGNORECASE)
        if not m:
            return ""
        lines = [l.strip() for l in m.group(1).split("\n") if l.strip()]
        if lines:
            first = lines[0]
            if not re.search(r"\d{3,}", first) and len(first) > 2:
                return _clean_value(first)
        return ""

    def _extract_vendor_address(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        bill_to_block = fp.find_bill_to(text) or ""
        layout_addr = layout_fields.get("vendor_address", "")
        if layout_addr and _looks_like_address(layout_addr):
            candidates.append(Candidate(_clean_value(layout_addr), base_score("layout"), "layout"))
        for addr in nuext_ents.get("address", []):
            if not _looks_like_address(addr):
                continue
            clean = _clean_value(addr)
            if bill_to_block and clean in bill_to_block:
                continue
            candidates.append(Candidate(clean, base_score("nuextract"), "nuextract"))
        payable_addr = self._address_from_payable_to(text)
        if payable_addr:
            candidates.append(Candidate(payable_addr, base_score("heuristic") + 0.10, "heuristic"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    @staticmethod
    def _address_from_payable_to(text):
        m = re.search(r"(?:payable\s*to|remit\s*to)\s*[:\s]*\n?((?:.+\n?){1,10})", text, re.IGNORECASE)
        if not m:
            return ""
        lines = [l.strip() for l in m.group(1).split("\n") if l.strip()]
        addr_lines: list[str] = []
        for l in lines:
            if re.match(r"^(NOTES|PAYABLE|REMIT|PO\s*No|Invoice\s*No|Date)\s*:?\s*$", l, re.IGNORECASE):
                continue
            if re.match(r"(?:bank|account|sort|iban|swift)", l, re.IGNORECASE):
                continue
            if re.match(r"^\d{2,4}[\-\s]\d{2,4}[\-\s]\d{2,4}$", l):
                continue
            if re.match(r".*@", l) or re.match(r"(?:https?://|www\.)", l, re.IGNORECASE):
                continue
            if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", l):
                continue
            if re.match(r"^\d{4,}$", l):
                continue
            addr_lines.append(l)
        if len(addr_lines) > 1:
            return ", ".join(addr_lines[1:])
        return addr_lines[0] if addr_lines else ""

    def _extract_invoice_number(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_inv = layout_fields.get("invoice_number", "")
        if layout_inv and _has_digits(layout_inv):
            candidates.append(Candidate(_clean_value(layout_inv), base_score("layout"), "layout"))
        for v in nuext_ents.get("invoice number", []):
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        inv_num = fp.find_invoice_number(text)
        if inv_num:
            candidates.append(Candidate(inv_num, base_score("regex_context"), "regex_context"))
        candidates = [adjust_invoice_number(c) for c in candidates]
        return pick_best(candidates)

    def _extract_po_number(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_po = layout_fields.get("po_number", "")
        if layout_po and _has_digits(layout_po):
            candidates.append(Candidate(_clean_value(layout_po), base_score("layout"), "layout"))
        for v in nuext_ents.get("purchase order number", []):
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        flat_text = layout_fields.get("_flat_text", text)
        for t in ([text, flat_text] if flat_text != text else [text]):
            po = fp.find_po_number(t)
            if po:
                candidates.append(Candidate(po, base_score("regex_context"), "regex_context"))
                break
        candidates = [adjust_invoice_number(c) for c in candidates]
        return pick_best(candidates)

    def _extract_invoice_date(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_date = layout_fields.get("invoice_date", "")
        if layout_date and _looks_like_date(layout_date):
            candidates.append(Candidate(_clean_value(layout_date), base_score("layout"), "layout"))
        for d in nuext_ents.get("date", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("nuextract"), "nuextract"))
        for d in spacy_ents.get("DATE", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("spacy"), "spacy"))
        inv_date = fp.find_invoice_date(text)
        if inv_date:
            candidates.append(Candidate(inv_date, base_score("regex_context"), "regex_context"))
        all_dates = fp.find_all_dates(text)
        for d in all_dates:
            candidates.append(Candidate(d, base_score("regex_fallback"), "regex_fallback"))
        candidates = [adjust_date(c) for c in candidates]
        return pick_best(candidates)

    def _extract_due_date(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_due = layout_fields.get("due_date", "")
        if layout_due and _looks_like_date(layout_due):
            candidates.append(Candidate(_clean_value(layout_due), base_score("layout"), "layout"))
        for d in nuext_ents.get("due date", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("nuextract"), "nuextract"))
        flat_text = layout_fields.get("_flat_text", text)
        for t in ([text, flat_text] if flat_text != text else [text]):
            due = fp.find_due_date(t)
            if due:
                candidates.append(Candidate(due, base_score("regex_context"), "regex_context"))
                break
        candidates = [adjust_date(c) for c in candidates]
        return pick_best(candidates)

    def _extract_currency(self, text, nuext_ents):
        candidates: list[Candidate] = []
        for c in nuext_ents.get("currency", []):
            code = c.strip().upper()
            if code in ("GBP", "USD", "EUR", "JPY", "INR", "CAD", "AUD"):
                candidates.append(Candidate(code, base_score("nuextract"), "nuextract"))
        detected = fp.detect_currency(text)
        if detected:
            candidates.append(Candidate(detected, base_score("regex_context"), "regex_context"))
        return pick_best(candidates)

    def _extract_bill_to(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_bt = layout_fields.get("bill_to", "")
        if layout_bt and _looks_like_address(layout_bt):
            candidates.append(Candidate(_clean_value(layout_bt), base_score("layout") + 0.05, "layout"))
        for addr in nuext_ents.get("bill_to_address", []):
            if _looks_like_address(addr):
                candidates.append(Candidate(_clean_value(addr), base_score("nuextract") + 0.05, "nuextract"))
        for name in nuext_ents.get("company name", []):
            if _is_clean_ner(name):
                for addr in nuext_ents.get("bill_to_address", []):
                    if _looks_like_address(addr):
                        combined = f"{_clean_value(name)}, {_clean_value(addr)}"
                        candidates.append(Candidate(combined, base_score("nuextract"), "nuextract"))
        bill_to = fp.find_bill_to(text)
        if bill_to:
            candidates.append(Candidate(bill_to, base_score("regex_context") + 0.10, "regex_context"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    def _extract_ship_to(self, text, spacy_ents, nuext_ents):
        candidates: list[Candidate] = []
        ship = fp.find_ship_to(text)
        if ship:
            candidates.append(Candidate(ship, base_score("regex_context"), "regex_context"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    def _extract_subtotal(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_sub = layout_fields.get("subtotal", "")
        if layout_sub and _looks_like_money(layout_sub):
            candidates.append(Candidate(_clean_value(layout_sub), base_score("layout"), "layout"))
        for v in nuext_ents.get("subtotal", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        sub = fp.find_subtotal(text)
        if sub:
            candidates.append(Candidate(sub, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_tax(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_tax = layout_fields.get("tax", "")
        if layout_tax and _looks_like_money(layout_tax):
            candidates.append(Candidate(_clean_value(layout_tax), base_score("layout"), "layout"))
        for v in nuext_ents.get("tax amount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        tax = fp.find_tax(text)
        if tax:
            candidates.append(Candidate(tax, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_discount(self, text, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_disc = layout_fields.get("discount", "")
        if layout_disc and _looks_like_money(layout_disc):
            candidates.append(Candidate(_clean_value(layout_disc), base_score("layout"), "layout"))
        disc = fp.find_discount(text)
        if disc:
            candidates.append(Candidate(disc, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_total(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_total = layout_fields.get("total", "")
        if layout_total and _looks_like_money(layout_total):
            candidates.append(Candidate(_clean_value(layout_total), base_score("layout") + 0.10, "layout"))
        for v in nuext_ents.get("total amount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for v in spacy_ents.get("MONEY", []):
            v = v.strip()
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(v, base_score("spacy") - 0.10, "spacy"))
        regex_total = fp.find_total(text)
        if regex_total:
            candidates.append(Candidate(regex_total, base_score("regex_context") + 0.15, "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_line_items(self, text, word_boxes, layout_fields=None):
        layout_fields = layout_fields or {}
        scored_sets: list[tuple[float, str, list[dict[str, str]]]] = []

        layout_headers = layout_fields.get("_table_headers", [])
        layout_rows = layout_fields.get("_table_rows", [])
        if layout_headers and layout_rows:
            layout_items = self._layout_rows_to_items(layout_headers, layout_rows)
            if layout_items:
                score = self._score_line_item_set(layout_items, base_score("layout"))
                scored_sets.append((score, "layout", layout_items))

        text_items = fp.find_line_items_from_text(text)
        if text_items:
            score = self._score_line_item_set(text_items, base_score("text_table"))
            scored_sets.append((score, "text", text_items))

        flat_text = layout_fields.get("_flat_text", "")
        if flat_text and flat_text != text:
            flat_items = fp.find_line_items_from_text(flat_text)
            if flat_items:
                score = self._score_line_item_set(flat_items, base_score("text_table"))
                scored_sets.append((score, "flat_text", flat_items))

        if word_boxes:
            ocr_items = extract_table_from_boxes(word_boxes)
            if ocr_items:
                score = self._score_line_item_set(ocr_items, base_score("ocr_table"))
                scored_sets.append((score, "ocr", ocr_items))

        if not scored_sets:
            logger.warning("No line items could be extracted.")
            return []

        scored_sets.sort(key=lambda x: x[0], reverse=True)
        best_score, best_source, best_items = scored_sets[0]
        logger.debug("line_items: picked %s set with score=%.2f (%d items)", best_source, best_score, len(best_items))
        return self._clean_line_items(best_items)

    @staticmethod
    def _score_line_item_set(items, base):
        if not items:
            return 0.0
        score = base
        completeness_total = 0.0
        for item in items:
            fields_filled = sum(1 for k in ("item", "quantity", "unit_price", "amount") if item.get(k, "").strip())
            completeness_total += fields_filled / 4.0
            qty = item.get("quantity", "").strip()
            try:
                int(float(qty.replace(",", "")))
                score += 0.02
            except (ValueError, TypeError):
                score -= 0.03
            amt = item.get("amount", "").strip()
            if re.search(r"[£$€¥₹\u00a3]", amt):
                score += 0.02
        avg_completeness = completeness_total / len(items)
        score += avg_completeness * 0.15
        if len(items) >= 2:
            score += 0.05
        if len(items) >= 4:
            score += 0.03
        return min(score, 1.0)

    @staticmethod
    def _clean_line_items(items):
        _skip_exact = {"total", "sub-total", "subtotal", "tax", "vat", "gst",
                       "discount", "grand total", "net", "balance", "amount due"}
        _skip_starts = ("total amount", "vat ", "vat(", "sales tax",
                        "grand total", "tax ", "tax(", "tax:")

        filtered: list[dict[str, str]] = []
        for item in items:
            desc = item.get("item", "").strip()
            desc_lower = desc.lower().rstrip(":")
            if desc_lower in _skip_exact:
                continue
            if any(desc_lower.startswith(p) for p in _skip_starts):
                continue
            filtered.append(item)

        merged: list[dict[str, str]] = []
        i = 0
        while i < len(filtered):
            item = filtered[i]
            has_data = any(item.get(f, "").strip() for f in ("quantity", "unit_price", "amount"))
            if not has_data and i + 1 < len(filtered):
                next_item = filtered[i + 1]
                combined_desc = item.get("item", "").strip()
                next_desc = next_item.get("item", "").strip()
                if next_desc:
                    combined_desc = f"{combined_desc} - {next_desc}"
                merged.append({
                    "item": combined_desc,
                    "quantity": next_item.get("quantity", ""),
                    "unit_price": next_item.get("unit_price", ""),
                    "amount": next_item.get("amount", ""),
                })
                i += 2
            else:
                merged.append(item)
                i += 1

        cleaned: list[dict[str, str]] = []
        for item in merged:
            desc = item.get("item", "").strip()
            qty = item.get("quantity", "").strip()
            qty = re.sub(r"[£$€¥₹\u00a3]", "", qty).strip()
            try:
                qty_int = int(float(qty.replace(",", "")))
            except (ValueError, TypeError):
                qty_int = None
            cleaned.append({
                "item": desc,
                "quantity": str(qty_int) if qty_int is not None else qty,
                "unit_price": item.get("unit_price", "").strip(),
                "amount": item.get("amount", "").strip(),
            })
        return cleaned

    @staticmethod
    def _layout_rows_to_items(headers, rows):
        header_map: dict[str, str] = {}
        desc_keys = {"description", "item", "particular", "particulars", "product", "service", "details"}
        qty_keys = {"qty", "quantity", "units", "count", "months", "hours", "days"}
        price_keys = {"price", "rate", "unit price", "cost"}
        amount_keys = {"amount", "subtotal", "sub-total", "total", "line total", "extended"}

        for i, h in enumerate(headers):
            h_lower = h.strip().lower()
            if h_lower in desc_keys:
                header_map[i] = "item"
            elif h_lower in qty_keys:
                header_map[i] = "quantity"
            elif h_lower in price_keys:
                header_map[i] = "unit_price"
            elif h_lower in amount_keys:
                header_map[i] = "amount"
            else:
                if any(kw in h_lower for kw in desc_keys):
                    header_map[i] = "item"
                elif any(kw in h_lower for kw in qty_keys):
                    header_map[i] = "quantity"
                elif any(kw in h_lower for kw in price_keys):
                    header_map[i] = "unit_price"
                elif any(kw in h_lower for kw in amount_keys):
                    header_map[i] = "amount"

        if "item" not in header_map.values():
            if len(headers) >= 3:
                header_map = {0: "item", 1: "quantity", 2: "unit_price"}
                if len(headers) >= 4:
                    header_map[3] = "amount"
            elif len(headers) == 2:
                header_map = {0: "item", 1: "amount"}

        items: list[dict[str, str]] = []
        pending_desc: str = ""
        for row in rows:
            item: dict[str, str] = {"item": "", "quantity": "", "unit_price": "", "amount": ""}
            for col_idx, field_name in header_map.items():
                if col_idx < len(row):
                    item[field_name] = row[col_idx].strip()
            qty_val = item.get("quantity", "")
            if qty_val:
                _qty_money = re.match(r"^(\d+)\s+([£$€¥₹]?\s*[\d,]+\.?\d*)$", qty_val)
                if _qty_money:
                    item["quantity"] = _qty_money.group(1)
                    money = _qty_money.group(2)
                    if not item.get("amount"):
                        item["amount"] = money
                    elif not item.get("unit_price"):
                        item["unit_price"] = money
            has_values = any(item[k] for k in ("quantity", "unit_price", "amount"))
            if item["item"] and not has_values:
                if pending_desc:
                    pending_desc += " " + item["item"]
                else:
                    pending_desc = item["item"]
                continue
            if item["item"] or has_values:
                if pending_desc:
                    if item["item"]:
                        item["item"] = pending_desc + " " + item["item"]
                    else:
                        item["item"] = pending_desc
                    pending_desc = ""
                if item["item"]:
                    items.append(item)
        if pending_desc:
            items.append({"item": pending_desc, "quantity": "", "unit_price": "", "amount": ""})
        return items

    @staticmethod
    def _extract_payment_info(text, nuext_ents):
        info: dict[str, str] = {}
        for bn in nuext_ents.get("bank name", []):
            if _is_clean_ner(bn):
                info["bank"] = _clean_value(bn)
                break
        if "bank" not in info:
            bank = re.search(r"(?:bank\s*name)\s*[:\s]*(.+)", text, re.IGNORECASE)
            if bank:
                info["bank"] = bank.group(1).strip().split("\n")[0].strip()
        acct_candidates: list[Candidate] = []
        for ba in nuext_ents.get("bank account", []):
            if _has_digits(ba) and _is_clean_ner(ba):
                acct_candidates.append(Candidate(_clean_value(ba), base_score("nuextract"), "nuextract"))
        acct_match = re.search(r"(?:account\s*(?:no|number|#)?)\s*[:\s]*([\d][\d\s\-]{4,25})", text, re.IGNORECASE)
        if acct_match:
            acct_candidates.append(Candidate(acct_match.group(1).strip(), base_score("regex_context"), "regex_context"))
        acct = pick_best(acct_candidates)
        if acct:
            info["account_number"] = acct
        for sc in nuext_ents.get("sort code", []):
            if _has_digits(sc) and _is_clean_ner(sc):
                info["sort_code"] = _clean_value(sc)
                break
        if "sort_code" not in info:
            sort_code = re.search(r"(?:sort\s*code)\s*[:\s]*([\d\-\s]+)", text, re.IGNORECASE)
            if sort_code:
                info["sort_code"] = sort_code.group(1).strip()
        for phone in nuext_ents.get("phone number", []):
            if _has_digits(phone) and _is_clean_ner(phone):
                info["phone"] = _clean_value(phone)
                break
        for email in nuext_ents.get("email address", []):
            if "@" in email and _is_clean_ner(email):
                info["email"] = _clean_value(email)
                break
        return info


class PurchaseOrderExtractor:

    def extract(self, context: dict[str, Any]) -> dict[str, Any]:
        text: str = context.get("text", "")
        spacy_ents: dict[str, list[str]] = context.get("spacy_ents", {})
        nuext_ents: dict[str, list[str]] = context.get("nuext_ents", {})
        word_boxes: list[dict] = context.get("word_boxes", [])
        layout_fields: dict[str, Any] = context.get("layout_fields", {})

        result: dict[str, Any] = {"document_type": "purchase_order"}

        result["po_number"] = self._extract_po_number(text, spacy_ents, nuext_ents, layout_fields)
        result["supplier_name"] = self._extract_supplier_name(text, spacy_ents, nuext_ents, layout_fields)
        result["supplier_address"] = self._extract_supplier_address(text, nuext_ents, layout_fields)
        result["buyer_name"] = self._extract_buyer_name(text, spacy_ents, nuext_ents, layout_fields)
        result["buyer_address"] = self._extract_buyer_address(text, nuext_ents, layout_fields)
        result["order_date"] = self._extract_order_date(text, spacy_ents, nuext_ents, layout_fields)
        result["quote_number"] = self._extract_quote_number(text, nuext_ents, layout_fields)
        result["currency"] = self._extract_currency(text, nuext_ents)
        result["subtotal"] = self._extract_subtotal(text, spacy_ents, nuext_ents, layout_fields)
        result["discount"] = self._extract_discount(text, nuext_ents, layout_fields)
        result["tax"] = self._extract_tax(text, spacy_ents, nuext_ents, layout_fields)
        result["total_amount"] = self._extract_total(text, spacy_ents, nuext_ents, layout_fields)
        result["line_items"] = self._extract_line_items(text, word_boxes, layout_fields)
        result["payment_terms"] = self._extract_payment_terms(text, nuext_ents)
        result["delivery_address"] = self._extract_delivery_address(text, nuext_ents, layout_fields)
        result["ship_to"] = find_ship_to(text) or ""

        return result

    def _extract_po_number(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_po = layout_fields.get("po_number", "")
        if layout_po and _has_digits(layout_po):
            candidates.append(Candidate(_clean_value(layout_po), base_score("layout"), "layout"))
        for v in nuext_ents.get("po_number", []):
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        po = find_po_number(text)
        if po:
            candidates.append(Candidate(po, base_score("regex_context"), "regex_context"))
        candidates = [adjust_invoice_number(c) for c in candidates]
        return pick_best(candidates)

    def _extract_supplier_name(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        recipient = find_recipient_po(text)
        if recipient:
            return recipient
        for v in nuext_ents.get("supplier_name", []):
            if _is_clean_ner(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for org in spacy_ents.get("ORG", []):
            if _is_clean_ner(org):
                candidates.append(Candidate(_clean_value(org), base_score("spacy"), "spacy"))
        candidates = [adjust_vendor_name(c) for c in candidates]
        return pick_best(candidates)

    def _extract_supplier_address(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_bt = layout_fields.get("bill_to", "")
        if layout_bt and _looks_like_address(layout_bt):
            candidates.append(Candidate(_clean_value(layout_bt), base_score("layout") + 0.05, "layout"))
        for addr in nuext_ents.get("supplier_address", []):
            if _looks_like_address(addr):
                candidates.append(Candidate(_clean_value(addr), base_score("nuextract"), "nuextract"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    def _extract_buyer_name(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        for v in nuext_ents.get("buyer_name", []):
            if _is_clean_ner(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        top = self._buyer_from_top_lines(text)
        if top:
            candidates.append(Candidate(top, base_score("heuristic") + 0.15, "heuristic"))
        text_lower = text.lower()
        recipient_pos = text_lower.find("recipient")
        for org in spacy_ents.get("ORG", []):
            if _is_clean_ner(org):
                clean = _clean_value(org)
                pos = text_lower.find(clean.lower())
                if 0 <= pos < 100 and (recipient_pos < 0 or pos < recipient_pos):
                    candidates.append(Candidate(clean, base_score("spacy") + 0.10, "spacy"))
        candidates = [adjust_vendor_name(c) for c in candidates]
        return pick_best(candidates)

    @staticmethod
    def _buyer_from_top_lines(text, max_lines=5):
        skip = {"purchase", "order", "purchase order", "po number",
                "description", "price", "qty", "subtotal", "amount",
                "recipient", "bill to", "invoice"}
        for line in text.split("\n")[:max_lines]:
            clean = line.strip()
            if not clean or len(clean) <= 1:
                continue
            lower = clean.lower().rstrip(":")
            if lower in skip:
                continue
            if re.match(r"^[\d/\-\s.,]+$", clean):
                continue
            if re.search(r"[£$€¥₹\u00a3]\s*\d", clean):
                continue
            if re.match(r"^[\d\+\-\s\(\)]{7,}$", clean):
                continue
            if "@" in clean or clean.startswith("www.") or clean.startswith("http"):
                continue
            if re.search(r"[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}", clean, re.IGNORECASE):
                continue
            if clean.count(",") >= 2:
                continue
            return clean
        return ""

    def _extract_buyer_address(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        addr = self._extract_buyer_addr_from_layout(layout_fields)
        if addr:
            candidates.append(Candidate(addr, base_score("layout") + 0.05, "layout"))
        for a in nuext_ents.get("buyer_address", []):
            if _looks_like_address(a):
                candidates.append(Candidate(_clean_value(a), base_score("nuextract"), "nuextract"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    @staticmethod
    def _extract_buyer_addr_from_layout(layout_fields):
        raw = layout_fields.get("vendor_address", "")
        if not raw:
            raw = layout_fields.get("vendor_name", "")
        if not raw:
            return ""
        parts: list[str] = []
        for segment in raw.split(", "):
            seg = segment.strip()
            if re.match(r"^(recipient|supplier|vendor)\s*:?\s*$", seg, re.IGNORECASE):
                break
            if re.match(r"^[\d\+\-\(\)\s]{7,}$", seg):
                break
            if "@" in seg:
                break
            parts.append(seg)
        result = ", ".join(parts)
        return _clean_value(result) if _looks_like_address(result) else ""

    def _extract_order_date(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_date = layout_fields.get("po_date", "") or layout_fields.get("invoice_date", "")
        if layout_date and _looks_like_date(layout_date):
            candidates.append(Candidate(_clean_value(layout_date), base_score("layout"), "layout"))
        for d in nuext_ents.get("order_date", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("nuextract"), "nuextract"))
        for d in spacy_ents.get("DATE", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("spacy"), "spacy"))
        po_date = find_po_date(text)
        if po_date:
            candidates.append(Candidate(po_date, base_score("regex_context") + 0.05, "regex_context"))
        all_dates = find_all_dates(text)
        for d in all_dates:
            candidates.append(Candidate(d, base_score("regex_fallback"), "regex_fallback"))
        candidates = [adjust_date(c) for c in candidates]
        return pick_best(candidates)

    def _extract_quote_number(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_quote = layout_fields.get("quote_number", "")
        if layout_quote:
            candidates.append(Candidate(_clean_value(layout_quote), base_score("layout"), "layout"))
        for v in nuext_ents.get("quote_number", []):
            if _is_clean_ner(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        quote = find_quote_number(text)
        if quote:
            candidates.append(Candidate(quote, base_score("regex_context"), "regex_context"))
        return pick_best(candidates)

    def _extract_currency(self, text, nuext_ents):
        candidates: list[Candidate] = []
        for c in nuext_ents.get("currency", []):
            code = c.strip().upper()
            if code in ("GBP", "USD", "EUR", "JPY", "INR", "CAD", "AUD"):
                candidates.append(Candidate(code, base_score("nuextract"), "nuextract"))
        detected = detect_currency(text)
        if detected:
            candidates.append(Candidate(detected, base_score("regex_context"), "regex_context"))
        return pick_best(candidates)

    def _extract_subtotal(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_sub = layout_fields.get("subtotal", "")
        if layout_sub and _looks_like_money(layout_sub):
            candidates.append(Candidate(_clean_value(layout_sub), base_score("layout"), "layout"))
        for v in nuext_ents.get("subtotal", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        sub = find_subtotal(text)
        if sub:
            candidates.append(Candidate(sub, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_discount(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_disc = layout_fields.get("discount", "")
        if layout_disc and _looks_like_money(layout_disc):
            candidates.append(Candidate(_clean_value(layout_disc), base_score("layout"), "layout"))
        for v in nuext_ents.get("discount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        disc = find_discount(text)
        if disc:
            candidates.append(Candidate(disc, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_tax(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_tax = layout_fields.get("tax", "")
        if layout_tax and _looks_like_money(layout_tax):
            candidates.append(Candidate(_clean_value(layout_tax), base_score("layout"), "layout"))
        for v in nuext_ents.get("tax_amount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        tax = find_tax(text)
        if tax:
            candidates.append(Candidate(tax, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_total(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_total = layout_fields.get("total", "")
        if layout_total and _looks_like_money(layout_total):
            candidates.append(Candidate(_clean_value(layout_total), base_score("layout") + 0.10, "layout"))
        for v in nuext_ents.get("total_amount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for v in spacy_ents.get("MONEY", []):
            v = v.strip()
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(v, base_score("spacy") - 0.10, "spacy"))
        regex_total = find_total(text)
        if regex_total:
            candidates.append(Candidate(regex_total, base_score("regex_context") + 0.15, "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_payment_terms(self, text, nuext_ents):
        for v in nuext_ents.get("payment_terms", []):
            if _is_clean_ner(v):
                return _clean_value(v)
        days = find_payment_terms(text)
        if days:
            return f"Net {days}"
        return ""

    def _extract_delivery_address(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        addr = self._extract_buyer_addr_from_layout(layout_fields)
        if addr:
            candidates.append(Candidate(addr, base_score("layout") + 0.05, "layout"))
        for addr in nuext_ents.get("delivery_address", []):
            if _looks_like_address(addr):
                candidates.append(Candidate(_clean_value(addr), base_score("nuextract"), "nuextract"))
        delivery = find_delivery_address(text)
        if delivery:
            candidates.append(Candidate(delivery, base_score("regex_context"), "regex_context"))
        ship = find_ship_to(text)
        if ship:
            candidates.append(Candidate(ship, base_score("regex_context"), "regex_context"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    def _extract_line_items(self, text, word_boxes, layout_fields=None):
        layout_fields = layout_fields or {}
        scored_sets: list[tuple[float, list[dict[str, str]]]] = []

        layout_headers = layout_fields.get("_table_headers", [])
        layout_rows = layout_fields.get("_table_rows", [])
        if layout_headers and layout_rows:
            layout_items = self._layout_rows_to_items(layout_headers, layout_rows)
            layout_items = self._filter_invalid_items(layout_items)
            if layout_items:
                score = self._score_line_item_set(layout_items, base_score("layout"))
                scored_sets.append((score, layout_items))

        text_items = find_line_items_from_text(text)
        text_items = self._filter_invalid_items(text_items)
        if text_items:
            score = self._score_line_item_set(text_items, base_score("text_table"))
            scored_sets.append((score, text_items))

        if word_boxes:
            ocr_items = extract_table_from_boxes(word_boxes)
            ocr_items = self._filter_invalid_items(ocr_items)
            if ocr_items:
                score = self._score_line_item_set(ocr_items, base_score("ocr_table"))
                scored_sets.append((score, ocr_items))

        if not scored_sets:
            logger.warning("No line items could be extracted.")
            return []

        scored_sets.sort(key=lambda x: x[0], reverse=True)
        best_score, best_items = scored_sets[0]
        return self._clean_line_items(best_items)

    @staticmethod
    def _filter_invalid_items(items):
        filtered = []
        for item in items:
            desc = item.get("item", "").strip()
            if re.match(r"^[\+\d\-\(\)\s]{1,5}$", desc):
                continue
            if re.match(r"^[\d\+\-\(\)\s]{7,}$", desc):
                continue
            if "@" in desc:
                continue
            if not desc:
                continue
            filtered.append(item)
        return filtered

    @staticmethod
    def _score_line_item_set(items, base):
        if not items:
            return 0.0
        score = base
        completeness_total = 0.0
        for item in items:
            fields_filled = sum(1 for k in ("item", "quantity", "unit_price", "amount") if item.get(k, "").strip())
            completeness_total += fields_filled / 4.0
            qty = item.get("quantity", "").strip()
            try:
                int(float(qty.replace(",", "")))
                score += 0.02
            except (ValueError, TypeError):
                score -= 0.03
            amt = item.get("amount", "").strip()
            if re.search(r"[£$€¥₹\u00a3]", amt):
                score += 0.02
        avg_completeness = completeness_total / len(items)
        score += avg_completeness * 0.15
        if len(items) >= 2:
            score += 0.05
        if len(items) >= 4:
            score += 0.03
        return min(score, 1.0)

    @staticmethod
    def _clean_line_items(items):
        merged: list[dict[str, str]] = []
        pending_desc = ""
        for item in items:
            has_any_numeric = any(item.get(k, "").strip() for k in ("quantity", "unit_price", "amount"))
            if not has_any_numeric:
                desc = item.get("item", "").strip()
                if desc:
                    pending_desc = (pending_desc + " " + desc).strip() if pending_desc else desc
            else:
                if pending_desc:
                    item = dict(item)
                    item["item"] = (pending_desc + " " + item.get("item", "")).strip()
                    pending_desc = ""
                merged.append(item)
        if pending_desc:
            merged.append({"item": pending_desc, "quantity": "", "unit_price": "", "amount": ""})

        cleaned: list[dict[str, str]] = []
        for item in merged:
            desc = item.get("item", "").strip()
            if re.match(r"^[\+\d\-\(\)\s]{1,5}$", desc):
                continue
            if re.match(r"^[\d\+\-\(\)\s]{7,}$", desc):
                continue
            if "@" in desc:
                continue
            qty = item.get("quantity", "").strip()
            qty = re.sub(r"[£$€¥₹\u00a3]", "", qty).strip()
            try:
                qty_int = int(float(qty.replace(",", "")))
            except (ValueError, TypeError):
                qty_int = None
            cleaned.append({
                "item": item.get("item", "").strip(),
                "quantity": str(qty_int) if qty_int is not None else qty,
                "unit_price": item.get("unit_price", "").strip(),
                "amount": item.get("amount", "").strip(),
            })
        return cleaned

    @staticmethod
    def _layout_rows_to_items(headers, rows):
        header_map: dict[int, str] = {}
        desc_keys = {"description", "item", "particular", "particulars", "product", "service", "details"}
        qty_keys = {"qty", "quantity", "units", "count"}
        price_keys = {"price", "rate", "unit price", "cost"}
        amount_keys = {"amount", "subtotal", "sub-total", "total", "line total", "extended"}
        for i, h in enumerate(headers):
            h_lower = h.strip().lower()
            if h_lower in desc_keys:
                header_map[i] = "item"
            elif h_lower in qty_keys:
                header_map[i] = "quantity"
            elif h_lower in price_keys:
                header_map[i] = "unit_price"
            elif h_lower in amount_keys:
                header_map[i] = "amount"
        if "item" not in header_map.values():
            if len(headers) >= 3:
                header_map = {0: "item", 1: "quantity", 2: "unit_price"}
                if len(headers) >= 4:
                    header_map[3] = "amount"
        items: list[dict[str, str]] = []
        for row in rows:
            item: dict[str, str] = {"item": "", "quantity": "", "unit_price": "", "amount": ""}
            for col_idx, field_name in header_map.items():
                if col_idx < len(row):
                    item[field_name] = row[col_idx].strip()
            if item["item"]:
                items.append(item)
        return items


class ExtractionEngine_invoice:

    def __init__(self) -> None:
        self._invoice_extractor = InvoiceExtractor()

    def run(self, text, word_boxes=None, layout_fields=None):
        if not text.strip():
            logger.warning("Empty text passed to ExtractionEngine_invoice.")
            return {"document_type": "invoice", "error": "No text could be extracted"}
        logger.info("Running spaCy NER...")
        spacy_ents = spacy_extract_entities(text)
        logger.info("Running NuExtract NER...")
        try:
            nuext_ents = nuextract_extract_entities_invoice(text)
        except Exception as exc:
            logger.warning("NuExtract extraction failed: %s -- continuing without it.", exc)
            nuext_ents = {}
        context = {"text": text, "spacy_ents": spacy_ents, "nuext_ents": nuext_ents,
                    "word_boxes": word_boxes or [], "layout_fields": layout_fields or {}}
        logger.info("Running InvoiceExtractor...")
        return self._invoice_extractor.extract(context)


class ExtractionEngine_po:

    def __init__(self) -> None:
        self._po_extractor = PurchaseOrderExtractor()

    def run(self, text, word_boxes=None, layout_fields=None):
        if not text.strip():
            logger.warning("Empty text passed to ExtractionEngine_po.")
            return {"document_type": "purchase_order", "error": "No text could be extracted"}
        logger.info("Running spaCy NER...")
        spacy_ents = spacy_extract_entities(text)
        logger.info("Running NuExtract NER...")
        try:
            nuext_ents = nuextract_extract_entities_po(text)
        except Exception as exc:
            logger.warning("NuExtract extraction failed: %s – continuing without it.", exc)
            nuext_ents = {}
        context = {"text": text, "spacy_ents": spacy_ents, "nuext_ents": nuext_ents,
                    "word_boxes": word_boxes or [], "layout_fields": layout_fields or {}}
        logger.info("Running PurchaseOrderExtractor...")
        return self._po_extractor.extract(context)


class AgentController_invoice:

    def __init__(self) -> None:
        self._engine = ExtractionEngine_invoice()
        self._layout_analyzer = LayoutAnalyzer()

    def process(self, file_path: str) -> dict[str, Any]:
        file_type = FileDetector.detect(file_path)
        logger.info("Detected file type: %s", file_type)
        text = ""
        word_boxes: list[dict] = []
        layout_fields: dict = {}
        if file_type == "pdf":
            text, word_boxes, layout_fields = self._extract_from_pdf(file_path)
        elif file_type == "docx":
            text = self._extract_from_docx(file_path)
        if not text.strip():
            logger.error("No text could be extracted from %s", file_path)
            return {"document_type": "invoice", "error": "No text could be extracted from the document."}
        logger.info("Extracted %d characters of text, %d word boxes, %d layout fields.", len(text), len(word_boxes), len(layout_fields))
        result = self._engine.run(text, word_boxes, layout_fields)
        result["_source_file"] = file_path
        return result

    def _extract_from_pdf(self, file_path):
        flat_text = PDFParser.extract_text(file_path)
        text = flat_text
        word_boxes: list[dict] = []
        layout_fields: dict = {}
        if len(text.strip()) >= OCR_MIN_TEXT_LENGTH:
            logger.info("Digital PDF – running block-level layout extraction.")
            try:
                blocks = PDFBlockParser.extract_blocks(file_path)
                page_w = PDFBlockParser.page_width(file_path)
                layout = self._layout_analyzer.analyze(blocks, page_w)
                if layout.reconstructed_text.strip():
                    text = layout.reconstructed_text
                layout_fields = self._layout_to_fields(layout)
                layout_fields["_flat_text"] = flat_text
            except Exception as exc:
                logger.warning("Block extraction failed, using flat text: %s", exc)
            try:
                ocr_text, word_boxes = run_ocr_with_boxes(file_path)
                if ocr_text:
                    layout_fields["_ocr_text"] = ocr_text
            except Exception as exc:
                logger.warning("OCR box extraction failed (non-critical): %s", exc)
        else:
            logger.info("Scanned PDF detected – running full OCR.")
            try:
                text, word_boxes = run_ocr_with_boxes(file_path)
            except Exception as exc:
                logger.error("OCR failed: %s", exc)
                text = PDFParser.extract_text(file_path)
        return text, word_boxes, layout_fields

    @staticmethod
    def _layout_to_fields(layout):
        fields: dict = {}
        if layout.vendor_zone:
            address_parts: list[str] = []
            for line in layout.vendor_zone:
                stripped = line.strip()
                is_address = bool(re.search(r"\d.*[A-Z]{1,2}\d", stripped))
                if not is_address and ", " in stripped and re.search(r"\d", stripped):
                    is_address = True
                if is_address:
                    address_parts.append(stripped)
            if address_parts:
                fields["vendor_address"] = ", ".join(address_parts)
        if layout.bill_to_zone:
            fields["bill_to"] = ", ".join(layout.bill_to_zone)
        fields.update(layout.header_fields)
        if layout.totals:
            for key, val in layout.totals.items():
                fields[key] = val
        if layout.table_headers and layout.table_rows:
            fields["_table_headers"] = layout.table_headers
            fields["_table_rows"] = layout.table_rows
        if layout.payment_zone:
            fields["_payment_zone"] = layout.payment_zone
        return fields

    @staticmethod
    def _extract_from_docx(file_path):
        logger.info("DOCX – extracting text via python-docx.")
        return DOCXParser.extract_text(file_path)


class AgentController_po:

    def __init__(self) -> None:
        self._engine = ExtractionEngine_po()
        self._layout_analyzer = LayoutAnalyzer()

    def process(self, file_path: str) -> dict[str, Any]:
        file_type = FileDetector.detect(file_path)
        logger.info("Detected file type: %s", file_type)
        text = ""
        word_boxes: list[dict] = []
        layout_fields: dict = {}
        if file_type == "pdf":
            text, word_boxes, layout_fields = self._extract_from_pdf(file_path)
        elif file_type == "docx":
            text = self._extract_from_docx(file_path)
        if not text.strip():
            logger.error("No text could be extracted from %s", file_path)
            return {"document_type": "purchase_order", "error": "No text could be extracted from the document."}
        logger.info("Extracted %d characters of text, %d word boxes, %d layout fields.", len(text), len(word_boxes), len(layout_fields))
        result = self._engine.run(text, word_boxes, layout_fields)
        result["_source_file"] = file_path
        return result

    def _extract_from_pdf(self, file_path):
        text = PDFParser.extract_text(file_path)
        word_boxes: list[dict] = []
        layout_fields: dict = {}
        if len(text.strip()) >= OCR_MIN_TEXT_LENGTH:
            logger.info("Digital PDF – running block-level layout extraction.")
            try:
                blocks = PDFBlockParser.extract_blocks(file_path)
                page_w = PDFBlockParser.page_width(file_path)
                layout = self._layout_analyzer.analyze(blocks, page_w)
                if layout.reconstructed_text.strip():
                    text = layout.reconstructed_text
                layout_fields = self._layout_to_fields(layout)
            except Exception as exc:
                logger.warning("Block extraction failed, using flat text: %s", exc)
            try:
                _, word_boxes = run_ocr_with_boxes(file_path)
            except Exception as exc:
                logger.warning("OCR box extraction failed (non-critical): %s", exc)
        else:
            logger.info("Scanned PDF detected – running full OCR.")
            try:
                text, word_boxes = run_ocr_with_boxes(file_path)
            except Exception as exc:
                logger.error("OCR failed: %s", exc)
                text = PDFParser.extract_text(file_path)
        return text, word_boxes, layout_fields

    @staticmethod
    def _layout_to_fields(layout):
        fields: dict = {}
        if layout.vendor_zone:
            fields["vendor_name"] = layout.vendor_zone[0]
            if len(layout.vendor_zone) > 1:
                fields["vendor_address"] = ", ".join(layout.vendor_zone[1:])
        if layout.bill_to_zone:
            fields["bill_to"] = ", ".join(layout.bill_to_zone)
        fields.update(layout.header_fields)
        if layout.totals:
            for key, val in layout.totals.items():
                fields[key] = val
        if layout.table_headers and layout.table_rows:
            fields["_table_headers"] = layout.table_headers
            fields["_table_rows"] = layout.table_rows
        if layout.payment_zone:
            fields["_payment_zone"] = layout.payment_zone
        return fields

    @staticmethod
    def _extract_from_docx(file_path):
        logger.info("DOCX – extracting text via python-docx.")
        return DOCXParser.extract_text(file_path)


def _parse_numeric(val: str) -> str:

    if not val:
        return ""
    cleaned = re.sub(r"[£$€¥₹\u00a3\s]", "", val).strip()
    cleaned = cleaned.replace(",", "")
    try:
        float(cleaned)
        return cleaned
    except ValueError:
        return val


def _round_money(val: str) -> str:
    if not val:
        return ""
    try:
        return f"{float(val):.2f}"
    except (ValueError, TypeError):
        return val


_DATE_FORMATS = [
    "%d/%m/%Y",
    "%d-%m-%Y",
    "%Y-%m-%d",
    "%d/%m/%y",
    "%d-%m-%y",
    "%d %B %Y",
    "%d %b %Y",
    "%d %b, %Y",
    "%d %B, %Y",
    "%B %d, %Y",
    "%b %d, %Y",
    "%B %Y",
    "%b %Y",
]


# Cached supplier data from Postgres (with TTL)
_SUPPLIER_CACHE: Optional[dict] = None
_SUPPLIER_CACHE_TIME: float = 0.0
_SUPPLIER_CACHE_TTL: float = 300.0  # 5 minutes

_BIZ_SUFFIXES = re.compile(
    r"\b(ltd|limited|llc|inc|plc|corp|corporation|gmbh|"
    r"s\.?a\.?|pty|co\.?|company|group|partners|llp|lp|"
    r"and\s+sons|& sons)\b\.?",
    re.IGNORECASE,
)
_NOISE_WORDS = {"and", "the", "of", "for", "a", "an", "&", "etc"}


def _normalize(name):
    s = name.lower().strip()
    s = _BIZ_SUFFIXES.sub("", s)
    s = re.sub(r"[,.\-&'\"]+", " ", s)
    return " ".join(s.split()).strip()


def _core_tokens(name):
    norm = _normalize(name)
    return {t for t in norm.split() if t and t not in _NOISE_WORDS}


def _load_suppliers_from_db() -> dict:
    """Load supplier names from proc.supplier Postgres table with TTL cache."""
    global _SUPPLIER_CACHE, _SUPPLIER_CACHE_TIME
    import time

    now = time.time()
    if _SUPPLIER_CACHE and (now - _SUPPLIER_CACHE_TIME) < _SUPPLIER_CACHE_TTL:
        return _SUPPLIER_CACHE

    suppliers: list[str] = []
    lower_map: dict[str, str] = {}
    norm_map: dict[str, str] = {}
    core_map: dict[str, set[str]] = {}

    if _db_connection_func is None:
        logger.warning("No database connection function set for supplier lookup")
        return {"suppliers": [], "lower_map": {}, "norm_map": {}, "core_tokens_map": {}}

    try:
        conn = _db_connection_func()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT supplier_name, trading_name FROM proc.supplier "
                    "WHERE supplier_name IS NOT NULL"
                )
                rows = cur.fetchall()
                for row in rows:
                    supplier_name = (row[0] or "").strip()
                    trading_name = (row[1] or "").strip() if len(row) > 1 else ""
                    for name in [supplier_name, trading_name]:
                        if name and name not in suppliers:
                            suppliers.append(name)
                            lower_map[name.lower()] = name
                            norm_map[_normalize(name)] = name
                            core_map[name] = _core_tokens(name)
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as exc:
        logger.warning("Failed to load suppliers from database: %s", exc)
        return {"suppliers": [], "lower_map": {}, "norm_map": {}, "core_tokens_map": {}}

    _SUPPLIER_CACHE = {
        "suppliers": suppliers,
        "lower_map": lower_map,
        "norm_map": norm_map,
        "core_tokens_map": core_map,
    }
    _SUPPLIER_CACHE_TIME = now
    logger.info("Loaded %d supplier names from proc.supplier", len(suppliers))
    return _SUPPLIER_CACHE


def _match_supplier(extracted_name):
    if not extracted_name:
        return ""
    data = _load_suppliers_from_db()
    suppliers = data["suppliers"]
    lower_map = data["lower_map"]
    norm_map = data["norm_map"]
    core_map = data["core_tokens_map"]
    if not suppliers:
        return ""
    query = extracted_name.strip()
    query_lower = query.lower()
    query_norm = _normalize(query)
    query_core = _core_tokens(query)
    if query_lower in lower_map:
        return lower_map[query_lower]
    best_contain = ""
    best_contain_len = 0
    for s_norm, s_orig in norm_map.items():
        if not s_norm or len(s_norm) < 3:
            continue
        if re.search(r"(?:^|\s)" + re.escape(s_norm) + r"(?:\s|$)", query_norm) or \
           re.search(r"(?:^|\s)" + re.escape(query_norm) + r"(?:\s|$)", s_norm):
            if len(s_norm) > best_contain_len:
                best_contain = s_orig
                best_contain_len = len(s_norm)
    if best_contain:
        return best_contain
    if not query_core:
        return ""
    best_score = 0.0
    best_name = ""
    for s_orig, s_tokens in core_map.items():
        if not s_tokens:
            continue
        overlap = query_core & s_tokens
        n_overlap = len(overlap)
        if n_overlap == 0:
            continue
        all_query_matched = query_core <= s_tokens
        all_supplier_matched = s_tokens <= query_core
        if n_overlap == 1 and not (all_query_matched or all_supplier_matched):
            continue
        union = len(query_core | s_tokens)
        jaccard = n_overlap / union if union else 0.0
        score = jaccard
        if all_query_matched and len(query_core) >= 2:
            score = max(score, 0.7)
        if score > best_score:
            best_score = score
            best_name = s_orig
    if best_score >= 0.5:
        return best_name
    return ""


def _scan_text_for_supplier(text):
    if not text:
        return ""
    data = _load_suppliers_from_db()
    suppliers = data["suppliers"]
    lower_map = data["lower_map"]
    if not suppliers:
        return ""
    text_lower = text.lower()
    best_match = ""
    best_len = 0
    for s_lower, s_orig in lower_map.items():
        if len(s_lower) < 4:
            continue
        pattern = r"\b" + re.escape(s_lower) + r"\b"
        if re.search(pattern, text_lower):
            if len(s_lower) > best_len:
                best_match = s_orig
                best_len = len(s_lower)
    if best_match:
        return best_match
    for s_orig in suppliers:
        s_norm = _normalize(s_orig)
        if not s_norm or len(s_norm) < 4:
            continue
        norm_words = [w for w in s_norm.split() if w not in _NOISE_WORDS]
        if len(norm_words) < 2:
            continue
        pattern = r"\b" + re.escape(s_norm) + r"\b"
        if re.search(pattern, text_lower):
            if len(s_norm) > best_len:
                best_match = s_orig
                best_len = len(s_norm)
    return best_match


def _compute_tax_percent_invoice(subtotal: str, tax: str) -> str:
    try:
        sub_val = float(_parse_numeric(subtotal))
        tax_val = float(_parse_numeric(tax))
        if sub_val > 0:
            pct = (tax_val / sub_val) * 100
            return f"{pct:.1f}"
    except (ValueError, ZeroDivisionError):
        pass
    return ""


def _parse_date_invoice(date_str: str) -> datetime | None:
    if not date_str:
        return None
    cleaned = re.sub(r"(\d{1,2})(?:st|nd|rd|th)\b", r"\1", date_str)
    cleaned = re.sub(r"\bSept\b", "Sep", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(
        r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.",
        r"\1", cleaned, flags=re.IGNORECASE,
    )
    cleaned = " ".join(cleaned.split()).strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
    m = re.match(r"(\d{1,2})\s+(\w+)\s+(\d{4})", cleaned)
    if m:
        day, month_str, year_str = int(m.group(1)), m.group(2), int(m.group(3))
        for mfmt in ("%B", "%b"):
            try:
                month_num = datetime.strptime(month_str, mfmt).month
                max_day = calendar.monthrange(year_str, month_num)[1]
                clamped_day = min(day, max_day)
                return datetime(year_str, month_num, clamped_day)
            except ValueError:
                continue
    return None


def _format_date(dt: datetime, original: str) -> str:
    if "/" in original:
        return dt.strftime("%d/%m/%Y")
    if "-" in original and re.match(r"\d{4}-", original):
        return dt.strftime("%Y-%m-%d")
    if "-" in original:
        return dt.strftime("%d-%m-%Y")
    return dt.strftime("%d/%m/%Y")


def _normalize_date_invoice(date_str: str) -> str:
    parsed = _parse_date_invoice(date_str)
    if parsed is None:
        return ""
    return parsed.strftime("%d/%m/%Y")


def _default_due_date(invoice_date: str, days: int = 90) -> str:
    parsed = _parse_date_invoice(invoice_date)
    if parsed is None:
        return ""
    due = parsed + timedelta(days=days)
    return due.strftime("%d/%m/%Y")


_TITLE_RE = re.compile(
    r"^(managing\s+director|director|ceo|cfo|cto|coo|president|vp|"
    r"vice\s+president|manager|head\s+of|partner|founder|owner|"
    r"accountant|accounts?\s+(?:payable|receivable)|procurement|"
    r"finance\s+(?:director|manager)|secretary|administrator)\b",
    re.IGNORECASE,
)

_LABEL_WORDS = {"invoice", "bill", "billed", "to", "bill to", "billed to",
                "billing to", "invoice to", "receipt", "statement",
                "customer", "customer information", "name", "address"}

_METADATA_RE = re.compile(
    r"^(purchase\s*order|po\s*(?:no|number|#|id|ref)?|"
    r"invoice\s*(?:no|number|#)?)\s*[:\s]",
    re.IGNORECASE,
)


def _extract_buyer_from_bill_to(bill_to):
    if not bill_to:
        return ""
    parts = [p.strip() for p in bill_to.split(",") if p.strip()]
    _sublabel_re = re.compile(r"^(?:name|address|company|contact|phone|tel|email|fax)\s*:\s*", re.IGNORECASE)
    parts = [_sublabel_re.sub("", p).strip() for p in parts]
    parts = [p for p in parts if p]
    _company_re = re.compile(
        r"\b(ltd|limited|llc|inc|plc|corp|corporation|gmbh|"
        r"s\.?a\.?|pty|co\.|company|group|partners|llp|lp)\b",
        re.IGNORECASE,
    )
    _url_or_domain_re = re.compile(
        r".*@|^\+?\d[\d\s\-()]{6,}$|^https?://|^www\."
        r"|\.(?:co\.uk|com|org|net|io|co|uk|gov|edu|info)\b",
        re.IGNORECASE,
    )
    for part in parts:
        lower = part.lower().strip()
        if lower in _LABEL_WORDS or _METADATA_RE.match(lower):
            continue
        if _url_or_domain_re.search(part.strip()):
            continue
        if _company_re.search(part):
            return part.strip()
    for part in parts:
        lower = part.lower().strip()
        if lower in _LABEL_WORDS:
            continue
        if _METADATA_RE.match(lower):
            continue
        if _TITLE_RE.match(lower):
            continue
        if _url_or_domain_re.search(part.strip()):
            continue
        if re.search(r"\b(road|street|lane|way|avenue|drive|park|crescent|close"
                     r"|court|place|square)\b", lower) and re.search(r"\d", part):
            continue
        if re.match(r"^[A-Z]{1,2}\d.*\d[A-Z]{2}$", part.strip(), re.IGNORECASE):
            continue
        alpha_count = sum(1 for c in part if c.isalpha())
        if alpha_count >= 3:
            return part.strip()
    return ""



def _ai_identify_supplier(text: str, extracted_name: str = "") -> str:
    """Use AI (Ollama) as final gate to identify the supplier name."""
    if not text:
        return ""
    text_snippet = text[:3000]
    prompt = (
        "You are a procurement document analysis expert. "
        "Extract the supplier/vendor company name from this document text. "
        "The supplier is the company that is SELLING goods or services (not the buyer). "
        "Return ONLY the company name, nothing else. "
        "If you cannot determine the supplier name, return 'UNKNOWN'.\n\n"
    )
    if extracted_name:
        prompt += f"A previous extraction attempt found this name: '{extracted_name}'. Verify or correct it.\n\n"
    prompt += f"Document text:\n{text_snippet}\n\nSupplier name:"
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={"model": NUEXTRACT_MODEL, "prompt": prompt, "stream": False,
                  "options": {"temperature": 0, "num_predict": 100, "num_gpu": 0}},
            timeout=60,
        )
        response.raise_for_status()
        result = response.json().get("response", "").strip()
        result = result.split("\n")[0].strip()
        result = re.sub(r"^(the\s+)?supplier\s*(name\s*)?(is\s*)?:?\s*", "", result, flags=re.IGNORECASE).strip()
        result = result.strip("'\"")
        if result and result.upper() != "UNKNOWN" and len(result) > 1:
            logger.info("AI identified supplier: %s", result)
            return result
    except Exception as exc:
        logger.warning("AI supplier identification failed: %s", exc)
    return ""


def resolve_supplier_name(extracted_name: str, raw_text: str = "") -> str:
    """Resolve supplier name: 1) Postgres lookup, 2) text scan, 3) AI fallback."""
    if extracted_name:
        matched = _match_supplier(extracted_name)
        if matched:
            return matched
    if raw_text:
        scanned = _scan_text_for_supplier(raw_text)
        if scanned:
            return scanned
    if raw_text:
        ai_name = _ai_identify_supplier(raw_text, extracted_name)
        if ai_name:
            matched = _match_supplier(ai_name)
            if matched:
                return matched
            return ai_name
    return extracted_name or ""


def map_invoice(result: dict[str, Any]) -> dict[str, str]:
    subtotal = result.get("subtotal", "")
    tax = result.get("tax", "")
    total = result.get("total_amount", "")
    discount = result.get("discount", "")
    if not subtotal and total and tax:
        try:
            total_val = float(_parse_numeric(total))
            tax_val = float(_parse_numeric(tax))
            derived = total_val - tax_val
            if derived > 0:
                subtotal = f"{derived:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    effective_subtotal = subtotal
    if discount and subtotal:
        try:
            sub_val = float(_parse_numeric(subtotal))
            disc_val = float(_parse_numeric(discount))
            effective_subtotal = f"{sub_val - disc_val:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    tax_pct = _compute_tax_percent_invoice(effective_subtotal, tax)
    invoice_date_raw = result.get("invoice_date", "")
    due_date_raw = result.get("due_date", "")
    terms_days = result.get("payment_terms_days", "")
    invoice_date = _normalize_date_invoice(invoice_date_raw)
    due_date = _normalize_date_invoice(due_date_raw)
    if not due_date and invoice_date_raw:
        days = int(terms_days) if terms_days else 90
        due_date = _default_due_date(invoice_date_raw, days)
    payment_terms = ""
    if due_date and invoice_date:
        parsed_inv = _parse_date_invoice(invoice_date)
        parsed_due = _parse_date_invoice(due_date)
        if parsed_inv and parsed_due:
            diff = (parsed_due - parsed_inv).days
            if diff >= 0:
                payment_terms = str(diff)
    bill_to = result.get("bill_to", "")
    recipient = result.get("recipient", "")
    buyer_name = _extract_buyer_from_bill_to(bill_to)
    if not buyer_name and recipient:
        buyer_name = _extract_buyer_from_bill_to(recipient)
    supplier_address = result.get("supplier_address", "") or result.get("vendor_address", "")
    buyer_address = result.get("buyer_address", "") or bill_to or recipient
    addr_to_parse = supplier_address or buyer_address
    all_addresses = " ".join(filter(None, [supplier_address, buyer_address]))
    addr_parts = _extract_address_parts(addr_to_parse)
    raw_text = result.get("_raw_text", "")
    region = _detect_region(all_addresses)
    if not region and raw_text:
        region = _detect_region(raw_text)
    if not addr_parts["country"]:
        raw_lower = raw_text.lower() if raw_text else ""
        combined_lower = all_addresses.lower()
        if "united kingdom" in raw_lower or "united kingdom" in combined_lower:
            addr_parts["country"] = "United Kingdom"
        elif "united states" in raw_lower or "united states" in combined_lower:
            addr_parts["country"] = "United States"
        elif re.search(r"[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}", all_addresses, re.IGNORECASE):
            addr_parts["country"] = "United Kingdom"
    return {
        "invoice_id": result.get("invoice_number", ""),
        "po_id": result.get("po_number", ""),
        "supplier_name": resolve_supplier_name(
            result.get("vendor_name", ""),
            result.get("_raw_text", ""),
        ),
        "buyer_id": buyer_name,
        "requisition_id": "",
        "requested_by": "",
        "requested_date": "",
        "invoice_date": invoice_date,
        "due_date": due_date,
        "invoice_paid_date": "",
        "payment_terms": payment_terms,
        "currency": result.get("currency", ""),
        "invoice_amount": _round_money(_parse_numeric(effective_subtotal)),
        "tax_percent": tax_pct,
        "tax_amount": _round_money(_parse_numeric(tax)),
        "invoice_total_incl_tax": _round_money(_parse_numeric(total)),
        "exchange_rate_to_usd": "",
        "converted_amount_usd": "",
        "country": addr_parts["country"],
        "region": region,
        "invoice_status": "",
        "ai_flag_required": "",
        "trigger_type": "",
        "trigger_context_description": "",
        "created_date": "",
        "created_by": "",
        "last_modified_by": "",
        "last_modified_date": "",
    }


def _merge_continuation_items(items: list[dict[str, str]]) -> list[dict[str, str]]:
    if len(items) <= 1:
        return items

    merged: list[dict[str, str]] = []
    for item in items:
        desc = item.get("item", "").strip()
        amount = item.get("amount", "").strip()
        qty = item.get("quantity", "").strip()
        unit_price = item.get("unit_price", "").strip()

        is_continuation = (
            not amount and not qty and not unit_price
            and merged
            and desc
        )

        if is_continuation:
            prev = merged[-1]
            prev_desc = prev.get("item", "")
            prev["item"] = f"{prev_desc} {desc}".strip()
        else:
            merged.append(dict(item))

    return merged


def map_line_items_invoice(result: dict[str, Any]) -> list[dict[str, str]]:
    invoice_id = result.get("invoice_number", "")
    po_id = result.get("po_number", "")
    items = _merge_continuation_items(result.get("line_items", []))
    subtotal = result.get("subtotal", "")
    tax = result.get("tax", "")
    total = result.get("total_amount", "")
    discount = result.get("discount", "")
    if not subtotal and total and tax:
        try:
            total_val = float(_parse_numeric(total))
            tax_val = float(_parse_numeric(tax))
            derived = total_val - tax_val
            if derived > 0:
                subtotal = f"{derived:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    effective_subtotal = subtotal
    if discount and subtotal:
        try:
            sub_val = float(_parse_numeric(subtotal))
            disc_val = float(_parse_numeric(discount))
            effective_subtotal = f"{sub_val - disc_val:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    tax_pct = _compute_tax_percent_invoice(effective_subtotal, tax)
    mapped: list[dict[str, str]] = []
    for i, item in enumerate(items, start=1):
        unit_price = item.get("unit_price", "")
        quantity = item.get("quantity", "")
        line_total = item.get("amount", "")
        if not line_total and len(items) == 1 and subtotal:
            line_total = subtotal
        if not line_total and unit_price and not quantity:
            line_total = unit_price
        line_tax = ""
        total_with_tax = ""
        try:
            lt_val = float(_parse_numeric(line_total))
            if tax_pct:
                pct = float(tax_pct) / 100
                line_tax_val = lt_val * pct
                line_tax = f"{line_tax_val:.2f}"
                total_with_tax = f"{lt_val + line_tax_val:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
        mapped.append({
            "invoice_id": invoice_id,
            "line_no": str(i),
            "item_id": item.get("item", ""),
            "quantity": quantity,
            "unit_price": _round_money(_parse_numeric(unit_price)),
            "tax_percent": tax_pct,
            "po_id": po_id,
            "line_total": _round_money(_parse_numeric(line_total)),
            "tax_amount": _round_money(line_tax),
            "total_with_tax": _round_money(total_with_tax),
            "created_date": "",
            "created_by": "",
            "last_modified_by": "",
            "last_modified_date": "",
        })
    return mapped


def _validate_tax(subtotal, discount, tax, total):
    try:
        sub_val = float(_parse_numeric(subtotal))
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        tot_val = float(_parse_numeric(total)) if total else 0.0
        tax_val = float(_parse_numeric(tax)) if tax else 0.0
        if sub_val > 0 and tot_val > 0:
            derived_tax = tot_val - (sub_val - disc_val)
            if derived_tax >= 0:
                if abs(derived_tax - tax_val) > 1.0:
                    return f"{derived_tax:.2f}"
    except (ValueError, ZeroDivisionError):
        pass
    return _parse_numeric(tax) if tax else ""


def _compute_tax_percent_po(subtotal: str, tax: str, discount: str = "") -> str:
    try:
        sub_val = float(_parse_numeric(subtotal))
        tax_val = float(_parse_numeric(tax))
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        taxable = sub_val - disc_val
        if taxable > 0:
            pct = (tax_val / taxable) * 100
            return f"{pct:.1f}"
    except (ValueError, ZeroDivisionError):
        pass
    return ""


def _parse_date_po(date_str: str) -> datetime | None:
    if not date_str:
        return None
    cleaned = re.sub(r"(\d{1,2})(?:st|nd|rd|th)\b", r"\1", date_str)
    cleaned = re.sub(r"\bSept\b", "Sep", cleaned)
    cleaned = " ".join(cleaned.split()).strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
    return None


def _normalize_date_po(date_str: str) -> str:
    parsed = _parse_date_po(date_str)
    if parsed is None:
        return ""
    return parsed.strftime("%d/%m/%Y")


def _extract_address_parts(address):
    parts = {"line1": "", "line2": "", "city": "", "postal_code": "", "country": ""}
    if not address:
        return parts
    _non_addr = re.compile(r"^(recipient|supplier|vendor|bill\s*to|ship\s*to)\s*:?\s*$", re.IGNORECASE)
    segments = [s.strip() for s in address.split(",")
                if s.strip() and not _non_addr.match(s.strip())
                and not re.match(r"^[\d\+\-\(\)\s]{7,}$", s.strip())
                and "@" not in s]
    if len(segments) >= 2:
        parts["line1"] = segments[0]
        parts["line2"] = ", ".join(segments[1:])
    elif len(segments) == 1:
        addr = segments[0]
        pc = re.search(r"([A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2})", addr, re.IGNORECASE)
        if pc:
            before_pc = addr[:pc.start()].rstrip()
            postcode_str = pc.group(1)
            city_match = re.search(
                r"\b(London|Birmingham|Manchester|Leeds|Liverpool|Bristol|Brighton|"
                r"Newport|Horsham|Sheffield|Edinburgh|Glasgow|Cardiff|Belfast|"
                r"York|Oxford|Cambridge|Nottingham|Leicester|Southampton|"
                r"New York|Los Angeles|Chicago|Houston|Phoenix)\b",
                before_pc, re.IGNORECASE,
            )
            if city_match:
                split_pos = city_match.start()
                parts["line1"] = before_pc[:split_pos].rstrip(" ,")
                parts["line2"] = before_pc[split_pos:].rstrip() + " " + postcode_str
            else:
                words = before_pc.split()
                mid = len(words) // 2
                if mid > 0:
                    parts["line1"] = " ".join(words[:mid])
                    parts["line2"] = " ".join(words[mid:]) + " " + postcode_str
                else:
                    parts["line1"] = before_pc + " " + postcode_str
        else:
            parts["line1"] = addr
    postcode_match = re.search(r"\b([A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2})\b", address, re.IGNORECASE)
    if postcode_match:
        parts["postal_code"] = postcode_match.group(1).upper()
    if not parts["postal_code"]:
        zip_match = re.search(r"\b(\d{5}(?:-\d{4})?)\b", address)
        if zip_match:
            parts["postal_code"] = zip_match.group(1)
    address_lower = address.lower()
    if "united kingdom" in address_lower or "uk" in address_lower:
        parts["country"] = "United Kingdom"
    elif "united states" in address_lower or "usa" in address_lower:
        parts["country"] = "United States"
    elif postcode_match:
        parts["country"] = "United Kingdom"
    city_match = re.search(
        r"\b(London|Birmingham|Manchester|Leeds|Liverpool|Bristol|Brighton|"
        r"Newport|Horsham|Sheffield|Edinburgh|Glasgow|Cardiff|Belfast|"
        r"York|Oxford|Cambridge|Nottingham|Leicester|Southampton|"
        r"New York|Los Angeles|Chicago|Houston|Phoenix)\b",
        address, re.IGNORECASE,
    )
    if city_match:
        parts["city"] = city_match.group(1).title()
    return parts


def _detect_region(address):
    normalized = re.sub(r",\s*", " ", address)
    normalized = " ".join(normalized.split())
    region_match = re.search(
        r"\b(West Sussex|East Sussex|West Yorkshire|South Yorkshire|North Yorkshire|"
        r"West Midlands|East Midlands|Greater London|Surrey|Kent|Essex|Hampshire|"
        r"Devon|Cornwall|Somerset|Norfolk|Suffolk|Berkshire|Oxfordshire|"
        r"Cambridgeshire|Hertfordshire|Lancashire|Cheshire|Wales|Scotland)\b",
        normalized, re.IGNORECASE,
    )
    if region_match:
        return region_match.group(1).title()
    return ""


def map_purchase_order(result: dict[str, Any]) -> dict[str, str]:
    subtotal = result.get("subtotal", "")
    discount = result.get("discount", "")
    tax_raw = result.get("tax", "")
    total_raw = result.get("total_amount", "")
    tax = _validate_tax(subtotal, discount, tax_raw, total_raw)
    tax_pct = _compute_tax_percent_po(subtotal, tax, discount)
    try:
        sub_val = float(_parse_numeric(subtotal)) if subtotal else 0.0
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        tax_val = float(tax) if tax else 0.0
        computed_total = (sub_val - disc_val) + tax_val
        final_total = f"{computed_total:.2f}"
    except (ValueError, ZeroDivisionError):
        final_total = _parse_numeric(total_raw)
    order_date_raw = result.get("order_date", "")
    order_date = _normalize_date_po(order_date_raw)
    currency = result.get("currency", "")
    delivery_addr = result.get("delivery_address", "") or result.get("ship_to", "")
    buyer_addr = result.get("buyer_address", "")
    addr_to_parse = delivery_addr or buyer_addr
    addr_parts = _extract_address_parts(addr_to_parse)
    region = _detect_region(addr_to_parse)
    payment_terms = result.get("payment_terms", "")
    return {
        "po_id": result.get("po_number", ""),
        "supplier_name": resolve_supplier_name(
            result.get("vendor_name", "") or result.get("supplier_name", ""),
            result.get("_raw_text", ""),
        ),
        "buyer_id": result.get("buyer_name", ""),
        "requisition_id": "",
        "requested_by": "",
        "requested_date": "",
        "currency_code": currency,
        "order_date": order_date,
        "expected_delivery_date": "",
        "ship_to_country": addr_parts["country"],
        "delivery_region": region,
        "incoterm_code": "",
        "incoterm_responsibility": "",
        "total_amount": final_total,
        "delivery_address_line1": addr_parts["line1"],
        "delivery_address_line2": addr_parts["line2"],
        "delivery_city": addr_parts["city"],
        "postal_code": addr_parts["postal_code"],
        "base_currency": currency,
        "po_status": "",
        "payment_terms": payment_terms,
        "exchange_rate_to_usd": "",
        "total_amount_usd": "",
        "ai_flag_required": "",
        "trigger_type": "",
        "trigger_context_description": "",
        "created_date": "",
        "created_by": "",
        "contract_id": "",
    }


def map_po_line_items(result: dict[str, Any]) -> list[dict[str, str]]:
    po_id = result.get("po_number", "")
    quote_number = result.get("quote_number", "")
    currency = result.get("currency", "")
    items = result.get("line_items", [])
    subtotal = result.get("subtotal", "")
    discount = result.get("discount", "")
    tax_raw = result.get("tax", "")
    total_raw = result.get("total_amount", "")
    tax = _validate_tax(subtotal, discount, tax_raw, total_raw)
    tax_pct = _compute_tax_percent_po(subtotal, tax, discount)
    try:
        sub_val = float(_parse_numeric(subtotal))
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        discount_ratio = disc_val / sub_val if sub_val > 0 else 0.0
    except (ValueError, ZeroDivisionError):
        discount_ratio = 0.0
    mapped: list[dict[str, str]] = []
    for i, item in enumerate(items, start=1):
        unit_price = item.get("unit_price", "")
        quantity = item.get("quantity", "")
        line_total = item.get("amount", "")
        line_after_disc_str = ""
        line_tax = ""
        total_with_tax = ""
        try:
            lt_val = float(_parse_numeric(line_total))
            line_after_disc = lt_val * (1 - discount_ratio)
            line_after_disc_str = f"{line_after_disc:.2f}"
            if tax_pct:
                pct = float(tax_pct) / 100
                line_tax_val = line_after_disc * pct
                line_tax = f"{line_tax_val:.2f}"
                total_with_tax = f"{line_after_disc + line_tax_val:.2f}"
            else:
                total_with_tax = f"{line_after_disc:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
        mapped.append({
            "po_id": po_id,
            "po_line_id": "",
            "line_number": str(i),
            "item_id": "",
            "item_description": item.get("item", ""),
            "quote_number": quote_number,
            "quantity": quantity,
            "unit_price": _parse_numeric(unit_price),
            "unit_of_measure": "",
            "currency_code": currency,
            "line_total_amount": line_after_disc_str or _parse_numeric(line_total),
            "tax_percentage": tax_pct,
            "tax_amount": line_tax,
            "total_amount": total_with_tax,
            "created_date": "",
            "created_by": "",
        })
    return mapped


QUOTE_TEMPLATE = {
    "quote number": "",
    "quote date": "",
    "validity or pay by date": "",
    "purchase order number": "",
    "buyer or billed to company name": "",
    "buyer or billed to company address": "",
    "supplier or quoting company name": "",
    "supplier or quoting company address": "",
    "subtotal amount before tax": "",
    "discount amount": "",
    "tax amount": "",
    "total amount after tax": "",
    "currency code": "",
    "payment terms": "",
    "phone number": "",
    "email address": "",
}

_QUOTE_LABEL_MAP = {
    "quote number": "quote_number",
    "quote date": "quote_date",
    "validity or pay by date": "validity_date",
    "purchase order number": "po_number",
    "buyer or billed to company name": "buyer_name",
    "buyer or billed to company address": "buyer_address",
    "supplier or quoting company name": "supplier_name",
    "supplier or quoting company address": "supplier_address",
    "subtotal amount before tax": "subtotal",
    "discount amount": "discount",
    "tax amount": "tax_amount",
    "total amount after tax": "total_amount",
    "currency code": "currency",
    "payment terms": "payment_terms",
    "phone number": "phone_number",
    "email address": "email_address",
}


def nuextract_extract_entities_quote(text: str) -> dict[str, list[str]]:
    chunks = _split_text(text, max_chars=3000)
    entities: dict[str, list[str]] = {}
    for chunk in chunks:
        try:
            extracted = _call_nuextract_quote(chunk)
        except Exception as exc:
            logger.warning("NuExtract extraction failed on chunk: %s", exc)
            continue
        for key, value in extracted.items():
            if key not in _QUOTE_LABEL_MAP:
                continue
            label = _QUOTE_LABEL_MAP[key]
            values = value if isinstance(value, list) else [value]
            for v in values:
                v = str(v).strip()
                if not v:
                    continue
                entities.setdefault(label, [])
                if v not in entities[label]:
                    entities[label].append(v)
    return entities


def _call_nuextract_quote(text: str) -> dict:
    template_str = json.dumps(QUOTE_TEMPLATE, indent=2)
    prompt = f"""<|input|>
{text}
<|template|>
{template_str}
<|output|>"""
    response = requests.post(
        f"{OLLAMA_BASE_URL}/api/generate",
        json={
            "model": NUEXTRACT_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0,
                "num_predict": 2048,
                "num_gpu": 0,
            },
        },
        timeout=120,
    )
    response.raise_for_status()
    raw_output = response.json().get("response", "").strip()
    logger.debug("NuExtract raw output: %s", raw_output[:500])
    return _parse_response(raw_output)


_QUOTE_NUM_PATTERNS_Q = [
    re.compile(
        r"(?:quote|quotation)\s*(?:no|number|#|num|ref)?\.?\s*[:\s]+\n?\s*"
        r"(QUT[\-]?\d[\d\-]{2,20})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:quote|quotation)\s*(?:no|number|#|num|ref)?\.?\s*[:\s]+\n?\s*"
        r"(\d{3,20})",
        re.IGNORECASE,
    ),
    re.compile(r"\b(QUT[\-]?\d[\d\-]{2,20})\b"),
    re.compile(r"\b(DHA[\-]\d[\d\-]{2,20})\b"),
    re.compile(r"(?:^|\n)\s*QUOTE\s*\n\s*(\d{3,20})\s*(?:\n|$)"),
    re.compile(
        r"(?:quote|quotation)\s*(?:no|number|#|num|ref)?\.?\s*[:\s]+\n?\s*"
        r"([A-Z][\w\-/]*\d[\w\-/]*)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:inv|invoice)\s*(?:no|number|#)\.?\s*[:\s./]+\s*"
        r"(\d{4,20})",
        re.IGNORECASE,
    ),
]


def find_quote_number_q(text: str) -> Optional[str]:
    for pat in _QUOTE_NUM_PATTERNS_Q:
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]", val):
                continue
            lower = val.lower()
            if lower in ("bill", "to", "no", "item", "product", "service",
                         "wade", "assurity", "number"):
                continue
            if len(val) < 2 or re.match(r"^[\W]+$", val):
                continue
            end_pos = m.end(1)
            if end_pos < len(text) and re.match(r"[\-]\d{2,}", text[end_pos:]):
                continue
            return val
    return None


_QUOTE_DATE_PATTERNS_Q = [
    re.compile(
        r"(?:^|\n)\s*(?:quote\s*)?date\s*(?:issued)?\s*[:\s]+\n?([\s\S]{0,200})",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:issue\s*date|date\s*(?:of\s*)?(?:issue|quote|quotation))"
        r"[\s.:]+\n?([\s\S]{0,200})",
        re.IGNORECASE,
    ),
    re.compile(r"(?:date\s*issued)[\s.:]+\n?([\s\S]{0,200})", re.IGNORECASE),
    re.compile(r"(?:date)[\s.:]+(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})", re.IGNORECASE),
]


def find_quote_date_q(text: str) -> Optional[str]:
    for pat in _QUOTE_DATE_PATTERNS_Q:
        m = pat.search(text)
        if m:
            raw = m.group(1).strip()
            for dp in _DATE_PATTERNS:
                dm = dp.search(raw)
                if dm:
                    return dm.group(1)
    return None


_PAY_BY_PATTERNS_Q = [
    re.compile(
        r"(?:pay\s*by|valid\s*(?:until|till|through)|validity\s*date|expiry\s*date|expires?\s*on)"
        r"\s*[:\s]*\n?([\s\S]{0,120})",
        re.IGNORECASE,
    ),
]


def find_pay_by_date_q(text: str) -> Optional[str]:
    for pat in _PAY_BY_PATTERNS_Q:
        m = pat.search(text)
        if m:
            raw = m.group(1).strip()
            for dp in _DATE_PATTERNS:
                dm = dp.search(raw)
                if dm:
                    return dm.group(1)
    return None


_QUOTE_TO_PATTERNS_Q = [
    re.compile(
        r"(?:quote\s*to|quoted?\s*for|billing\s*to|recipient)\s*[:\s]*\n\s*(.+?)(?:\n|$)",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:quote\s*to|quoted?\s*for|billing\s*to|recipient)\s*[:\s]+(.+?)(?:\n|$)",
        re.IGNORECASE,
    ),
]


def find_quote_to_q(text: str) -> Optional[str]:
    for pat in _QUOTE_TO_PATTERNS_Q:
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if val and len(val) > 1:
                val = re.sub(r"\s+\S+@\S+", "", val).strip()
                val = re.sub(r"\s+[\d\+\-\(\)]{7,}", "", val).strip()
                return val
    return None


_PAYABLE_TO_PATTERNS_Q = [
    re.compile(r"(?:payable\s*to|sender)\s*[:\s]*\n\s*(.+?)(?:\n|$)", re.IGNORECASE),
    re.compile(r"(?:payable\s*to|sender)\s*[:\s]+(.+?)(?:\n|$)", re.IGNORECASE),
    re.compile(r"(?:pay\s*to)\s*[:\s]*\n\s*(.+?)(?:\n|$)", re.IGNORECASE),
]


def find_payable_to_q(text: str) -> Optional[str]:
    for pat in _PAYABLE_TO_PATTERNS_Q:
        m = pat.search(text)
        if m:
            val = m.group(1).strip()
            if val and len(val) > 1:
                val = re.sub(r"\s+\S+@\S+", "", val).strip()
                val = re.sub(r"\s+[\d\+\-\(\)]{7,}", "", val).strip()
                val = val.split(",")[0].strip()
                lower = val.lower()
                if lower.startswith(("bank ", "notes", "terms", "account", "payment")):
                    continue
                if re.match(r"^\d+\s+\w", val) and not re.search(r"\b(ltd|llc|inc|plc|corp)\b", val, re.IGNORECASE):
                    continue
                if re.match(r"^(suite|unit|floor|level|room)\s+\d", val, re.IGNORECASE):
                    continue
                return val
    return None


def find_payable_to_address_q(text: str) -> Optional[str]:
    pat = re.compile(
        r"(?:payable\s*to|pay\s*to|sender)\s*[:\s]*\n((?:.+\n?){1,8})",
        re.IGNORECASE,
    )
    m = pat.search(text)
    if m:
        block = m.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        addr_lines: list[str] = []
        for line in lines[1:]:
            if re.match(r"^(notes|payment|bank|account|total|description|item|quote)", line, re.IGNORECASE):
                break
            if re.match(r"^(https?://|www\.|.*@)", line, re.IGNORECASE):
                break
            addr_lines.append(line)
            if len(addr_lines) >= 4:
                break
        if addr_lines:
            return ", ".join(addr_lines)
    return None


_BILL_TO_PATTERN_Q = re.compile(
    r"(?:bill(?:ing)?\s*(?:to|ed\s*to)|customer\s*information)\s*[:\s]*\n?((?:.+\n?){1,10})",
    re.IGNORECASE,
)


def find_bill_to_q(text: str) -> Optional[str]:
    m = _BILL_TO_PATTERN_Q.search(text)
    if m:
        block = m.group(1).strip()
        lines = [l.strip() for l in block.split("\n") if l.strip()]
        filtered: list[str] = []
        for line in lines:
            if re.match(r"^(payable|notes|payment|bank|account|total|description|item\s+description|purchase\s*order|p\.?o\.?\s*(?:no|number|#|ref)|terms)", line, re.IGNORECASE):
                break
            if re.match(r"^(date\s*issued|invoice\s*no|po\s*no|quote\s*number)\s*[:.]?\s*$", line, re.IGNORECASE):
                continue
            if re.match(r"^(invoice(\s+to)?|bill(\s+to)?|billed(\s+to)?|quote(\s+to)?|quotation|receipt|statement)\s*[:.]?\s*$", line, re.IGNORECASE):
                continue
            if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}$", line):
                continue
            if re.match(r"^(https?://|www\.|.*@)", line, re.IGNORECASE):
                break
            if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", line):
                break
            if re.match(r"^(name|address|company|contact|phone|tel|email|fax)\s*:?\s*$", line, re.IGNORECASE):
                continue
            if re.match(r"^[:.\-]+$", line):
                continue
            cleaned = re.sub(
                r"^(?:name|address|company|contact)\s*:\s*",
                "", line, flags=re.IGNORECASE,
            ).strip()
            if cleaned:
                filtered.append(cleaned)
        if filtered:
            return ", ".join(filtered[:4])
    return None


BILL_TO_KEYWORDS_Q = {
    "bill to", "billed to", "billing to", "invoice to", "invoiced to",
    "sold to", "customer", "customer information",
    "quote to", "quoted for",
    "recipient",
}


class LayoutAnalyzer_quote:

    def analyze(self, blocks: list[TextBlock], page_width: float = 595.0) -> InvoiceLayout:
        if not blocks:
            return InvoiceLayout()
        layout = InvoiceLayout()
        rows = _group_into_rows(blocks)
        page_mid = _detect_page_mid(blocks, page_width)

        table_start = None
        for i, row in enumerate(rows):
            if _is_table_header_row(row):
                table_start = i
                layout.table_headers = [b.text.strip() for b in row.left_blocks]
                break

        table_end = None
        if table_start is not None:
            table_end = self._find_table_end(rows, table_start)

        header_rows = rows[:table_start] if table_start else rows
        self._process_header_zone(header_rows, page_mid, layout, rows)

        if table_start is not None and table_end is not None:
            self._process_table_zone(rows, table_start, table_end, layout)

        footer_start = table_end if table_end else (table_start or len(rows))
        if footer_start < len(rows):
            self._process_footer_zone(rows[footer_start:], layout)

        layout.reconstructed_text = self._reconstruct_text(layout, rows, page_mid)
        return layout

    def _process_header_zone(self, rows: list[LayoutRow], page_mid: float,
                             layout: InvoiceLayout,
                             all_rows: list[LayoutRow] | None = None) -> None:
        left_lines: list[str] = []
        right_lines: list[str] = []
        bill_to_found = False
        bill_to_row_idx = -1
        bill_to_side: str = ""

        _skip_vendor = re.compile(
            r"^(invoice|bill|receipt|statement|tax\s*invoice|credit\s*note"
            r"|quote|quotation"
            r"|invoice\s*no|inv\s*no|p\.?o\.?\s*(?:no|ref)|date|due\s*date"
            r"|purchase\s*order|description\s*of\s*service"
            r"|payment\s*(?:instruction|info|method|term))",
            re.IGNORECASE,
        )
        _date_line = re.compile(
            r"^\d{1,2}[/\-]\d{1,2}[/\-]"
            r"|^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
            r"|^\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)",
            re.IGNORECASE,
        )

        for row_i, row in enumerate(rows):
            left_text = row.column_text(page_mid, "left")
            right_text = row.column_text(page_mid, "right")

            for block in row.left_blocks:
                self._try_extract_header_field(block.text, row, layout,
                                               all_rows=all_rows, row_idx=row_i)

            for block in row.left_blocks:
                if _has_label(block.text.strip().lower(), BILL_TO_KEYWORDS_Q):
                    bill_to_found = True
                    bill_to_row_idx = row_i
                    bill_to_side = "left" if block.mid_x < page_mid else "right"
                    break

            full = row.full_text.lower()
            if _has_label(full, BILL_TO_KEYWORDS_Q):
                for block in row.left_blocks:
                    bt = block.text.strip()
                    bt_lower = bt.lower()
                    if _has_label(bt_lower, BILL_TO_KEYWORDS_Q):
                        colon_match = re.search(r":\s*(.+)", bt)
                        if colon_match:
                            inline_val = colon_match.group(1).strip()
                            if inline_val:
                                right_lines.append(inline_val)
                        continue
                    if any(val and val == bt for val in layout.header_fields.values()):
                        continue
                    if block.mid_x >= page_mid:
                        right_lines.append(bt)
                continue

            if left_text:
                stripped = left_text.strip()
                if _date_line.match(stripped):
                    if "invoice_date" not in layout.header_fields:
                        dates = re.findall(r"\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}", stripped)
                        if dates:
                            layout.header_fields["invoice_date"] = dates[0]
                        else:
                            layout.header_fields["invoice_date"] = stripped
                elif not _skip_vendor.match(stripped):
                    if (bill_to_found
                            and row_i > bill_to_row_idx
                            and bill_to_side == "left"):
                        right_lines.append(left_text)
                    else:
                        left_lines.append(left_text)

            if right_text:
                stripped_right = right_text.strip()
                skip_right = False
                for val in layout.header_fields.values():
                    if val and val in right_text:
                        skip_right = True
                        break
                if skip_right:
                    pass
                elif _date_line.match(stripped_right):
                    if "invoice_date" not in layout.header_fields:
                        layout.header_fields["invoice_date"] = stripped_right
                else:
                    if bill_to_found:
                        right_lines.append(right_text)

        layout.vendor_zone = [l for l in left_lines if l.strip()]
        _skip_bill_to = re.compile(
            r"^(invoice\s*to|bill\s*to|billed\s*to|billing\s*to|sold\s*to|"
            r"quote\s*to|quoted?\s*for|customer\s*(?:information)?|recipient)",
            re.IGNORECASE,
        )
        layout.bill_to_zone = [l for l in right_lines
                                if l.strip() and not _skip_bill_to.match(l.strip())]

    def _try_extract_header_field(self, text: str, row: LayoutRow,
                                   layout: InvoiceLayout,
                                   all_rows: list[LayoutRow] | None = None,
                                   row_idx: int = -1) -> bool:
        field_labels = {
            r"(?:quote|quotation)\s*(?:no|number|#|num|ref)": "quote_number",
            r"inv(?:oice)?\s*(?:no|number|#|num)": "invoice_number",
            r"p\.?o\.?\s*(?:no|number|#|num|ref(?:erence)?)?": "po_number",
            r"(?:quote\s*)?date\s*(?:issued)?|issue\s*date": "invoice_date",
            r"due\s*date|payment\s*due|pay\s*by|valid(?:ity)?\s*date": "due_date",
        }
        lower = text.strip().lower().rstrip(":.")
        for pattern, field_name in field_labels.items():
            if re.match(pattern, lower):
                sep_match = re.search(r"[.:]\s*(.+)", text)
                if sep_match:
                    val = sep_match.group(1).strip()
                    if val and re.search(r"\d", val):
                        if field_name in ("quote_number", "invoice_number", "po_number"):
                            parts = re.split(r"\s*/\s*", val)
                            val = parts[0].strip()
                        layout.header_fields[field_name] = val
                        return True
                blocks = row.left_blocks
                for i, b in enumerate(blocks):
                    if b.text == text:
                        for j in range(i + 1, len(blocks)):
                            next_text = blocks[j].text.strip()
                            if next_text in (":", ".", ":."):
                                continue
                            if next_text and re.search(r"\d", next_text):
                                layout.header_fields[field_name] = next_text
                                return True
                            break
                label_x = 0.0
                for b in row.left_blocks:
                    if b.text == text:
                        label_x = b.mid_x
                        break
                if all_rows and 0 <= row_idx < len(all_rows) - 1:
                    max_x_dist = 150.0
                    for offset in range(1, min(3, len(all_rows) - row_idx)):
                        candidate_row = all_rows[row_idx + offset]
                        best_block = None
                        best_dist = float("inf")
                        for nb in candidate_row.left_blocks:
                            nt = nb.text.strip()
                            if nt and nt not in (":", ".") and re.search(r"\d", nt):
                                dist = abs(nb.mid_x - label_x)
                                if dist < best_dist and dist <= max_x_dist:
                                    best_dist = dist
                                    best_block = nt
                        if best_block:
                            layout.header_fields[field_name] = best_block
                            return True
                return True
        return False

    def _find_table_end(self, rows: list[LayoutRow], table_start: int) -> int:
        typical_gap = 50.0
        if table_start + 2 < len(rows):
            first_gap = rows[table_start + 2].y - rows[table_start + 1].y
            if first_gap > 5:
                typical_gap = first_gap
        gap_threshold = max(typical_gap * 2.0, 60.0)
        for i in range(table_start + 1, len(rows)):
            for block in rows[i].left_blocks:
                text_lower = block.text.strip().lower().rstrip(":")
                text_lower = re.sub(r"\s*\(?[\d.]+%\)?", "", text_lower).strip()
                if text_lower in TOTALS_KEYWORDS and i > table_start + 1:
                    return i
                if _has_label(text_lower, PAYMENT_KEYWORDS):
                    return i
            if i > table_start + 2:
                prev_y = rows[i - 1].y
                curr_y = rows[i].y
                if curr_y - prev_y > gap_threshold:
                    return i
        return len(rows)

    def _process_table_zone(self, rows: list[LayoutRow], table_start: int,
                            table_end: int, layout: InvoiceLayout) -> None:
        if table_start + 1 >= table_end:
            return
        header_row = rows[table_start]
        header_xs = [b.mid_x for b in header_row.left_blocks]
        n_cols = len(header_xs)
        for i in range(table_start + 1, table_end):
            row = rows[i]
            row_blocks = row.left_blocks
            if not row_blocks:
                continue
            first_text = row_blocks[0].text.strip().lower()
            if first_text in TOTALS_KEYWORDS or first_text.startswith("sub"):
                continue
            cells = [""] * n_cols
            for block in row_blocks:
                best_col = min(range(n_cols), key=lambda c: abs(block.mid_x - header_xs[c]))
                if cells[best_col]:
                    cells[best_col] += " " + block.text
                else:
                    cells[best_col] = block.text
            if not cells[0].strip():
                if layout.table_rows:
                    for c in range(1, n_cols):
                        if cells[c].strip() and not layout.table_rows[-1][c].strip():
                            layout.table_rows[-1][c] = cells[c]
                continue
            has_numeric = any(
                _LA_MONEY_RE.search(cells[c]) or re.match(r"^\d+$", cells[c].strip())
                for c in range(1, n_cols)
                if cells[c].strip()
            )
            if not has_numeric and layout.table_rows:
                desc_text = cells[0].strip()
                has_upcoming_amount = False
                for lookahead in range(i + 1, min(i + 3, table_end)):
                    la_row = rows[lookahead]
                    la_blocks = la_row.left_blocks
                    if not la_blocks:
                        continue
                    la_cells = [""] * n_cols
                    for block in la_blocks:
                        bc = min(range(n_cols), key=lambda c: abs(block.mid_x - header_xs[c]))
                        if la_cells[bc]:
                            la_cells[bc] += " " + block.text
                        else:
                            la_cells[bc] = block.text
                    if not la_cells[0].strip() and any(
                        _LA_MONEY_RE.search(la_cells[c]) or re.match(r"^\d+$", la_cells[c].strip())
                        for c in range(1, n_cols) if la_cells[c].strip()
                    ):
                        has_upcoming_amount = True
                        break
                    if la_cells[0].strip():
                        break
                if has_upcoming_amount:
                    layout.table_rows.append(cells)
                else:
                    layout.table_rows[-1][0] += " " + desc_text
            else:
                layout.table_rows.append(cells)

    def _process_footer_zone(self, rows: list[LayoutRow], layout: InvoiceLayout) -> None:
        total_labels = {
            "sub-total": "subtotal", "sub total": "subtotal", "subtotal": "subtotal",
            "net": "subtotal", "discount": "discount",
            "tax": "tax", "vat": "tax", "gst": "tax",
            "total": "total", "grand total": "total", "grand-total": "total",
            "amount due": "total", "balance due": "total",
        }
        for row in rows:
            for block in row.left_blocks:
                text = block.text.strip()
                text_lower = text.lower().rstrip(":")
                text_lower = re.sub(r"\s*\(?[\d.]+%\)?", "", text_lower).strip()
                matched_field = None
                for kw, field_name in total_labels.items():
                    if text_lower == kw:
                        matched_field = field_name
                        break
                if matched_field is None:
                    for kw, field_name in total_labels.items():
                        if text_lower.startswith(kw):
                            money = _LA_MONEY_RE.search(text)
                            if money:
                                if (field_name == "subtotal"
                                        and "subtotal" in layout.totals
                                        and "tax" in layout.totals):
                                    layout.totals["total"] = money.group(0)
                                else:
                                    layout.totals[field_name] = money.group(0)
                                break
                            matched_field = field_name
                            break
                if matched_field is None:
                    continue
                if (matched_field == "subtotal"
                        and "subtotal" in layout.totals
                        and "tax" in layout.totals):
                    matched_field = "total"
                for other in row.left_blocks:
                    if other.x0 > block.x1 and _LA_MONEY_RE.search(other.text):
                        layout.totals[matched_field] = other.text.strip()
                        break
            for block in row.left_blocks:
                text_lower = block.text.strip().lower()
                if _has_label(text_lower, PAYMENT_KEYWORDS):
                    layout.payment_zone.append(row.full_text)
                    break

    def _reconstruct_text(self, layout: InvoiceLayout, rows: list[LayoutRow],
                          page_mid: float) -> str:
        parts: list[str] = []
        if layout.vendor_zone:
            parts.extend(layout.vendor_zone)
            parts.append("")
        if layout.bill_to_zone:
            parts.append("Bill To:")
            parts.extend(layout.bill_to_zone)
            parts.append("")
        field_labels = {
            "quote_number": "Quote No",
            "invoice_number": "Invoice No",
            "po_number": "PO No",
            "invoice_date": "Date Issued",
            "due_date": "Due Date",
        }
        for key, label in field_labels.items():
            val = layout.header_fields.get(key, "")
            if val:
                parts.append(f"{label}: {val}")
        if layout.header_fields:
            parts.append("")
        if layout.table_headers:
            parts.append("\t".join(layout.table_headers))
            for row_cells in layout.table_rows:
                parts.append("\t".join(row_cells))
            parts.append("")
        if layout.totals:
            for key in ("subtotal", "discount", "tax", "total"):
                val = layout.totals.get(key, "")
                if val:
                    label = {"subtotal": "Sub-Total", "discount": "Discount",
                             "tax": "Tax", "total": "TOTAL"}[key]
                    parts.append(f"{label}: {val}")
            parts.append("")
        if layout.payment_zone:
            parts.extend(layout.payment_zone)
        text = "\n".join(parts)
        if len(text.strip()) < 50:
            text = "\n".join(row.full_text for row in rows)
        return text


class QuoteExtractor(BaseExtractor):

    def extract(self, context: dict[str, Any]) -> dict[str, Any]:
        text: str = context.get("text", "")
        raw_text: str = context.get("raw_text", text)
        spacy_ents: dict[str, list[str]] = context.get("spacy_ents", {})
        nuext_ents: dict[str, list[str]] = context.get("nuext_ents", {})
        word_boxes: list[dict] = context.get("word_boxes", [])
        layout_fields: dict[str, Any] = context.get("layout_fields", {})

        result: dict[str, Any] = {"document_type": "quote"}

        result["quote_number"] = self._extract_quote_number(text, nuext_ents, layout_fields, raw_text)
        result["supplier_name"] = self._extract_supplier_name(text, spacy_ents, nuext_ents, layout_fields, raw_text)
        result["buyer_name"] = self._extract_buyer_name(text, spacy_ents, nuext_ents, layout_fields, raw_text)
        buyer_name = result["buyer_name"]
        supplier_name = result["supplier_name"]
        result["supplier_address"] = self._extract_supplier_address(text, nuext_ents, layout_fields, raw_text, buyer_name)
        result["buyer_address"] = self._extract_buyer_address(text, nuext_ents, layout_fields, raw_text, buyer_name)

        if (not result["supplier_address"]
                or result["supplier_address"] == result["buyer_address"]):
            search_text = raw_text if raw_text and raw_text != text else text
            other_addr = self._find_other_address(search_text, result["buyer_address"], supplier_name)
            if other_addr:
                result["supplier_address"] = other_addr
            elif result["supplier_address"] == result["buyer_address"]:
                result["supplier_address"] = ""

        result["quote_date"] = self._extract_quote_date(text, spacy_ents, nuext_ents, layout_fields, raw_text)
        result["validity_date"] = self._extract_validity_date(text, nuext_ents, raw_text)
        result["po_number"] = self._extract_po_number(text, nuext_ents, layout_fields)
        result["currency"] = self._extract_currency(text, nuext_ents)
        result["subtotal"] = self._extract_subtotal(text, spacy_ents, nuext_ents, layout_fields)
        result["discount"] = self._extract_discount(text, nuext_ents, layout_fields)
        result["tax"] = self._extract_tax(text, spacy_ents, nuext_ents, layout_fields)
        result["total_amount"] = self._extract_total(text, spacy_ents, nuext_ents, layout_fields)
        result["line_items"] = self._extract_line_items(text, word_boxes, layout_fields, raw_text)
        result["payment_terms"] = self._extract_payment_terms(text, nuext_ents)

        flat_text = layout_fields.get("_flat_text", text)
        ocr_text = layout_fields.get("_ocr_text", "")
        parts = [text]
        if flat_text != text:
            parts.append(flat_text)
        if ocr_text:
            parts.append(ocr_text)
        result["_raw_text"] = "\n".join(parts)

        return result

    def _extract_quote_number(self, text, nuext_ents, layout_fields=None, raw_text=""):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_quote = layout_fields.get("quote_number", "")
        if layout_quote and _has_digits(layout_quote):
            candidates.append(Candidate(_clean_value(layout_quote), base_score("layout"), "layout"))
        for v in nuext_ents.get("quote_number", []):
            if _is_clean_ner(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        quote = find_quote_number_q(text)
        if quote:
            candidates.append(Candidate(quote, base_score("regex_context") + 0.10, "regex_context"))
        if raw_text and raw_text != text:
            quote_raw = find_quote_number_q(raw_text)
            if quote_raw and quote_raw != quote:
                candidates.append(Candidate(quote_raw, base_score("regex_context") + 0.05, "regex_context"))
        search_text = raw_text or text
        inv = find_invoice_number(search_text)
        if inv and not quote:
            candidates.append(Candidate(inv, base_score("regex_fallback"), "regex_fallback"))
        candidates = [adjust_invoice_number(c) for c in candidates]
        return pick_best(candidates)

    def _extract_supplier_name(self, text, spacy_ents, nuext_ents, layout_fields=None, raw_text=""):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        search_text = raw_text or text
        scanned = _scan_text_for_supplier(search_text)
        if scanned:
            candidates.append(Candidate(scanned, 0.98, "supplier_list"))
        for st in [text, raw_text] if raw_text and raw_text != text else [text]:
            payable = find_payable_to_q(st)
            if payable and self._is_valid_name(payable):
                matched = _match_supplier(payable)
                if matched:
                    candidates.append(Candidate(matched, 0.95, "supplier_list"))
                else:
                    candidates.append(Candidate(payable, base_score("regex_context") + 0.15, "regex_context"))
                break
        vendor = layout_fields.get("vendor_name", "")
        if vendor and self._is_valid_name(vendor):
            matched = _match_supplier(vendor)
            if matched:
                candidates.append(Candidate(matched, 0.95, "supplier_list"))
            else:
                candidates.append(Candidate(_clean_value(vendor), base_score("layout") + 0.05, "layout"))
        for v in nuext_ents.get("supplier_name", []):
            if _is_clean_ner(v) and self._is_valid_name(v):
                matched = _match_supplier(v)
                if matched:
                    candidates.append(Candidate(matched, 0.95, "supplier_list"))
                else:
                    candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        top = self._supplier_from_top_lines(search_text)
        if top and self._is_valid_name(top):
            matched = _match_supplier(top)
            if matched:
                candidates.append(Candidate(matched, 0.95, "supplier_list"))
            else:
                candidates.append(Candidate(top, base_score("heuristic") + 0.10, "heuristic"))
        company = self._find_company_with_suffix(search_text)
        if company and "assurity" not in company.lower():
            matched = _match_supplier(company)
            if matched:
                candidates.append(Candidate(matched, 0.95, "supplier_list"))
            else:
                candidates.append(Candidate(company, base_score("heuristic") + 0.12, "heuristic"))
        for org in spacy_ents.get("ORG", []):
            if _is_clean_ner(org):
                clean = _clean_value(org)
                if "assurity" not in clean.lower() and self._is_valid_name(clean):
                    matched = _match_supplier(clean)
                    if matched:
                        candidates.append(Candidate(matched, 0.93, "supplier_list"))
                    else:
                        candidates.append(Candidate(clean, base_score("spacy"), "spacy"))
        candidates = [adjust_vendor_name(c) for c in candidates]
        return pick_best(candidates)

    @staticmethod
    def _is_valid_name(s):
        clean = s.strip()
        if not clean or len(clean) < 2:
            return False
        if re.match(r"^(https?://|www\.)", clean, re.IGNORECASE):
            return False
        if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", clean):
            return False
        if "@" in clean:
            return False
        if re.match(r"^[\d\+\-\(\)\s]{7,}$", clean):
            return False
        if re.match(r"^\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}$", clean):
            return False
        if re.match(r"^[\d\s,.\-]+$", clean):
            return False
        lower = clean.lower()
        if lower.startswith(("date", "issue date", "sender:", "payable")):
            return False
        return True

    @staticmethod
    def _find_company_with_suffix(text):
        pattern = re.compile(
            r"([A-Z][\w\s,.'&\-]{2,40}?"
            r"\b(?:Ltd|Limited|LLC|Inc|Incorporated|Corp|Corporation|PLC|"
            r"GmbH|Pvt|Co|Company|Trading\s+Ltd))\b\.?",
        )
        _skip_words = {"assurity", "bank", "rolands", "fauget", "borcelle",
                       "wells fargo", "metro bank", "midland", "khan-woods",
                       "walker-vasquez", "becker"}
        for m in pattern.finditer(text):
            name = m.group(1).strip()
            lower = name.lower()
            if any(w in lower for w in _skip_words):
                continue
            start = max(0, m.start() - 100)
            context = text[start:m.start()].lower()
            if any(w in context for w in ("bank", "account", "sort code", "iban", "swift")):
                continue
            return name
        return ""

    @staticmethod
    def _supplier_from_top_lines(text, max_lines=8):
        skip = {"quote", "quotation", "invoice", "bill to", "billed to",
                "customer", "recipient", "description", "item",
                "price", "qty", "subtotal", "amount", "total",
                "payment", "terms", "service", "services"}
        for line in text.split("\n")[:max_lines]:
            clean = line.strip()
            if not clean or len(clean) <= 1:
                continue
            lower = clean.lower().rstrip(":")
            if lower in skip:
                continue
            if "assurity" in lower:
                continue
            if re.match(r"^[\d/\-\s.,]+$", clean):
                continue
            if re.search(r"[£$€¥₹\u00a3]\s*\d", clean):
                continue
            if re.match(r"^[\d\+\-\s\(\)]{7,}$", clean):
                continue
            if "@" in clean or clean.startswith("www.") or clean.startswith("http"):
                continue
            if re.search(r"[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}", clean, re.IGNORECASE):
                continue
            if clean.count(",") >= 2:
                continue
            return clean
        return ""

    def _extract_supplier_address(self, text, nuext_ents, layout_fields=None, raw_text="", buyer_name=""):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        _buyer_keywords = {"assurity"}
        if buyer_name:
            _buyer_keywords.add(buyer_name.lower().split(",")[0].split()[0])

        def _is_buyer_address(addr):
            return any(kw in addr.lower() for kw in _buyer_keywords)

        vendor_addr = layout_fields.get("vendor_address", "")
        if vendor_addr and _looks_like_address(vendor_addr) and not _is_buyer_address(vendor_addr):
            candidates.append(Candidate(_clean_value(vendor_addr), base_score("layout") + 0.05, "layout"))
        for addr in nuext_ents.get("supplier_address", []):
            if _looks_like_address(addr) and not _is_buyer_address(addr):
                candidates.append(Candidate(_clean_value(addr), base_score("nuextract"), "nuextract"))
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            payable_addr = find_payable_to_address_q(search_text)
            if payable_addr and _looks_like_address(payable_addr) and not _is_buyer_address(payable_addr):
                candidates.append(Candidate(_clean_value(payable_addr), base_score("regex_context"), "regex_context"))
                break
        supplier_name = layout_fields.get("vendor_name", "")
        if supplier_name:
            first_word = supplier_name.strip().split()[0] if supplier_name.strip() else ""
            if first_word and len(first_word) >= 3 and first_word.lower() not in _buyer_keywords:
                search_text = raw_text or text
                supplier_addr = self._find_address_near_keyword(search_text, first_word)
                if supplier_addr and not _is_buyer_address(supplier_addr):
                    candidates.append(Candidate(supplier_addr, base_score("heuristic") + 0.08, "heuristic"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    def _extract_buyer_name(self, text, spacy_ents, nuext_ents, layout_fields=None, raw_text=""):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            bill_to = find_bill_to_q(search_text)
            if bill_to:
                name = bill_to.split(",")[0].strip()
                if self._is_valid_name(name):
                    candidates.append(Candidate(_clean_value(name), base_score("regex_context") + 0.15, "regex_context"))
                break
        bt = layout_fields.get("bill_to", "")
        if bt:
            name = bt.split(",")[0].strip()
            if _is_clean_ner(name) and self._is_valid_name(name):
                candidates.append(Candidate(_clean_value(name), base_score("layout") + 0.05, "layout"))
        for v in nuext_ents.get("buyer_name", []):
            if _is_clean_ner(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            quote_to = find_quote_to_q(search_text)
            if quote_to:
                name = quote_to.split(",")[0].strip()
                if self._is_valid_name(name):
                    candidates.append(Candidate(_clean_value(name), base_score("regex_context") + 0.10, "regex_context"))
                break
        search_all = text + "\n" + (raw_text or "")
        if re.search(r"\bassurity\s+ltd\b", search_all, re.IGNORECASE):
            candidates.append(Candidate("Assurity Ltd", base_score("regex_context") + 0.10, "regex_context"))
        candidates = [adjust_vendor_name(c) for c in candidates]
        return pick_best(candidates)

    def _extract_buyer_address(self, text, nuext_ents, layout_fields=None, raw_text="", buyer_name=""):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        bt = layout_fields.get("bill_to", "")
        if bt and _looks_like_address(bt):
            candidates.append(Candidate(_clean_value(bt), base_score("layout") + 0.05, "layout"))
        for a in nuext_ents.get("buyer_address", []):
            if _looks_like_address(a):
                candidates.append(Candidate(_clean_value(a), base_score("nuextract"), "nuextract"))
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            bill_to = find_bill_to_q(search_text)
            if bill_to and _looks_like_address(bill_to):
                candidates.append(Candidate(_clean_value(bill_to), base_score("regex_context"), "regex_context"))
                break
        for search_text in [raw_text, text] if raw_text else [text]:
            assurity_addr = self._find_address_near_keyword(search_text, "assurity")
            if assurity_addr:
                candidates.append(Candidate(assurity_addr, base_score("heuristic") + 0.12, "heuristic"))
                break
        if buyer_name and "assurity" not in buyer_name.lower():
            first_word = buyer_name.strip().split()[0] if buyer_name.strip() else ""
            if first_word and len(first_word) >= 3:
                for search_text in [raw_text, text] if raw_text else [text]:
                    buyer_addr = self._find_address_near_keyword(search_text, first_word)
                    if buyer_addr:
                        candidates.append(Candidate(buyer_addr, base_score("heuristic") + 0.08, "heuristic"))
                        break
        vendor_addr = layout_fields.get("vendor_address", "")
        if vendor_addr and _looks_like_address(vendor_addr):
            if "assurity" in vendor_addr.lower():
                parts = [p.strip() for p in vendor_addr.split(",")]
                cleaned_parts = []
                for p in parts:
                    if re.match(r"^[\d\+\-\(\)\s]{7,}$", p):
                        continue
                    if p.startswith("+") and re.match(r"^\+[\d\s\-()]{7,20}$", p):
                        continue
                    if re.search(r"\b(ltd|llc|inc|plc|corp)\b", p, re.IGNORECASE) and "assurity" in p.lower():
                        continue
                    cleaned_parts.append(p)
                if cleaned_parts:
                    candidates.append(Candidate(", ".join(cleaned_parts), base_score("layout") + 0.10, "layout"))
        candidates = [adjust_address(c) for c in candidates]
        return pick_best(candidates)

    @staticmethod
    def _find_address_near_keyword(text, keyword):
        lines = text.split("\n")
        for i, line in enumerate(lines):
            if keyword.lower() in line.lower():
                addr_lines: list[str] = []
                for j in range(i + 1, min(i + 6, len(lines))):
                    l = lines[j].strip()
                    if not l:
                        break
                    if re.match(r"^(quote|invoice|description|item|payment|payable|bank|date|total|sub|"
                                r"terms|notes|thank|conditions|tax|vat|discount|qty|price|amount)",
                                l, re.IGNORECASE):
                        break
                    if re.match(r"^(sort\s*code|account\s*(?:no|number|name)|iban|swift|bic|"
                                r"routing\s*number|bank\s*name|vat\s*no)",
                                l, re.IGNORECASE):
                        continue
                    if re.match(r"^[\d\+\-\(\)\s]{7,}$", l):
                        continue
                    if l.startswith("+") and re.match(r"^\+[\d\s\-()]{7,20}$", l):
                        continue
                    if "@" in l:
                        continue
                    if re.match(r"^(https?://|www\.)", l, re.IGNORECASE):
                        continue
                    if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", l):
                        continue
                    if (l.isupper() and len(l) > 3
                            and not re.search(r"\d", l)
                            and not re.search(r"\b(street|road|lane|way|avenue|drive|house|floor|suite|unit)\b", l, re.IGNORECASE)):
                        continue
                    addr_lines.append(l)
                if addr_lines:
                    return ", ".join(addr_lines)
        return ""

    @staticmethod
    def _find_other_address(text, buyer_address, supplier_name):
        if not text:
            return ""
        buyer_addr_lower = buyer_address.lower() if buyer_address else ""
        buyer_tokens = set(re.findall(r"\w{3,}", buyer_addr_lower))
        _buyer_kw = {"assurity", "recipient"}
        _postcode_re = re.compile(r"[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}", re.IGNORECASE)
        _street_kw = re.compile(
            r"\b(?:street|road|lane|avenue|drive|crescent|close|"
            r"court|place|square|boulevard|terrace)\b"
            r"|\b(?:floor|suite)\s+\d"
            r"|\b(?:unit)\s+\d"
            r"|\b(?:\d+\w?\s+(?:house|park|estate|way))\b",
            re.IGNORECASE,
        )
        _reject_line = re.compile(
            r"^(qty|quantity|price|amount|unit price|cost|description|"
            r"item|product|service|subtotal|total|tax|vat|discount|"
            r"bill\s*to|billed\s*to|quote\s*to|quote\s*no|date|"
            r"intel|acer|lenov|thinkcentre|thinkpad|macbook|samsung|"
            r"ssd|ram|win\s*\d|gb\s|ghz|core\s*i)",
            re.IGNORECASE,
        )
        lines = text.split("\n")
        found_addresses: list[str] = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            if _reject_line.search(line):
                i += 1
                continue
            has_postcode = bool(_postcode_re.search(line))
            has_street = bool(_street_kw.search(line))
            is_addr_line = has_postcode or has_street
            if not is_addr_line and re.match(r"^\d+\s+[A-Z]", line):
                if not re.search(r"\b(gb|ram|ssd|intel|acer|lenov)\b", line, re.IGNORECASE):
                    is_addr_line = True
            if is_addr_line:
                addr_lines: list[str] = []
                start = max(0, i - 2)
                for j in range(start, min(i + 4, len(lines))):
                    l = lines[j].strip()
                    if not l:
                        if j > i:
                            break
                        continue
                    if re.match(r"^(quote|invoice|description|item|payment|payable|bank|"
                                r"subtotal|total|tax|notes|terms|thank|qty|price|amount)",
                                l, re.IGNORECASE):
                        if j > i:
                            break
                        else:
                            continue
                    if _reject_line.search(l):
                        continue
                    if re.match(r"^[\d\+\-\(\)\s]{7,}$", l) or l.startswith("+"):
                        continue
                    if "@" in l or re.match(r"^(https?://|www\.)", l, re.IGNORECASE):
                        continue
                    if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", l):
                        continue
                    if re.match(r"^[£$€¥₹]\s*[\d,]+\.?\d*$", l):
                        continue
                    addr_lines.append(l)
                if addr_lines:
                    non_buyer_lines = []
                    _label_re = re.compile(
                        r"^(recipient|bill\s*to|billed\s*to|quote\s*to|customer|"
                        r"sender|payable\s*to|payment|bank)\s*:?\s*$",
                        re.IGNORECASE,
                    )
                    for al in addr_lines:
                        if _label_re.match(al.strip()):
                            continue
                        al_tokens = set(re.findall(r"\w{3,}", al.lower()))
                        if not al_tokens:
                            continue
                        line_overlap = al_tokens & buyer_tokens
                        line_overlap_ratio = len(line_overlap) / len(al_tokens)
                        if line_overlap_ratio > 0.6:
                            continue
                        if any(kw in al.lower() for kw in _buyer_kw):
                            continue
                        non_buyer_lines.append(al)
                    if non_buyer_lines:
                        addr_text = ", ".join(non_buyer_lines)
                        if _postcode_re.search(addr_text) or len(non_buyer_lines) >= 2:
                            found_addresses.append(addr_text)
                i += 1
            else:
                i += 1
        if not found_addresses:
            return ""
        if supplier_name:
            supplier_lower = supplier_name.lower()
            supplier_tokens = set(re.findall(r"\w{3,}", supplier_lower))
            for addr in found_addresses:
                addr_tokens = set(re.findall(r"\w{3,}", addr.lower()))
                if supplier_tokens & addr_tokens:
                    return addr
        return found_addresses[0] if found_addresses else ""

    def _extract_quote_date(self, text, spacy_ents, nuext_ents, layout_fields=None, raw_text=""):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_date = layout_fields.get("invoice_date", "") or layout_fields.get("po_date", "")
        if layout_date and _looks_like_date(layout_date):
            candidates.append(Candidate(_clean_value(layout_date), base_score("layout"), "layout"))
        for d in nuext_ents.get("quote_date", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("nuextract"), "nuextract"))
        qd = find_quote_date_q(text)
        if qd:
            candidates.append(Candidate(qd, base_score("regex_context") + 0.10, "regex_context"))
        if raw_text and raw_text != text:
            qd_raw = find_quote_date_q(raw_text)
            if qd_raw:
                candidates.append(Candidate(qd_raw, base_score("regex_context") + 0.05, "regex_context"))
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            inv_date = find_invoice_date(search_text)
            if inv_date:
                candidates.append(Candidate(inv_date, base_score("regex_context") + 0.05, "regex_context"))
                break
        for d in spacy_ents.get("DATE", []):
            if _looks_like_date(d):
                candidates.append(Candidate(_clean_value(d), base_score("spacy"), "spacy"))
        for search_text_fb in [text, raw_text] if raw_text and raw_text != text else [text]:
            all_dates = find_all_dates(search_text_fb)
            for d in all_dates:
                candidates.append(Candidate(d, base_score("regex_fallback"), "regex_fallback"))
        candidates = [adjust_date(c) for c in candidates]
        return pick_best(candidates)

    def _extract_validity_date(self, text, nuext_ents, raw_text=""):
        candidates: list[Candidate] = []
        for v in nuext_ents.get("validity_date", []):
            if _looks_like_date(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            pay_by = find_pay_by_date_q(search_text)
            if pay_by:
                candidates.append(Candidate(pay_by, base_score("regex_context") + 0.10, "regex_context"))
                break
        for search_text in [text, raw_text] if raw_text and raw_text != text else [text]:
            due = find_due_date(search_text)
            if due:
                candidates.append(Candidate(due, base_score("regex_context"), "regex_context"))
                break
        candidates = [adjust_date(c) for c in candidates]
        return pick_best(candidates)

    def _extract_po_number(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_po = layout_fields.get("po_number", "")
        if layout_po and _has_digits(layout_po):
            candidates.append(Candidate(_clean_value(layout_po), base_score("layout"), "layout"))
        for v in nuext_ents.get("po_number", []):
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        po = find_po_number(text)
        if po:
            candidates.append(Candidate(po, base_score("regex_context"), "regex_context"))
        candidates = [adjust_invoice_number(c) for c in candidates]
        return pick_best(candidates)

    def _extract_currency(self, text, nuext_ents):
        candidates: list[Candidate] = []
        for c in nuext_ents.get("currency", []):
            code = c.strip().upper()
            if code in ("GBP", "USD", "EUR", "JPY", "INR", "CAD", "AUD"):
                candidates.append(Candidate(code, base_score("nuextract"), "nuextract"))
        detected = detect_currency(text)
        if detected:
            candidates.append(Candidate(detected, base_score("regex_context"), "regex_context"))
        return pick_best(candidates)

    def _extract_subtotal(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_sub = layout_fields.get("subtotal", "")
        if layout_sub and _looks_like_money(layout_sub):
            candidates.append(Candidate(_clean_value(layout_sub), base_score("layout"), "layout"))
        for v in nuext_ents.get("subtotal", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        sub = find_subtotal(text)
        if sub:
            candidates.append(Candidate(sub, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_discount(self, text, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_disc = layout_fields.get("discount", "")
        if layout_disc and _looks_like_money(layout_disc):
            candidates.append(Candidate(_clean_value(layout_disc), base_score("layout"), "layout"))
        for v in nuext_ents.get("discount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        disc = find_discount(text)
        if disc:
            candidates.append(Candidate(disc, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_tax(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_tax = layout_fields.get("tax", "")
        if layout_tax and _looks_like_money(layout_tax):
            candidates.append(Candidate(_clean_value(layout_tax), base_score("layout"), "layout"))
        for v in nuext_ents.get("tax_amount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        tax = find_tax(text)
        if tax:
            candidates.append(Candidate(tax, base_score("regex_context"), "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_total(self, text, spacy_ents, nuext_ents, layout_fields=None):
        candidates: list[Candidate] = []
        layout_fields = layout_fields or {}
        layout_total = layout_fields.get("total", "")
        if layout_total and _looks_like_money(layout_total):
            candidates.append(Candidate(_clean_value(layout_total), base_score("layout") + 0.10, "layout"))
        for v in nuext_ents.get("total_amount", []):
            if _looks_like_money(v):
                candidates.append(Candidate(_clean_value(v), base_score("nuextract"), "nuextract"))
        for v in spacy_ents.get("MONEY", []):
            v = v.strip()
            if _is_clean_ner(v) and _has_digits(v):
                candidates.append(Candidate(v, base_score("spacy") - 0.10, "spacy"))
        regex_total = find_total(text)
        if regex_total:
            candidates.append(Candidate(regex_total, base_score("regex_context") + 0.15, "regex_context"))
        candidates = [adjust_money(c) for c in candidates]
        return pick_best_money(candidates)

    def _extract_payment_terms(self, text, nuext_ents):
        for v in nuext_ents.get("payment_terms", []):
            if _is_clean_ner(v):
                return _clean_value(v)
        days = find_payment_terms(text)
        if days:
            return f"Net {days}"
        return ""

    def _extract_line_items(self, text, word_boxes, layout_fields=None, raw_text=""):
        layout_fields = layout_fields or {}
        scored_sets: list[tuple[float, list[dict[str, str]]]] = []
        layout_headers = layout_fields.get("_table_headers", [])
        layout_rows = layout_fields.get("_table_rows", [])
        if layout_headers and layout_rows:
            layout_items = self._layout_rows_to_items(layout_headers, layout_rows)
            layout_items = self._filter_invalid_items(layout_items)
            if layout_items:
                score = self._score_line_item_set(layout_items, base_score("layout"))
                scored_sets.append((score, layout_items))
        text_items = find_line_items_from_text(text)
        text_items = self._filter_invalid_items(text_items)
        if text_items:
            score = self._score_line_item_set(text_items, base_score("text_table"))
            scored_sets.append((score, text_items))
        if raw_text and raw_text != text:
            raw_items = find_line_items_from_text(raw_text)
            raw_items = self._filter_invalid_items(raw_items)
            if raw_items:
                score = self._score_line_item_set(raw_items, base_score("text_table") - 0.05)
                scored_sets.append((score, raw_items))
        if word_boxes:
            ocr_items = extract_table_from_boxes(word_boxes)
            ocr_items = self._filter_invalid_items(ocr_items)
            if ocr_items:
                score = self._score_line_item_set(ocr_items, base_score("ocr_table"))
                scored_sets.append((score, ocr_items))
        if not scored_sets:
            return []
        scored_sets.sort(key=lambda x: x[0], reverse=True)
        best_score, best_items = scored_sets[0]
        return self._clean_line_items(best_items)

    @staticmethod
    def _filter_invalid_items(items):
        filtered = []
        for item in items:
            desc = item.get("item", "").strip()
            if re.match(r"^[\+\d\-\(\)\s]{1,5}$", desc):
                continue
            if re.match(r"^[\d\+\-\(\)\s]{7,}$", desc):
                continue
            if "@" in desc:
                continue
            if not desc:
                continue
            filtered.append(item)
        return filtered

    @staticmethod
    def _score_line_item_set(items, base):
        if not items:
            return 0.0
        score = base
        completeness_total = 0.0
        for item in items:
            fields_filled = sum(
                1 for k in ("item", "quantity", "unit_price", "amount")
                if item.get(k, "").strip()
            )
            completeness_total += fields_filled / 4.0
            qty = item.get("quantity", "").strip()
            try:
                int(float(qty.replace(",", "")))
                score += 0.02
            except (ValueError, TypeError):
                score -= 0.03
            amt = item.get("amount", "").strip()
            if re.search(r"[£$€¥₹\u00a3]", amt):
                score += 0.02
        avg_completeness = completeness_total / len(items)
        score += avg_completeness * 0.15
        if len(items) >= 2:
            score += 0.05
        if len(items) >= 4:
            score += 0.03
        return min(score, 1.0)

    @staticmethod
    def _clean_line_items(items):
        _skip_exact = {"total", "sub-total", "subtotal", "tax", "vat", "gst",
                       "discount", "grand total", "net", "balance", "amount due"}
        _skip_starts = ("total amount", "vat ", "vat(", "sales tax",
                        "grand total", "tax ", "tax(", "tax:")
        filtered: list[dict[str, str]] = []
        for item in items:
            desc = item.get("item", "").strip()
            desc_lower = desc.lower().rstrip(":")
            if desc_lower in _skip_exact:
                continue
            if any(desc_lower.startswith(p) for p in _skip_starts):
                continue
            filtered.append(item)
        merged: list[dict[str, str]] = []
        pending_desc = ""
        for item in filtered:
            has_any_numeric = any(
                item.get(k, "").strip()
                for k in ("quantity", "unit_price", "amount")
            )
            if not has_any_numeric:
                desc = item.get("item", "").strip()
                if desc:
                    pending_desc = (pending_desc + " " + desc).strip() if pending_desc else desc
            else:
                if pending_desc:
                    item = dict(item)
                    item["item"] = (pending_desc + " " + item.get("item", "")).strip()
                    pending_desc = ""
                merged.append(item)
        if pending_desc:
            merged.append({"item": pending_desc, "quantity": "", "unit_price": "", "amount": ""})
        cleaned: list[dict[str, str]] = []
        for item in merged:
            qty = item.get("quantity", "").strip()
            qty = re.sub(r"[£$€¥₹\u00a3]", "", qty).strip()
            try:
                qty_int = int(float(qty.replace(",", "")))
            except (ValueError, TypeError):
                qty_int = None
            cleaned.append({
                "item": item.get("item", "").strip(),
                "quantity": str(qty_int) if qty_int is not None else qty,
                "unit_price": item.get("unit_price", "").strip(),
                "amount": item.get("amount", "").strip(),
            })
        return cleaned

    @staticmethod
    def _layout_rows_to_items(headers, rows):
        header_map: dict[int, str] = {}
        desc_keys = {"description", "item", "particular", "particulars",
                      "product", "service", "details", "item description",
                      "item descripsion"}
        qty_keys = {"qty", "quantity", "units", "count", "months", "hours", "days"}
        price_keys = {"price", "rate", "unit price", "cost", "monthly cost",
                      "price p/m", "unit price ($)"}
        amount_keys = {"amount", "subtotal", "sub-total", "total",
                       "line total", "extended", "total cost", "total ($)"}
        id_keys = {"item id", "id", "sku", "code"}
        for i, h in enumerate(headers):
            h_lower = h.strip().lower()
            if h_lower in desc_keys:
                header_map[i] = "item"
            elif h_lower in qty_keys:
                header_map[i] = "quantity"
            elif h_lower in price_keys:
                header_map[i] = "unit_price"
            elif h_lower in amount_keys:
                header_map[i] = "amount"
            elif h_lower in id_keys:
                header_map[i] = "item_id"
            else:
                if any(kw in h_lower for kw in desc_keys):
                    header_map[i] = "item"
                elif any(kw in h_lower for kw in qty_keys):
                    header_map[i] = "quantity"
                elif any(kw in h_lower for kw in price_keys):
                    header_map[i] = "unit_price"
                elif any(kw in h_lower for kw in amount_keys):
                    header_map[i] = "amount"
        if "item" not in header_map.values():
            if len(headers) >= 3:
                header_map = {0: "item", 1: "quantity", 2: "unit_price"}
                if len(headers) >= 4:
                    header_map[3] = "amount"
        items: list[dict[str, str]] = []
        pending_desc = ""
        for row in rows:
            item: dict[str, str] = {"item": "", "quantity": "", "unit_price": "", "amount": ""}
            for col_idx, field_name in header_map.items():
                if col_idx < len(row):
                    item[field_name] = row[col_idx].strip()
            qty_val = item.get("quantity", "")
            if qty_val:
                _qty_money = re.match(r"^(\d+)\s+([£$€¥₹]?\s*[\d,]+\.?\d*)$", qty_val)
                if _qty_money:
                    item["quantity"] = _qty_money.group(1)
                    money = _qty_money.group(2)
                    if not item.get("amount"):
                        item["amount"] = money
                    elif not item.get("unit_price"):
                        item["unit_price"] = money
            has_values = any(item[k] for k in ("quantity", "unit_price", "amount"))
            if item["item"] and not has_values:
                if pending_desc:
                    pending_desc += " " + item["item"]
                else:
                    pending_desc = item["item"]
                continue
            if item["item"] or has_values:
                if pending_desc:
                    if item["item"]:
                        item["item"] = pending_desc + " " + item["item"]
                    else:
                        item["item"] = pending_desc
                    pending_desc = ""
                if item["item"]:
                    items.append(item)
        if pending_desc:
            items.append({"item": pending_desc, "quantity": "", "unit_price": "", "amount": ""})
        return items


class ExtractionEngine_quote:
    def __init__(self) -> None:
        self._quote_extractor = QuoteExtractor()

    def run(self, text: str, word_boxes: list[dict] | None = None,
            layout_fields: dict[str, Any] | None = None) -> dict[str, Any]:
        if not text.strip():
            logger.warning("Empty text passed to ExtractionEngine_quote.")
            return {"document_type": "quote", "error": "No text could be extracted"}
        logger.info("Running spaCy NER...")
        spacy_ents = spacy_extract_entities(text)
        logger.info("Running NuExtract NER (quote)...")
        try:
            nuext_ents = nuextract_extract_entities_quote(text)
        except Exception as exc:
            logger.warning("NuExtract extraction failed: %s – continuing without it.", exc)
            nuext_ents = {}
        context: dict[str, Any] = {
            "text": text,
            "raw_text": layout_fields.get("_flat_text", text) if layout_fields else text,
            "spacy_ents": spacy_ents,
            "nuext_ents": nuext_ents,
            "word_boxes": word_boxes or [],
            "layout_fields": layout_fields or {},
        }
        logger.info("Running QuoteExtractor...")
        result = self._quote_extractor.extract(context)
        return result


class AgentController_quote:
    def __init__(self) -> None:
        self._engine = ExtractionEngine_quote()
        self._layout_analyzer = LayoutAnalyzer_quote()

    def process(self, file_path: str) -> dict[str, Any]:
        file_type = FileDetector.detect(file_path)
        logger.info("Detected file type: %s", file_type)
        text = ""
        word_boxes: list[dict] = []
        layout_fields: dict = {}
        if file_type == "pdf":
            text, word_boxes, layout_fields = self._extract_from_pdf(file_path)
        elif file_type == "docx":
            text = DOCXParser.extract_text(file_path)
        if not text.strip():
            logger.error("No text could be extracted from %s", file_path)
            return {"document_type": "quote", "error": "No text could be extracted from the document."}
        logger.info("Extracted %d characters, %d word boxes, %d layout fields.",
                     len(text), len(word_boxes), len(layout_fields))
        result = self._engine.run(text, word_boxes, layout_fields)
        result["_source_file"] = file_path
        return result

    def _extract_from_pdf(self, file_path):
        flat_text = PDFParser.extract_text(file_path)
        text = flat_text
        word_boxes: list[dict] = []
        layout_fields: dict = {}
        if len(text.strip()) >= OCR_MIN_TEXT_LENGTH:
            logger.info("Digital PDF – running block-level layout extraction (quote).")
            try:
                blocks = PDFBlockParser.extract_blocks(file_path)
                page_w = PDFBlockParser.page_width(file_path)
                layout = self._layout_analyzer.analyze(blocks, page_w)
                if layout.reconstructed_text.strip():
                    text = layout.reconstructed_text
                layout_fields = self._layout_to_fields(layout)
                layout_fields["_flat_text"] = flat_text
            except Exception as exc:
                logger.warning("Block extraction failed, using flat text: %s", exc)
            try:
                ocr_text, word_boxes = run_ocr_with_boxes(file_path)
                if ocr_text:
                    layout_fields["_ocr_text"] = ocr_text
            except Exception as exc:
                logger.warning("OCR box extraction failed (non-critical): %s", exc)
        else:
            logger.info("Scanned PDF detected – running full OCR.")
            try:
                text, word_boxes = run_ocr_with_boxes(file_path)
            except Exception as exc:
                logger.error("OCR failed: %s", exc)
                text = PDFParser.extract_text(file_path)
        return text, word_boxes, layout_fields

    @staticmethod
    def _layout_to_fields(layout) -> dict:
        fields: dict = {}
        _sublabel_re = re.compile(
            r"^(?:name|address|company|contact|phone|tel|email|fax)\s*:\s*",
            re.IGNORECASE,
        )

        def _clean_zone(zone_lines):
            cleaned = []
            for line in zone_lines:
                stripped = line.strip()
                if re.match(r"^(name|address|company|contact)\s*:?\s*$", stripped, re.IGNORECASE):
                    continue
                if re.match(r"^[:.\-]+$", stripped):
                    continue
                if re.match(r"^(quote\s*(?:no|number|#|date)|invoice\s*(?:no|number|date)|"
                            r"date\s*(?:issued)?|po\s*(?:no|number|ref)|due\s*date)\s*[:.#\s]",
                            stripped, re.IGNORECASE):
                    continue
                if re.match(r"^(description\s*(?:of\s*)?(?:services?)?|"
                            r"provision\s+of|scope\s+of|terms|notes)\s*:?\s*$",
                            stripped, re.IGNORECASE):
                    break
                word_count = len(stripped.split())
                _sentence_words = {"the", "will", "can", "for", "over", "including",
                                   "provision", "freelance", "graphic", "marketing",
                                   "campaign", "branding", "strategy", "ongoing",
                                   "content", "creation", "design", "support"}
                lower = stripped.lower()
                if word_count >= 5 and any(w in lower for w in _sentence_words):
                    continue
                stripped = _sublabel_re.sub("", stripped).strip()
                if stripped:
                    cleaned.append(stripped)
            return cleaned

        if layout.vendor_zone:
            fields["vendor_name"] = layout.vendor_zone[0]
            if len(layout.vendor_zone) > 1:
                addr_lines = []
                for line in layout.vendor_zone[1:]:
                    stripped = line.strip()
                    if re.match(r"^(https?://|www\.)", stripped, re.IGNORECASE):
                        continue
                    if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", stripped):
                        continue
                    if "@" in stripped:
                        continue
                    if re.match(r"^[\d\+\-\(\)\s]{7,}$", stripped):
                        continue
                    if stripped.startswith("+") and re.match(r"^\+[\d\s\-()]{7,}$", stripped):
                        continue
                    if re.match(r"^(recipient|bill\s*to|billed\s*to|quote\s*to|customer)", stripped, re.IGNORECASE):
                        break
                    if (stripped.isupper() and len(stripped) > 3
                            and not re.search(r"\d", stripped)
                            and not re.search(r"\b(street|road|lane|way|avenue|drive|house|floor|suite|unit)\b", stripped, re.IGNORECASE)):
                        continue
                    if "assurity" in stripped.lower():
                        break
                    lower = stripped.lower()
                    word_count = len(stripped.split())
                    _sentence_words = {"the", "will", "can", "please", "provide",
                                       "includes", "including", "consists", "listed",
                                       "purchased", "beginning", "process", "request",
                                       "revision", "revisions", "payment", "thank"}
                    if word_count >= 5 and any(w in lower for w in _sentence_words):
                        continue
                    if stripped.endswith(".") and word_count >= 4:
                        continue
                    addr_lines.append(stripped)
                if addr_lines:
                    fields["vendor_address"] = ", ".join(addr_lines)
        if layout.bill_to_zone:
            clean_bt = _clean_zone(layout.bill_to_zone)
            fields["bill_to"] = ", ".join(clean_bt)
        fields.update(layout.header_fields)
        if layout.totals:
            for key, val in layout.totals.items():
                fields[key] = val
        if layout.table_headers and layout.table_rows:
            fields["_table_headers"] = layout.table_headers
            fields["_table_rows"] = layout.table_rows
        if layout.payment_zone:
            fields["_payment_zone"] = layout.payment_zone
        return fields


def _compute_tax_percent_quote(subtotal: str, tax: str, discount: str = "") -> str:
    try:
        sub_val = float(_parse_numeric(subtotal))
        tax_val = float(_parse_numeric(tax))
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        taxable = sub_val - disc_val
        if taxable > 0:
            pct = (tax_val / taxable) * 100
            return f"{pct:.1f}"
    except (ValueError, ZeroDivisionError):
        pass
    return ""


def _parse_date_quote(date_str: str) -> datetime | None:
    if not date_str:
        return None
    cleaned = re.sub(r"(\d{1,2})(?:st|nd|rd|th)\b", r"\1", date_str)
    cleaned = re.sub(r"\bSept\b", "Sep", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(
        r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
        r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.",
        r"\1", cleaned, flags=re.IGNORECASE,
    )
    cleaned = " ".join(cleaned.split()).strip()
    for fmt in _DATE_FORMATS:
        try:
            return datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
    m = re.match(r"(\d{1,2})\s+(\w+)\s+(\d{4})", cleaned)
    if m:
        day, month_str, year_str = int(m.group(1)), m.group(2), int(m.group(3))
        for mfmt in ("%B", "%b"):
            try:
                month_num = datetime.strptime(month_str, mfmt).month
                max_day = calendar.monthrange(year_str, month_num)[1]
                clamped_day = min(day, max_day)
                return datetime(year_str, month_num, clamped_day)
            except ValueError:
                continue
    return None


def _normalize_date_quote(date_str: str) -> str:
    parsed = _parse_date_quote(date_str)
    if parsed is None:
        return ""
    return parsed.strftime("%d/%m/%Y")


def _extract_address_parts_quote(address: str) -> dict[str, str]:
    parts = {"line1": "", "line2": "", "city": "", "postal_code": "", "country": ""}
    if not address:
        return parts
    _non_addr = re.compile(
        r"^(recipient|supplier|vendor|bill\s*to|quote\s*to)\s*:?\s*$",
        re.IGNORECASE,
    )
    segments = [s.strip() for s in address.split(",")
                if s.strip() and not _non_addr.match(s.strip())
                and not re.match(r"^[\d\+\-\(\)\s]{7,}$", s.strip())
                and "@" not in s]
    if len(segments) >= 2:
        parts["line1"] = segments[0]
        parts["line2"] = ", ".join(segments[1:])
    elif len(segments) == 1:
        parts["line1"] = segments[0]
    postcode_match = re.search(r"\b([A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2})\b", address, re.IGNORECASE)
    if postcode_match:
        parts["postal_code"] = postcode_match.group(1).upper()
    if not parts["postal_code"]:
        zip_match = re.search(r"\b(\d{5}(?:-\d{4})?)\b", address)
        if zip_match:
            parts["postal_code"] = zip_match.group(1)
    address_lower = address.lower()
    if "united kingdom" in address_lower or "uk" in address_lower:
        parts["country"] = "United Kingdom"
    elif "united states" in address_lower or "usa" in address_lower:
        parts["country"] = "United States"
    elif postcode_match:
        parts["country"] = "United Kingdom"
    city_match = re.search(
        r"\b(London|Birmingham|Manchester|Leeds|Liverpool|Bristol|Brighton|"
        r"Newport|Horsham|Sheffield|Edinburgh|Glasgow|Cardiff|Belfast|"
        r"York|Oxford|Cambridge|Nottingham|Leicester|Southampton|"
        r"New York|Los Angeles|Chicago|Houston|Phoenix)\b",
        address, re.IGNORECASE,
    )
    if city_match:
        parts["city"] = city_match.group(1).title()
    return parts


def _detect_region_quote(address: str) -> str:
    normalized = re.sub(r"[\n,]+", " ", address)
    normalized = " ".join(normalized.split())
    region_match = re.search(
        r"\b(West Sussex|East Sussex|West Yorkshire|South Yorkshire|North Yorkshire|"
        r"West Midlands|East Midlands|Greater London|Surrey|Kent|Essex|Hampshire|"
        r"Devon|Cornwall|Somerset|Norfolk|Suffolk|Berkshire|Oxfordshire|"
        r"Cambridgeshire|Hertfordshire|Lancashire|Cheshire|Wales|Scotland|"
        r"London)\b",
        normalized, re.IGNORECASE,
    )
    if region_match:
        return region_match.group(1).title()
    return ""


def map_quote(result: dict[str, Any]) -> dict[str, str]:
    subtotal = result.get("subtotal", "")
    discount = result.get("discount", "")
    tax_raw = result.get("tax", "")
    total_raw = result.get("total_amount", "")
    if not subtotal and total_raw and tax_raw:
        try:
            total_val = float(_parse_numeric(total_raw))
            tax_val = float(_parse_numeric(tax_raw))
            derived = total_val - tax_val
            if derived > 0:
                subtotal = f"{derived:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    if not subtotal:
        items = result.get("line_items", [])
        line_sum = 0.0
        for item in items:
            amt = item.get("amount", "")
            if amt:
                try:
                    line_sum += float(_parse_numeric(amt))
                except (ValueError, TypeError):
                    pass
        if line_sum > 0:
            subtotal = f"{line_sum:.2f}"
    tax = _validate_tax(subtotal, discount, tax_raw, total_raw)
    tax_pct = _compute_tax_percent_quote(subtotal, tax, discount)
    try:
        sub_val = float(_parse_numeric(subtotal)) if subtotal else 0.0
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        tax_val = float(tax) if tax else 0.0
        computed_total = (sub_val - disc_val) + tax_val
        final_total = f"{computed_total:.2f}"
    except (ValueError, ZeroDivisionError):
        final_total = _parse_numeric(total_raw)
    effective_subtotal = subtotal
    if discount and subtotal:
        try:
            sub_val = float(_parse_numeric(subtotal))
            disc_val = float(_parse_numeric(discount))
            effective_subtotal = f"{sub_val - disc_val:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    quote_date_raw = result.get("quote_date", "")
    quote_date = _normalize_date_quote(quote_date_raw)
    validity_date = ""
    if quote_date_raw:
        parsed_qd = _parse_date_quote(quote_date_raw)
        if parsed_qd:
            validity_date = (parsed_qd + timedelta(days=7)).strftime("%d/%m/%Y")
    currency = result.get("currency", "")
    supplier_name = resolve_supplier_name(
        result.get("supplier_name", ""),
        result.get("_raw_text", ""),
    )
    supplier_address = result.get("supplier_address", "") if supplier_name else ""
    buyer_address = result.get("buyer_address", "")
    if buyer_address:
        parts = [p.strip() for p in buyer_address.split(",")]
        cleaned = []
        supplier_name_raw = result.get("supplier_name", "").lower()
        supplier_addr_raw = result.get("supplier_address", "").lower()
        for p in parts:
            pl = p.lower().strip()
            if not pl:
                continue
            if re.match(r"^[\d\+\-\(\)\s]{7,}$", p.strip()):
                continue
            if p.strip().startswith("+") and re.match(r"^\+[\d\s\-()]{7,20}$", p.strip()):
                continue
            if "@" in p:
                continue
            if re.match(r"^(https?://|www\.)", p.strip(), re.IGNORECASE):
                continue
            if re.match(r"^[\w\-]+\.\w{2,3}(\.\w{2,3})?$", p.strip()):
                continue
            if re.match(r"^(payable\s*to|billing\s*to|inv\s*no)\s*:?", pl):
                continue
            if pl in ("service", "services"):
                continue
            if supplier_name_raw and pl == supplier_name_raw:
                continue
            if supplier_name and pl == supplier_name.lower():
                continue
            if supplier_addr_raw and len(pl) > 3:
                if pl in supplier_addr_raw and not re.search(r"\b(horsham|redkiln|assurity|west\s*sussex|rh13)\b", pl):
                    continue
            cleaned.append(p.strip())
        buyer_address = ", ".join(cleaned)
    all_addresses = " ".join(filter(None, [supplier_address, buyer_address]))
    raw_text = result.get("_raw_text", "")
    addr_to_check = supplier_address or buyer_address
    addr_parts = _extract_address_parts_quote(addr_to_check)
    region = _detect_region_quote(all_addresses)
    if not addr_parts["country"]:
        raw_lower = raw_text.lower() if raw_text else ""
        combined_lower = all_addresses.lower()
        if "united kingdom" in raw_lower or "united kingdom" in combined_lower:
            addr_parts["country"] = "United Kingdom"
        elif "united states" in raw_lower or "united states" in combined_lower:
            addr_parts["country"] = "United States"
        elif re.search(r"[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}", all_addresses, re.IGNORECASE):
            addr_parts["country"] = "United Kingdom"
    return {
        "quote_id": result.get("quote_number", ""),
        "deal_id": "",
        "supplier_id": supplier_name,
        "buyer_id": result.get("buyer_name", ""),
        "supplier_address": supplier_address,
        "buyer_address": buyer_address,
        "quote_date": quote_date,
        "validity_date": validity_date,
        "currency": currency,
        "total_amount": _parse_numeric(effective_subtotal),
        "tax_percent": tax_pct,
        "tax_amount": _parse_numeric(tax),
        "total_amount_incl_tax": final_total,
        "po_id": result.get("po_number", ""),
        "country": addr_parts["country"],
        "region": region,
        "ai_flag_required": "",
        "trigger_type": "",
        "trigger_context_description": "",
        "created_date": "",
        "created_by": "",
        "last_modified_by": "",
        "last_modified_date": "",
    }


def map_quote_line_items(result: dict[str, Any]) -> list[dict[str, str]]:
    quote_id = result.get("quote_number", "")
    currency = result.get("currency", "")
    items = result.get("line_items", [])
    subtotal = result.get("subtotal", "")
    discount = result.get("discount", "")
    tax_raw = result.get("tax", "")
    total_raw = result.get("total_amount", "")
    if not subtotal and total_raw and tax_raw:
        try:
            total_val = float(_parse_numeric(total_raw))
            tax_val = float(_parse_numeric(tax_raw))
            derived = total_val - tax_val
            if derived > 0:
                subtotal = f"{derived:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
    if not subtotal:
        line_sum = 0.0
        for item in items:
            amt = item.get("amount", "")
            if amt:
                try:
                    line_sum += float(_parse_numeric(amt))
                except (ValueError, TypeError):
                    pass
        if line_sum > 0:
            subtotal = f"{line_sum:.2f}"
    tax = _validate_tax(subtotal, discount, tax_raw, total_raw)
    tax_pct = _compute_tax_percent_quote(subtotal, tax, discount)
    try:
        sub_val = float(_parse_numeric(subtotal))
        disc_val = float(_parse_numeric(discount)) if discount else 0.0
        discount_ratio = disc_val / sub_val if sub_val > 0 else 0.0
    except (ValueError, ZeroDivisionError):
        discount_ratio = 0.0
    mapped: list[dict[str, str]] = []
    for i, item in enumerate(items, start=1):
        unit_price = item.get("unit_price", "")
        quantity = item.get("quantity", "")
        line_total = item.get("amount", "")
        if not line_total and len(items) == 1 and subtotal:
            line_total = subtotal
        if not line_total and unit_price and not quantity:
            line_total = unit_price
        line_after_disc_str = ""
        line_tax = ""
        total_with_tax = ""
        try:
            lt_val = float(_parse_numeric(line_total))
            line_after_disc = lt_val * (1 - discount_ratio)
            line_after_disc_str = f"{line_after_disc:.2f}"
            if tax_pct:
                pct = float(tax_pct) / 100
                line_tax_val = line_after_disc * pct
                line_tax = f"{line_tax_val:.2f}"
                total_with_tax = f"{line_after_disc + line_tax_val:.2f}"
            else:
                total_with_tax = f"{line_after_disc:.2f}"
        except (ValueError, ZeroDivisionError):
            pass
        mapped.append({
            "quote_line_id": "",
            "quote_id": quote_id,
            "line_number": str(i),
            "item_id": "",
            "item_description": item.get("item", ""),
            "quantity": quantity,
            "unit_of_measure": "",
            "unit_price": _parse_numeric(unit_price),
            "line_total": line_after_disc_str or _parse_numeric(line_total),
            "tax_percent": tax_pct,
            "tax_amount": line_tax,
            "total_amount": total_with_tax,
            "currency": currency,
            "created_date": "",
            "created_by": "",
            "last_modified_by": "",
            "last_modified_date": "",
        })
    return mapped


def detect_document_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    text = ""

    if ext == ".pdf":
        text = PDFParser.extract_text(file_path)
        if len(text.strip()) < OCR_MIN_TEXT_LENGTH:
            text = run_ocr(file_path)
    elif ext == ".docx":
        text = DOCXParser.extract_text(file_path)

    text_lower = text.lower()

    quote_score = 0
    if re.search(r"(?:^|\n)\s*(?:price\s+)?quot(?:e|ation)\s*(?:\n|$)", text_lower):
        quote_score += 3
    if re.search(r"\bquot(?:e|ation)\s*(?:no|number|#|num|ref)\b", text_lower):
        quote_score += 2
    if re.search(r"\bquot(?:e|ation)\s*date\b", text_lower):
        quote_score += 1
    if re.search(r"\bvalid(?:ity)?\s*(?:date|until|till)\b", text_lower):
        quote_score += 2
    if re.search(r"\bquot(?:e|ation)\s*to\b", text_lower):
        quote_score += 1
    if re.search(r"\bqut[\s\-]?\d", text_lower):
        quote_score += 1

    po_score = 0
    if re.search(r"(?:^|\n)\s*purchase\s+order\s*(?:\n|$)", text_lower):
        po_score += 3
    if re.search(r"\bpo\s*date\b", text_lower):
        po_score += 2
    if re.search(r"\border\s*date\b", text_lower):
        po_score += 1
    if re.search(r"\brecipient\s*:", text_lower):
        po_score += 2
    if re.search(r"\bquote\s*(?:number|no|#|ref)", text_lower):
        po_score += 1

    invoice_score = 0
    if re.search(r"(?:^|\n)\s*(?:tax\s+)?invoice\s*(?:\n|$)", text_lower):
        invoice_score += 3
    if re.search(r"\binvoice\s*(?:no|number|#|num)\b", text_lower):
        invoice_score += 2
    if re.search(r"\binvoice\s*date\b", text_lower):
        invoice_score += 1
    if re.search(r"\bdate\s*issued\b", text_lower):
        invoice_score += 1
    if re.search(r"\bbill(?:ed|ing)?\s*to\b", text_lower):
        invoice_score += 1
    if re.search(r"\bpayable\s*to\b", text_lower):
        invoice_score += 1
    if re.search(r"\binv[\s\-.]?\d", text_lower):
        invoice_score += 1

    best = max(quote_score, po_score, invoice_score)
    if best == quote_score and quote_score > po_score and quote_score > invoice_score:
        return "quote"
    if best == po_score and po_score > invoice_score:
        return "po"
    return "invoice"


def run_data_extraction(file_path: str):
    doc_type = detect_document_type(file_path)
    logger.info("Detected document type: %s", doc_type)

    if doc_type == "invoice":
        controller = AgentController_invoice()
        result = controller.process(file_path)
        invoice_data = map_invoice(result)
        line_items_data = map_line_items_invoice(result)
        return {
            "document_type": "invoice",
            "invoice_data": invoice_data,
            "line_items": line_items_data,
        }
    elif doc_type == "quote":
        controller = AgentController_quote()
        result = controller.process(file_path)
        quote_data = map_quote(result)
        line_items_data = map_quote_line_items(result)
        return {
            "document_type": "quote",
            "quote_data": quote_data,
            "line_items": line_items_data,
        }
    else:
        controller = AgentController_po()
        result = controller.process(file_path)
        po_data = map_purchase_order(result)
        line_items_data = map_po_line_items(result)
        return {
            "document_type": "po",
            "po_data": po_data,
            "line_items": line_items_data,
        }


def _setup_logging(verbose: bool = False) -> None:
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
        datefmt="%H:%M:%S",
    )


def _collect_files(path: str) -> list[str]:
    path = os.path.abspath(path)

    if os.path.isfile(path):
        return [path]

    if os.path.isdir(path):
        files: list[str] = []
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(glob.glob(os.path.join(path, f"*{ext}")))
        files.sort()
        return files

    return []


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract structured data from invoice, purchase order, or quote documents (PDF / DOCX)."
    )
    parser.add_argument("path", help="Path to a document file or folder.")
    parser.add_argument("--output-dir", "-o", default=None, help="Output directory.")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable debug logging.")
    args = parser.parse_args()

    _setup_logging(args.verbose)
    main_logger = logging.getLogger("final_data_extraction_agent")

    files = _collect_files(args.path)
    if not files:
        main_logger.error("No supported files found at: %s", args.path)
        sys.exit(1)

    main_logger.info("Found %d file(s) to process.", len(files))

    output_dir = args.output_dir or os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "output",
    )
    os.makedirs(output_dir, exist_ok=True)

    controller_invoice = AgentController_invoice()
    controller_po = AgentController_po()
    controller_quote = AgentController_quote()
    success = 0
    failed = 0

    for file_path in files:
        base = os.path.splitext(os.path.basename(file_path))[0]
        main_logger.info("Processing: %s", file_path)
        try:
            doc_type = detect_document_type(file_path)
            main_logger.info("Detected as: %s", doc_type)

            if doc_type == "invoice":
                result = controller_invoice.process(file_path)
                invoice_data = map_invoice(result)
                line_items_data = map_line_items_invoice(result)
                inv_dir = os.path.join(output_dir, "invoice")
                os.makedirs(inv_dir, exist_ok=True)
                JSONWriter.to_file(invoice_data, os.path.join(inv_dir, f"{base}_invoice.json"))
                JSONWriter.to_file(line_items_data, os.path.join(inv_dir, f"{base}_line_items.json"))

            elif doc_type == "quote":
                result = controller_quote.process(file_path)
                quote_data = map_quote(result)
                line_items_data = map_quote_line_items(result)
                quote_dir = os.path.join(output_dir, "quotes")
                os.makedirs(quote_dir, exist_ok=True)
                JSONWriter.to_file(quote_data, os.path.join(quote_dir, f"{base}_quote.json"))
                JSONWriter.to_file(line_items_data, os.path.join(quote_dir, f"{base}_line_items.json"))

            else:
                result = controller_po.process(file_path)
                po_data = map_purchase_order(result)
                line_items_data = map_po_line_items(result)
                po_dir = os.path.join(output_dir, "purchase_order")
                os.makedirs(po_dir, exist_ok=True)
                JSONWriter.to_file(po_data, os.path.join(po_dir, f"{base}_purchase_order.json"))
                JSONWriter.to_file(line_items_data, os.path.join(po_dir, f"{base}_line_items.json"))

            main_logger.info("Done: %s", base)
            success += 1
        except Exception as exc:
            main_logger.error("Failed: %s – %s", base, exc)
            failed += 1

    main_logger.info("Complete. %d succeeded, %d failed, %d total.", success, failed, len(files))


if __name__ == "__main__":
    main()
